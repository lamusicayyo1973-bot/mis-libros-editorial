import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_oni_vol2_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-2"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 2: El Examen de la Montaña Sombría", 2, 1, "subtitulo_vol2")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: La Cuvo-Espada de la Noche", "cap_1"),
        ("Capítulo 2: El Pueblo de los Faroles", "cap_2"),
        ("Capítulo 3: La Danza de la Tinta Roja", "cap_3"),
        ("Capítulo 4: El Tren de las Sombras", "cap_4"),
        ("Capítulo 5: La Trampa de los Sueños (Clímax del Volumen 2)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: La Cuvo-Espada de la Noche",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: El Regreso de la Selección",
                    "texto": """El amanecer iluminó los picos del monte Fujikane. De los más de sesenta aspirantes que ingresaron a la prueba, solo cuatro permanecieron de pie en la salida rodeada de glicinas púrpuras: Ren, el joven de las dos katanas cortas, la chica silenciosa de la máscara de zorro y un muchacho asustadizo que no paraba de temblar.

Las dos niñas con rostros de porcelana los recibieron con una reverencia formal.

—Felicitaciones a los sobrevivientes —dijeron al unísono—. Ahora son oficialmente Cazadores de Rango Mizunoto, el escalafón inicial del Gremio Cuervo.

Dos asistentes del gremio se acercaron portando una mesa de madera. Sobre ella descansaban varios bloques de mineral Tamahagane-Sombra, el metal especial que absorbe la luz del sol.

—Elijan el bloque que moldeará su katana definitiva —indicó una de las niñas—. Cada hoja responderá al color del alma de su portador.

Ren extendió su mano y eligió un bloque oscuro que emitía un calor sutil, idéntico al mineral de la antigua herrería de su padre.""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Cuervo Mensajero",
                    "texto": """Un graznido agudo resonó sobre el cielo de la mañana. Un cuervo de plumaje negro azabache, portando una pequeña placa de bronce en el pecho, descendió y se posó sobre el hombro derecho de Ren.

—¡Ren Hagane! ¡Ren Hagane! —graznó la ave con una voz metálica e inesperadamente humana—. ¡Tu primera misión te aguarda! ¡Misión de reconocimiento en el Pueblo de los Faroles! ¡Jóvenes están desapareciendo cada noche en la penumbra!

Ren ajustó la caja de madera de su espalda, donde Miyuki descansaba del sol de la mañana.

—Entendido —dijo Ren, apretando el mango de su katana provisoria—. Nos ponemos en marcha inmediatamente.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Entrega del Acero Negro",
                    "texto": """Antes de partir al Pueblo de los Faroles, un herrero excéntrico del gremio —cubierto con un sombrero de paja y una máscara de Hyottoko— alcanzó a Ren en el camino de la montaña.

Llevaba un envoltorio de tela roja. Al desplegarlo, reveló la nueva katana de Ren: la hoja era de un negro azabache profundo, pero a lo largo del filo corría una fina línea carmesí que brillaba al contacto con la respiración de Ren.

—Es una Hoja de Tinta y Fuego —explicó el herrero, señalando el metal—. Absorbió el fragmento de cristal de tu antigua arma rota. Si tu espíritu flaquea, la hoja te consumirá. Si tu espíritu es fuerte, cortará el acero de los demonios más antiguos.

Ren desenvainó la espada. El peso era perfecto; se sentía como una extensión natural de su propio brazo.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: El Pueblo de los Faroles",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: Las Sombras del Mercado",
                    "texto": """El Pueblo de los Faroles era una villa próspera conocida por sus festivales nocturnos. Cientos de faroles de papel rojo y amarillo iluminaban las calles de piedra mientras la gente caminaba vistiendo yukatas de colores.

Pero tras la fachada festiva, Ren sintió el olor fétido a azufre y sangre rancia flotando en el aire.

—Las desapariciones ocurren siempre cuando la última vela del festival se apaga —murmuró Ren, caminando por un callejón oscuro alejado de la multitud.

De pronto, la caja de madera en su espalda comenzó a golpear suavemente. Miyuki estaba despierta y alertando a su hermano sobre una amenaza inminente desde los tejados.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Demonio de la Cera",
                    "texto": """Desde la sombra de una vivienda, la pared de madera se derritió como si fuera cera caliente. Una criatura de tres metros hecha de sustancia viscosa y brillante emergió, atrapando a un joven poblador con sus tentáculos de cera hirviente.

—¡Un nuevo cazador de carne tierna! —siseó el demonio, mientras su rostro sin ojos se abría revelando hileras de dientes de cristal.

Ren no dudó. Desenvainó su katana negra y se lanzó al ataque. La cera del demonio intentó envolver su hoja, pero al contacto con la Respiración de Sangre, el calor de la katana evaporó la sustancia al instante.

Sin embargo, la criatura se regeneraba a una velocidad pasmosa: cada corte hecho por Ren se cerraba en una fracción de segundo.

—¡Las espadas comunes no pueden cortar lo que se derrite! —se burló el demonio, atrapando el tobillo de Ren.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Despertar de Miyuki",
                    "texto": """Antes de que el demonio pudiera aplastar a Ren contra el suelo, la tapa de la caja de madera voló por los aires.

Miyuki saltó al combate. Con el bambú aún en la boca, dio una patada voladora con una fuerza descomunal que destrozó el torso de cera del demonio, arrojándolo diez metros hacia atrás. La sangre de la joven, al caer sobre la cera, comenzó a arder en llamas púrpuras que impidieron la regeneración de la monstruosidad.

—¡Sangre de Demonio Explosiva! —comprendió Ren al ver la habilidad de su hermana.

Miyuki miró a Ren con sus ojos de pupila rasgada, dándole una señal con la cabeza.

Hermano y hermana estaban luchando juntos por primera vez en perfecta sincronía.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: La Danza de la Tinta Roja",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: La Purificación del Fuego",
                    "texto": """Aprovechando que las llamas púrpuras de Miyuki congelaron la regeneración del monstruo, Ren inhaló una bocanada profunda de aire, llenando sus pulmones hasta el límite.

Las venas negras de su rostro brillaron con intensidad.

—Estilo de la Sangre Negra... Tercera Postura: Espiral del Horno Encendido.

Ren giró sobre su propio eje a una velocidad cegadora. Su katana trazó una espiral de fuego rojo y sombras que cortó al demonio en tres partes principales, alcanzando el núcleo de su cuello.

La criatura emitió un alarido sónico antes de disolverse en cenizas de fuego rojo que el viento de la noche dispersó sobre los tejados.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: Los Recuerdos de la Cera",
                    "texto": """Antes de que las cenizas del demonio desaparecieran por completo, la Respiración de Sangre de Ren le permitió percibir los últimos pensamientos de la criatura.

Vio la visión de un anciano fabricante de velas de la era Edo, abandonado por su familia en la pobreza y tentado por un hombre elegante que vestía un sombrero blanco de época y trajo la maldición sobre su cuerpo.

—Muzan... —murmuró Ren, escuchando el nombre del Rey Oni por primera vez en la mente de la víctima.

Miyuki se acercó despacio a las cenizas y colocó su mano sobre el lugar donde cayó el demonio, demostrando compasión por el alma que alguna vez fue humana.

Ren guardó su katana en la funda. La primera misión había concluido, pero la verdadera cacería recién comenzaba.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Sombra en la Niebla",
                    "texto": """Desde lo alto de la torre del reloj del pueblo, una figura envuelta en una túnica oscura con bordados de telaraña observaba la escena.

Llevaba un biwa (instrumento musical japonés) en sus manos y la pupila de su ojo izquierdo estaba marcada con el número del Sexto Lunar Rojo.

—El chico de las venas negras y la niña que no consume carne... —susurró la entidad, tocando una cuerda del biwa—. El señor Estará muy complacido de saber que la estirpe sobrevivió.

Con un sonido seco de cuerda (¡TONG!), la figura desapareció en el aire, dejando solo un rastro de humo negro.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: El Tren de las Sombras",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Cuervo de la Segunda Misión",
                    "texto": """Al amanecer del día siguiente, el cuervo mensajero de Ren descendió agitado sobre la barandilla de la posada.

—¡Nuevas órdenes! ¡Nuevas órdenes! —gritó la ave—. ¡Abordar el Tren de la Noche en la estación central! ¡Más de cuarenta pasajeros han desaparecido dentro de los vagones! ¡Un Pilar del Gremio ya está a bordo!

Ren miró a Miyuki, quien descansaba dentro de la caja de madera.

—Un Pilar... —dijo Ren—. Los guerreros más poderosos del Gremio Cuervo. Esta es nuestra oportunidad de aprender cómo derrotar a los Seis Lunares.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Estación del Vapor",
                    "texto": """La estación de trenes de la ciudad estaba cubierta por una densa nube de vapor blanco producido por la gigantesca locomotora de hierro negro.

Ren caminaba por el andén entre la multitud, impresionado por la tecnología de la máquina de vapor. Al subir al vagón de pasajeros, el aire dentro del tren se sentía anormalmente pesado, como si el metal de la locomotora estuviera vivo.

En el asiento central del vagón de primera clase, un hombre de hombros anchos vestía un haori con patrones de llamas doradas. Comía bento de carne a una velocidad pasmosa, gritando "¡Sabroso!" con cada bocado.

Era Kenshin, el Pilar del Fuego.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Bienvenida del Pilar",
                    "texto": """Ren se acercó con cautela al asiento de Kenshin.

—Disculpe... ¿es usted el Pilar del Fuego? —preguntó Ren inclinándose levemente.

Kenshin se detuvo en seco, clavando sus ojos dorados y brillantes sobre Ren. En un instante, su mirada se volvió seria al notar la caja de madera en la espalda del joven y las marcas de venas en su rostro.

—Joven del acero negro... —dijo Kenshin con un tono de voz retumbante que hizo vibrar las ventanas del vagón—. Siento la presencia de un demonio a tu espalda, y sin embargo, tu espíritu no emite sed de sangre. Siéntate. La noche será larga y el enemigo ya está entre nosotros.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: La Trampa de los Sueños (Clímax del Volumen 2)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: El Revisor del Tren",
                    "texto": """Un hombre pálido con uniforme de revisor avanzó por el pasillo del vagón, picando los boletos de los pasajeros con una herramienta de metal.

Cuando picó el boleto de Ren y Kenshin, un sonido seco resonó en el aire. Sin darse cuenta, una niebla invisible de origen demoníaco envolvió las mentes de todos los presentes.

Ren sintió que sus párpados se volvían extremadamente pesados. Cayó sentado en su asiento, incapaz de mantener los ojos abiertos.

El boleto no era de papel común: era un arte demoníaco de ilusión que atrapaba a las víctimas dentro de sus recuerdos más profundos.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Sueño de la Herrería",
                    "texto": """Dentro de su mente, Ren despertó en el taller de su infancia. Su padre estaba vivo frente al horno de la herrería y su hermana Miyuki cantaba felizmente sin los cuernos ni la marca de demonio.

El calor del hogar se sentía tan real y reconfortante que Ren sintió el deseo de quedarse en ese lugar para siempre.

Pero al mirarse las manos en el reflejo del agua, notó la katana de cristal negro en su cintura.

—Esto no es real... —susurró Ren con dolor en el pecho—. Mi familia ya no está. No puedo quedarme en una mentira mientras Miyuki me necesita afuera.

Para romper la ilusión del demonio dentro del sueño, Ren tomó su katana y ejecutó el acto más difícil: cortar su propio cuello dentro de la visión.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Fusión de la Locomotora (Cierre del Tomo 2)",
                    "texto": """Ren abrió los ojos de golpe en el mundo real, jadeando por aire.

A su lado, Kenshin luchaba contra su propia ilusión mientras su cuerpo emitía chispas de fuego dorado. El tren completo comenzó a retorcerse: las paredes de metal y los asientos se transformaron en carne, venas y boca humana.

El demonio no estaba escondido dentro del tren: el demonio se había fusionado con la locomotora entera.

Desde el techo del primer vagón, el Sexto Lunar Rojo —un demonio de piel pálida con bocas en las palmas de las manos— reía con locura mientras el tren avanzaba a toda velocidad hacia un barranco destruido.

—¡Despierta, Pilar del Fuego! —gritó Ren, desenvainando su katana negra envuelta en llamas rojas—. ¡Si el tren se cae, nadie sobrevivirá!

[ CONTINUARÁ EN EL VOLUMEN 3 ]""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Vol 2 at {output_docx}")

if __name__ == "__main__":
    create_oni_vol2_manuscript()
