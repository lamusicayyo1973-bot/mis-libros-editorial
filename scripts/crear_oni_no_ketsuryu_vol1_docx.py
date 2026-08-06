import os
import glob
import shutil
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

# Copy generated images to target locations
src_dir = r"C:\Users\nicol\.gemini\antigravity\brain\6adf8ce5-9839-4292-a8a1-57beed4c3827"
target_root = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-1"
target_sys = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\oni-no-ketsuryu-volumen-1"

os.makedirs(target_root, exist_ok=True)
os.makedirs(target_sys, exist_ok=True)

mapping = {
    'oni_vol1_portada*.jpg': 'portada.jpg',
    'oni_vol1_thumbnail*.jpg': 'thumbnail.jpg',
    'oni_vol1_banner*.jpg': 'banner.jpg',
    'oni_vol1_escena_pacto*.jpg': 'escena_1.jpg',
    'oni_vol1_escena_climax*.jpg': 'escena_climax.jpg'
}

for pattern, target_name in mapping.items():
    matches = glob.glob(os.path.join(src_dir, pattern))
    if matches:
        matches.sort(key=os.path.getmtime, reverse=True)
        latest = matches[0]
        shutil.copy2(latest, os.path.join(target_root, target_name))
        shutil.copy2(latest, os.path.join(target_sys, target_name))
        print(f"Copied {target_name}")

def create_oni_vol1_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(220, 38, 38)
    
    run_sub = title_p.add_run("Volumen 1: La Noche de las Hojas Rotas\n\nMANGA / LIGHT NOVEL • FANTASÍA OSCURA & ACCIÓN SENGOKU\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(185, 28, 28)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    img_cover_path = os.path.join(target_root, "portada.jpg")
    if os.path.exists(img_cover_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_cover_path, width=Inches(4.5))
        
    doc.add_page_break()
    
    # TOC Header
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = toc_p.add_run("TABLA DE CONTENIDOS\n")
    r_toc.font.name = "Arial"
    r_toc.font.size = Pt(20)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(220, 38, 38)
    
    chapters_toc = [
        ("Capítulo 1", "El Olor a Nieve y Sangre", "chap1"),
        ("Capítulo 2", "La Hermana de las Sombras", "chap2"),
        ("Capítulo 3", "La Prueba del Cazador", "chap3"),
        ("Capítulo 4", "El Arte de la Hoja Maldita", "chap4"),
        ("Capítulo 5", "La Selección Final (Clímax del Volumen 1)", "chap5")
    ]
    
    for c_num, c_title, c_bkm in chapters_toc:
        p_t = doc.add_paragraph()
        r_num = p_t.add_run(f"• {c_num}: ")
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(185, 28, 28)
        r_name = p_t.add_run(c_title)
        r_name.font.italic = True
        
    doc.add_page_break()
    
    # Full Content Text
    content = [
        {
            "id": "chap1",
            "number": "Capítulo 1",
            "title": "El Olor a Nieve y Sangre",
            "scenes": [
                {
                    "num": "Escena 1: La Herrería de la Montaña",
                    "text": """La nieve caía en copos pesados sobre el monte Kurodake. Dentro de la herrería de la familia Hagane, el fuego del horno brillaba con un rojo intenso. Ren, un joven de quince años con cabello negro despeinado y ropas de trabajo manchadas de hollín, martillaba con ritmo constante un lingote de acero Tamahagane.

Su padre, el maestro herrero, observaba en silencio desde la esquina del taller mientras su hermana menor, Miyuki, atizaba el fuego cantando una antigua melodía de la aldea.

—El acero no se moldea solo con fuerza, Ren —dijo su padre, ajustando sus anteojos de madera—. Responde a la intención de tu corazón. Si tu mente está agitada, la hoja se romperá al primer impacto.

—Lo sé, padre —respondió Ren, secándose el sudor de la frente con el antebrazo—. Pero las historias dicen que los demonios han bajado hasta los valles vecinos. Necesitamos armas más resistentes.

—Los Oni pertenecen a las sombras, hijo. Mientras el fuego de este horno siga encendido, esta casa estará protegida.

Pero esa misma noche, la tormenta de nieve apagó la última llama del horno."""
                },
                {
                    "num": "Escena 2: La Masacre del Viento Helado",
                    "text": """A medianoche, el viento arrancó la puerta de madera del taller de cuajo. El frío helado de la montaña inundó la estancia, acompañado por un olor fétido a azufre y carne en descomposición.

Desde la penumbra de la tormenta emergió una figura grotesca de tres metros de altura. Su piel era gris como la piedra, cuatro cuernos afilados brotaban de su frente y sus ojos amarillos brillaban con la inscripción del rango demoníaco. Era un Oni de la Sangre.

En un abrir y cerrar de ojos, antes de que Ren pudiera reaccionar, la bestia atravesó el pecho de su padre con sus garras de metal.

—¡Padre! —gritó Ren, lanzándose hacia la katana en proceso de forja.

Miyuki corrió a proteger el cuerpo de su padre, pero el demonio la sujetó del cuello con una sola mano, alzándola en el aire mientras una risa gutural retumbaba en la casa.

—Sangre de herrero... la más dulce de todas —siseó el Oni, clavando sus colmillos en el hombro de la joven."""
                },
                {
                    "num": "Escena 3: El Pacto de la Hoja Rota",
                    "img": "escena_1.jpg",
                    "text": """Desesperado, Ren tomó la katana incandescente con las manos desnudas, ignorando el dolor de la piel quemándose en la empuñadura. Se arrojó contra el demonio y clavó la hoja en el cuello de la criatura.

El acero no estaba templado. Con un chasquido seco, la katana se partió en dos.

El Oni rugió de ira y golpeó a Ren, arrojándolo contra la pared de piedra. Pero la sangre del demonio goteó abundantemente sobre el pedazo de hoja rota que Ren aún apretaba en su mano derecha.

Una reacción alquímica violenta ocurrió al instante. La sangre del Oni no destruyó el metal: fue absorbida por los poros del acero Tamahagane. La katana rota comenzó a regenerarse, formando un filo hecho de cristal negro y venas carmesí que latían como un corazón humano.

Una voz antigua resonó en la mente de Ren:

«Entrega tu humanidad para cortar la carne de los inmortales...»

Ren apretó la empuñadura. Las venas de sus brazos se volvieron negras mientras la maldición se extendía por su cuello hasta la mejilla derecha. Con un movimiento veloz como el rayo, decapito al demonio de un solo tajo antes de que la criatura pudiera reaccionar."""
                }
            ]
        },
        {
            "id": "chap2",
            "number": "Capítulo 2",
            "title": "La Hermana de las Sombras",
            "scenes": [
                {
                    "num": "Escena 1: Las Venas de la Maldición",
                    "text": """El cuerpo del demonio se disolvió en cenizas negras que el viento de la tormenta dispersó en la noche.

Ren cayó de rodillas sobre la madera ensangrentada, respirando con dificultad. La katana de cristal negro se apagó, volviendo a parecer un pedazo de hierro gastado, pero las venas oscuras en el cuello de Ren no desaparecieron; quedaron grabadas bajo su piel como un tatuaje indeleble.

—Miyuki... —susurró Ren, arrastrándose hacia su hermana.

Miyuki yacía en el suelo. La herida de su hombro no sangraba; en su lugar, marcas púrpuras se extendían rápidamente por su cuello. Sus ojos se abrieron de golpe, pero ya no eran negros: sus pupilas eran rasgadas como las de una bestia y dos pequeños cuernos de hueso comenzaban a asomar entre su cabello.

La joven se abalanzó sobre Ren con un rugido primario, intentando morderlo.

Ren no usó la espada. La abrazó con fuerza, sujetando sus brazos mientras las garras de Miyuki le desgarraban la espalda.

—Miyuki, soy yo... Ren... Por favor, lucha —suplicó Ren, llorando contra el hombro de su hermana—. No me dejes solo."""
                },
                {
                    "num": "Escena 2: El Silencio del Bambú",
                    "text": """Sorprendentemente, al sentir las lágrimas de su hermano y escuchar su voz, el frenesí de Miyuki comenzó a disminuir. Las garras de la joven se retrayeron ligeramente y sus pupilas rasgadas volvieron a dilatarse, mostrando un destello de conciencia humana.

Miyuki comenzó a temblar, emitiendo un gemido de dolor mientras luchaba contra el instinto de devorar carne.

Para evitar que se hiciera daño o atacara a alguien, Ren cortó un trozo de bambú sagrado del altar del taller y lo colocó en la boca de Miyuki, sujetándolo con un cordón de cuero reforzado.

—Te voy a curar, Miyuki —prometió Ren, colocándole su capa sobre los hombros—. Voy a encontrar al demonio original que te hizo esto y voy a obligarlo a devolverte tu humanidad, sin importar cuántos demonios tenga que cortar en el camino.

En ese momento, el crujido de pasos sobre la nieve resonó fuera de la casa."""
                },
                {
                    "num": "Escena 3: El Filo del Cuervo",
                    "text": """De la niebla de la tormenta emergió un hombre alto con una haori azul oscuro decorada con patrones de olas de tormenta. Llevaba una katana de hoja carmesí desenvainada y una máscara de cuervo (Tengu) que cubría su rostro.

Era Kaito, un espadachín de alto rango del Gremio Cuervo de Cazadores de Demonios.

—Apártate del chico —dijo Kaito con una voz fría y calmada—. Esa niña ya no es tu hermana. Es una bestia hambrienta. Si no la ejecuto ahora, mañana devorará a los habitantes de la aldea del valle.

Kaito se desplazó a una velocidad sobrehumana. En menos de un segundo, la hoja carmesí de su katana estaba a milímetros del cuello de Miyuki.

Ren reaccionó por instinto puro: levantó la empuñadura de su katana rota e interceptó el ataque de Kaito. El choque de metales levantó una ráfaga de nieve alrededor de los tres.

Kaito abrió los ojos con sorpresa detrás de la máscara al ver que un simple campesino había bloqueado la velocidad de su Estilo del Trueno."""
                }
            ]
        },
        {
            "id": "chap3",
            "number": "Capítulo 3",
            "title": "La Prueba del Cazador",
            "scenes": [
                {
                    "num": "Escena 1: La Apuesta sobre la Nieve",
                    "text": """—Interesante... —murmuró Kaito, aumentando la presión de su espada sobre la de Ren—. Tienes los reflejos de un cazador, pero tu arma está rota y cargas con la marca de la sangre demoníaca en el rostro. ¿Sabes lo que el Gremio le hace a los que pactan con los Oni?

—Ella no ha lastimado a nadie —respondió Ren, apretando los dientes mientras sus pies se hundían en la nieve por la fuerza de Kaito—. ¡Ella luchó contra el instinto para no matarme! ¡Hay una forma de salvarla!

Miyuki, al ver a su hermano en peligro, se lanzó hacia Kaito no para morderlo, sino para colocar su propio cuerpo como escudo frente a Ren.

Kaito detuvo su ataque en el último instante. Miró a la niña demonio con el bambú en la boca, protegiendo a su hermano con los brazos abiertos, y luego miró la determinación en los ojos de Ren.

Kaito guardó su katana carmesí en la vaina con un sonido metálico seco.

—Tienen tres años —dijo Kaito en voz baja—. Si en tres años no logras entrar al Gremio Cuervo y encontrar la cura, yo mismo regresaré y cortaré sus cabezas. Tu primera prueba comienza ahora: si quieren sobrevivir a la noche, deben llegar a la cumbre de la Montaña Sombría antes del amanecer."""
                },
                {
                    "num": "Escena 2: El Ascenso a la Montaña Sombría",
                    "text": """El viaje hacia la Montaña Sombría estuvo lleno de trampas mortales. La cumbre no solo estaba infestada de trampas de bambú afilado y rocas rodantes puestas por Kaito, sino que el aire era tan enrarecido que cada respiración quemaba los pulmones de Ren.

Miyuki caminaba a su lado, demostrando una fuerza física sobrehumana pero necesitando esconderse dentro de las sombras de las cuevas cuando los primeros rayos del sol matutino comenzaban a rozar los picos.

—Los demonios no pueden tocar la luz del sol... —comprendió Ren, llevando a Miyuki sobre su espalda dentro de un cesto de mimbre especial cubierto con mantas gruesas para protegerla de los rayos solares.

Ren no se detuvo a descansar. Usó la técnica de respiración de la herrería que su padre le enseñó —la Respiración del Fuego del Horno— para mantener la temperatura de sus músculos y evitar colapsar por la falta de oxígeno."""
                },
                {
                    "num": "Escena 3: La Cabaña del Maestro",
                    "text": """Al mediodía, exhausto y con las manos sangrando, Ren alcanzó la cumbre de la Montaña Sombría. Allí, escondida entre un bosque de glicinas de floración púrpura —plantas cuyo olor repelía a los demonios de forma natural—, se encontraba una pequeña cabaña de madera.

Kaito los esperaba sentado en el porche, tomando una taza de té caliente sin la máscara de Tengu. Reveló el rostro de un hombre joven con cicatrices profundas alrededor de los ojos.

—Llegaron tarde por diez minutos —dijo Kaito sin mirarlos—. Pero la niña sigue viva y no has consumido carne humana para alimentar tu marca. Eso significa que pasaron la primera ronda.

Ren dejó la cesta con cuidado en la sombra del porche.

—Enséñame a luchar —exigió Ren, inclinándose profundamente sobre la madera—. Enséñame el arte de la esgrima para destruir a los Seis Lunares Rojos.

Kaito miró la katana rota de cristal negro que Ren llevaba en la cintura.

—No te enseñaré a luchar con una espada común, Ren. Te enseñaré el Estilo de la Sangre Negra, la técnica prohibida que destruye el cuerpo del usuario a cambio de rebanar el alma de las bestias."""
                }
            ]
        },
        {
            "id": "chap4",
            "number": "Capítulo 4",
            "title": "El Arte de la Hoja Maldita",
            "scenes": [
                {
                    "num": "Escena 1: El Entrenamiento del Acero",
                    "text": """Pasaron dos años de entrenamiento infernal en la Montaña Sombría.

Cada día, Ren debía esquivar cientos de trampas con los ojos vendados, aprender a cortar bloques de piedra maciza con katanas de madera y dominar el flujo de la Respiración de Sangre.

La técnica consistía en acelerar el ritmo cardíaco para bombear más oxígeno a los músculos, imitando la fuerza física de los demonios sin perder el control humano. Pero el costo era alto: cada vez que Ren usaba la respiración al máximo, las venas negras de su rostro se extendían unos milímetros más.

Mientras tanto, Miyuki permanecía dentro de la cabaña rodeada de flores de glicina. Pasaba la mayor parte del tiempo durmiendo en un trance profundo para recuperar energía sin necesidad de consumir carne humana."""
                },
                {
                    "num": "Escena 2: La Prueba del Peñasco",
                    "text": """Al final del segundo año, Kaito llevó a Ren frente a una roca gigante de cuatro metros de alto ubicada en el centro del bosque de glicinas.

—Este es tu examen final de graduación —dijo Kaito, cruzando los brazos—. Si puedes cortar este peñasco a la mitad con tu técnica de respiración, te daré permiso para ir a la Selección Final del Gremio Cuervo en el monte Fujikane. Si no lo logras, te quedarás aquí para siempre.

Ren desenvainó su katana. En estos dos años, la hoja de cristal negro había sido reforzada por él mismo en el horno de la cabaña, adquiriendo un filo tan afilado que cortaba las hojas que caían del aire sin tocarlas.

Ren cerró los ojos. Concentró toda la Respiración de Sangre en sus piernas y brazos. Su corazón latió como el golpe de un martillo de herrería: ¡BOOM! ¡BOOM!

—Estilo de la Sangre Negra... Primera Postura: Tajo del Horno Olvidado.

Ren se desplazó como una sombra roja. Un destello de luz carmesí cruzó la noche."""
                },
                {
                    "num": "Escena 3: La Partida hacia el Fujikane",
                    "text": """La enorme roca se dividió en dos mitades perfectas, cayendo a los lados con un estruendo pesado.

Kaito sonrió levemente detrás de su máscara. Se acercó a Ren y le entregó un haori tradicional de color negro con bordados de venas rojas en las mangas, además de una nueva funda para su katana reforzada.

—Estás listo —dijo Kaito—. En la montaña Fujikane te enfrentarás a docenas de demonios capturados por el Gremio durante la prueba de siete días. Si sobrevivís, te convertirás oficialmente en un Cazador de Demonios.

Miyuki salió de la cabaña, ajustándose el bambú de la boca. Se acercó a Ren y le sujetó la manga del haori con fuerza, demostrando que quería acompañarlo en la travesía.

—Iremos juntos, Miyuki —dijo Ren, colocándola dentro de la nueva caja de madera reforzada que había construido para llevarla en su espalda—. El viaje por la cura de tu humanidad comienza ahora."""
                }
            ]
        },
        {
            "id": "chap5",
            "number": "Capítulo 5",
            "title": "La Selección Final (Clímax del Volumen 1)",
            "scenes": [
                {
                    "num": "Escena 1: El Monte de las Glicinas",
                    "text": """El monte Fujikane era un espectáculo visual sobrecogedor. Toda la base de la montaña estaba cubierta por un bosque frondoso de flores de glicina púrpura que brillaban bajo la luz de la luna llena.

En la entrada de la prueba, decenas de jóvenes espadachines de todo el país se habían reunido. Entre ellos destacaba Kaito (un joven de cabello rubio desordenado que llevaba dos katanas cortas) y una chica de túnica blanca con una máscara de zorro que no hablaba con nadie.

Dos niñas vestidas con kimonos ceremoniales y rostros pálidos como muñecas de porcelana se presentaron ante la multitud:

—Bienvenidos a la Selección Final del Gremio Cuervo —dijeron las niñas al unísono—. En esta montaña habitan demonios capturados vivos por los maestros. Para aprobar, solo deben hacer una cosa: sobrevivir durante siete noches dentro del área sin glicinas.

Las puertas de madera de la montaña se abrieron, revelando un bosque oscuro y lúgubre donde el olor a sangre flotaba en el aire."""
                },
                {
                    "num": "Escena 2: El Demonio de las Manos",
                    "text": """Llegada la quinta noche de la prueba, más de la mitad de los aspirantes habían sido eliminados o devorados. Ren avanzaba con cautela por el bosque espeso, manteniendo la caja de Miyuki protegida en un refugio bajo las raíces de un árbol sagrado.

De pronto, la tierra tembló.

Un demonio gigantesco y grotesco emergió de las profundidades del suelo. Tenía decenas de brazos humanos cosidos a su cuerpo que se retorcían como tentáculos, y su rostro estaba oculto detrás de una masa de carne deforme.

—Otro discípulo de Kaito... —rugió el demonio con una voz cavernosa que hizo temblar las hojas de los árboles—. Reconozco el olor de la herrería en tu sangre. ¡Me he comido a todos los niños que ese hombre ha enviado a esta prueba durante los últimos cincuenta años!

Ren sintió un frío helado en la espalda al comprender la verdad: este demonio había sobrevivido en la montaña alimentándose de los alumnos de su maestro."""
                },
                {
                    "num": "Escena 3: La Postura Suprema (Cierre del Tomo 1)",
                    "img": "escena_climax.jpg",
                    "text": """El Demonio de las Manos atacó lanzando ocho de sus brazos gigantescos como lanzas de carne.

Ren esquivó las estocadas por milímetros, usando los árboles como apoyo para ganar altura. Las venas negras de su rostro se encendieron en un brillo carmesí brillante mientras inhalaba todo el aire de sus lungs.

—Respiración de Sangre... Segunda Postura: Corte del Horno de la Memoria.

Ren concentró toda su humanidad y su dolor por su familia caída en el filo de su katana de cristal negro. La hoja se envolvió en una espiral de fuego rojo y sombras oscuras que cortaron los brazos del monstruo como si fueran ramas secas.

En un estallido de velocidad pura, Ren cruzó la guardia del demonio y lanzó un tajo vertical directo al cuello del gigante de los brazos.

La cabeza del demonio salió volando por los aires, mientras el cuerpo de la bestia comenzaba a disolverse en cenizas rojas bajo la luz de la luna de la séptima noche.

Ren cayó de pie sobre la nieve, apoyándose en su katana mientras el amanecer iluminaba las flores de glicina del pico de la montaña.

Había aprobado la prueba.

[ CONTINUARÁ EN EL VOLUMEN 2 ]"""
                }
            ]
        }
    ]
    
    bkm_id = 1
    for chap in content:
        p_c = doc.add_paragraph()
        add_bookmark(p_c, chap["id"], bkm_id)
        
        r_cnum = p_c.add_run(f"{chap['number']}\n")
        r_cnum.font.name = "Arial"
        r_cnum.font.size = Pt(14)
        r_cnum.font.bold = True
        r_cnum.font.color.rgb = RGBColor(185, 28, 28)
        
        r_ctitle = p_c.add_run(chap["title"])
        r_ctitle.font.name = "Arial"
        r_ctitle.font.size = Pt(20)
        r_ctitle.font.bold = True
        r_ctitle.font.color.rgb = RGBColor(220, 38, 38)
        
        end_bookmark(p_c, bkm_id)
        bkm_id += 1
        
        for sc in chap["scenes"]:
            p_s = doc.add_paragraph()
            r_snum = p_s.add_run(f"\n◆ {sc['num']}\n")
            r_snum.font.name = "Arial"
            r_snum.font.size = Pt(13)
            r_snum.font.bold = True
            r_snum.font.color.rgb = RGBColor(71, 85, 105)
            
            p_stext = doc.add_paragraph()
            r_stext = p_stext.add_run(sc["text"])
            r_stext.font.name = "Georgia"
            r_stext.font.size = Pt(11)
            
            if "img" in sc:
                img_p = os.path.join(target_root, sc["img"])
                if os.path.exists(img_p):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_p, width=Inches(5))
                    
        doc.add_page_break()
        
    doc_dest_root = os.path.join(target_root, "libro.docx")
    doc_dest_sys = os.path.join(target_sys, "libro.docx")
    doc.save(doc_dest_root)
    shutil.copy2(doc_dest_root, doc_dest_sys)
    print(f"Generated docx at {doc_dest_root}")

create_oni_vol1_docx()

# Generate ficha_producto.json
ficha_data = {
    "titulo": "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre) - Volumen 1: La Noche de las Hojas Rotas",
    "autor": "Nicolás Noguera",
    "precio": 20.00,
    "moneda": "USD",
    "genero": "Manga / Light Novel / Fantasía Oscura / Acción Sengoku",
    "headline": "Una masacre en la nieve. Una katana rota de cristal negro. Y una hermana transformada en demonio que aún recuerda el amor.",
    "descripcion": "El monte Kurodake cae bajo la tormenta y los demonios de la sangre descienden. Ren, un joven herrero, pierde a su familia y toma una katana incandescente que se quiebra en la batalla. Pero al absorber la sangre del monstruo, el metal cobra vida convirtiéndose en una hoja de cristal negro y venas carmesí. Para salvar a su hermana Miyuki de la maldición demoníaca, Ren deberá dominar el temible Estilo de la Sangre Negra y sobrevivir a la Selección Final del Gremio Cuervo.",
    "beneficios": [
        "Manuscrito oficial ilustrado completo en formato .docx listo para eReaders y Amazon KDP.",
        "Ilustraciones de escenas de alta calidad en estética anime/manga dark fantasy.",
        "Incluye la primera entrega épica de la nueva saga de fantasía oscura de Nicolás Noguera."
    ],
    "capitulos": [
        "Capítulo 1: El Olor a Nieve y Sangre",
        "Capítulo 2: La Hermana de las Sombras",
        "Capítulo 3: La Prueba del Cazador",
        "Capítulo 4: El Arte de la Hoja Maldita",
        "Capítulo 5: La Selección Final"
    ]
}

with open(os.path.join(target_root, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)
with open(os.path.join(target_sys, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)

# Generate index.html landing page for Oni no Ketsuryu Vol 1
html_content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Oni no Ketsuryū Vol 1 - Nicolás Noguera | Tienda Oficial</title>
    <style>
        :root {
            --bg-color: #0b0f19;
            --card-bg: #111827;
            --accent: #dc2626;
            --accent-hover: #b91c1c;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
        }
        body {
            font-family: 'Segoe UI', system-ui, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 0;
            line-height: 1.6;
        }
        .header-banner {
            width: 100%;
            max-height: 400px;
            object-fit: cover;
            border-bottom: 3px solid var(--accent);
        }
        .container {
            max-width: 1000px;
            margin: 0 auto;
            padding: 40px 20px;
        }
        .product-grid {
            display: grid;
            grid-template-columns: 1fr 2fr;
            gap: 40px;
            margin-bottom: 50px;
        }
        @media (max-width: 768px) {
            .product-grid { grid-template-columns: 1fr; }
        }
        .cover-img {
            width: 100%;
            border-radius: 12px;
            box-shadow: 0 10px 30px rgba(220, 38, 38, 0.3);
            border: 1px solid rgba(220, 38, 38, 0.3);
        }
        .badge {
            background: var(--accent);
            color: white;
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            display: inline-block;
            margin-bottom: 15px;
        }
        h1 { font-size: 2.2rem; margin: 10px 0; color: #fff; }
        .author { color: var(--accent); font-weight: 600; font-size: 1.1rem; margin-bottom: 20px; }
        .price-tag {
            font-size: 2rem;
            font-weight: 800;
            color: #fff;
            margin: 20px 0;
        }
        .price-tag span { font-size: 1rem; color: var(--text-muted); }
        .buy-btn {
            background: linear-gradient(135deg, #dc2626, #991b1b);
            color: white;
            text-decoration: none;
            padding: 16px 36px;
            border-radius: 8px;
            font-weight: bold;
            font-size: 1.2rem;
            display: inline-block;
            box-shadow: 0 6px 20px rgba(220, 38, 38, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
        }
        .buy-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(220, 38, 38, 0.6);
        }
        .section-title {
            font-size: 1.5rem;
            border-left: 4px solid var(--accent);
            padding-left: 12px;
            margin: 40px 0 20px;
            color: #fff;
        }
        .chapters-list {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 25px;
            border: 1px solid rgba(255,255,255,0.05);
        }
        .chapters-list li {
            margin-bottom: 12px;
            color: var(--text-main);
        }
        .gallery-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin-top: 20px;
        }
        .gallery-grid img {
            width: 100%;
            border-radius: 8px;
            border: 1px solid rgba(255,255,255,0.1);
        }
    </style>
</head>
<body>
    <img src="banner.jpg" alt="Banner Oni no Ketsuryu Vol 1" class="header-banner">
    
    <div class="container">
        <div class="product-grid">
            <div>
                <img src="portada.jpg" alt="Portada Oni no Ketsuryu Vol 1" class="cover-img">
            </div>
            <div>
                <span class="badge">MANGA / LIGHT NOVEL</span>
                <h1>Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)</h1>
                <div class="author">Volumen 1: La Noche de las Hojas Rotas • Por Nicolás Noguera</div>
                <p>Una masacre en la nieve. Una katana rota de cristal negro. Y una hermana transformada en demonio que aún recuerda el amor.</p>
                <p>El monte Kurodake cae bajo la tormenta y los demonios de la sangre descienden. Ren, un joven herrero, pierde a su familia y toma una katana incandescente que se quiebra en la batalla. Pero al absorber la sangre del monstruo, el metal cobra vida convirtiéndose en una hoja de cristal negro y venas carmesí.</p>
                
                <div class="price-tag">$20.00 <span>USD</span></div>
                <a href="#" class="buy-btn">COMPRAR AHORA ($20 USD)</a>
            </div>
        </div>

        <h2 class="section-title">Contenido del Volumen 1</h2>
        <div class="chapters-list">
            <ul>
                <li><strong>Capítulo 1:</strong> El Olor a Nieve y Sangre (Escenas 1-3)</li>
                <li><strong>Capítulo 2:</strong> La Hermana de las Sombras (Escenas 1-3)</li>
                <li><strong>Capítulo 3:</strong> La Prueba del Cazador (Escenas 1-3)</li>
                <li><strong>Capítulo 4:</strong> El Arte de la Hoja Maldita (Escenas 1-3)</li>
                <li><strong>Capítulo 5:</strong> La Selección Final (Clímax del Volumen 1 - Escenas 1-3)</li>
            </ul>
        </div>

        <h2 class="section-title">Ilustraciones Interiores Destacadas</h2>
        <div class="gallery-grid">
            <img src="escena_1.jpg" alt="El Pacto de la Hoja Rota">
            <img src="escena_climax.jpg" alt="La Postura Suprema en Fujikane">
        </div>
    </div>
</body>
</html>
"""

with open(os.path.join(target_root, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)
with open(os.path.join(target_sys, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)

print("Generated HTML landing pages successfully")
