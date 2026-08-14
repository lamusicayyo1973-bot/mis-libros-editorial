# -*- coding: utf-8 -*-
import sys
import os
import io
import json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

vol10_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-10")
vol10_dir_downloads = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-10")

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(180, 20, 20)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(120, 15, 15)
    return p

def add_image_safe(doc, img_path, width_inches=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

def build_vol10_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    r_title = p_title.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n\nVOLUMEN 10: EL AMANECER DEL ACERO SANTO\n(Gran Final de Saga)")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(160, 20, 20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Obra Original por Nicolás Noguera\nEdición Ilustrada de Alta Definición (Dark Fantasy Anime)")
    r_sub.font.name = "Georgia"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Portada image
    add_image_safe(doc, vol10_dir / "portada.jpg", 5.0)

    doc.add_page_break()

    # Synopsis / Resumen
    add_heading_styled(doc, "SINOPSIS DEL VOLUMEN 10 (GRAN FINAL)", 1)
    p_syn = doc.add_paragraph()
    p_syn.paragraph_format.line_spacing = 1.15
    p_syn.paragraph_format.space_after = Pt(10)
    p_syn.add_run("Con la derrota definitiva de Kageyama, la luz de la mañana envuelve las ruinas de la capital imperial. Miyuki logra sanar el corazón de Ren con su energía purificada. Los guerreros disuelven la Hermandad del Sol en una emotiva ceremonia final. Ren y Miyuki regresan al monte Kurodake para convertir la herrería familiar en un espacio de paz. Cien años después, las reencarnaciones de los hermanos caminan bajo los cerezos del Tokio moderno, completando la saga épica por Nicolás Noguera.")

    add_image_safe(doc, vol10_dir / "banner.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 1
    add_heading_styled(doc, "CAPÍTULO 1: El Silencio del Alba", 1)
    add_image_safe(doc, vol10_dir / "escena_1.jpg", 6.0)
    p_c1 = doc.add_paragraph()
    p_c1.add_run("La luz dorada bañó las calles en ruinas. En medio del césped tranquilo, Miyuki sostuvo las manos de Ren mientras una energía luminosa cerraba las heridas del combate...")
    add_image_safe(doc, vol10_dir / "escena_c1_e1.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c1_e2.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c1_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 2
    add_heading_styled(doc, "CAPÍTULO 2: Las Lágrimas y la Paz", 1)
    add_image_safe(doc, vol10_dir / "escena_c2_e1.jpg", 6.0)
    p_c2 = doc.add_paragraph()
    p_c2.add_run("En la ceremonia de despedida, los guerreros supervivientes inclinaron sus cabezas en respeto y disolvieron la Hermandad del Sol. Los hermanos tomaron el camino de regreso a las montañas...")
    add_image_safe(doc, vol10_dir / "escena_c2_e2.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c2_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 3
    add_heading_styled(doc, "CAPÍTULO 3: La Forja de la Nueva Era", 1)
    add_image_safe(doc, vol10_dir / "escena_c3_e1.jpg", 6.0)
    p_c3 = doc.add_paragraph()
    p_c3.add_run("Guardando la katana solar en el relicario familiar, la antigua forja se transformó en un taller de herramientas de labranza rodeado de glicinas y rosas en flor...")
    add_image_safe(doc, vol10_dir / "escena_c3_e2.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c3_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 4
    add_heading_styled(doc, "CAPÍTULO 4: Las Memorias Grabadas", 1)
    add_image_safe(doc, vol10_dir / "escena_c4_e1.jpg", 6.0)
    p_c4 = doc.add_paragraph()
    p_c4.add_run("Las leyendas del sol se transmitieron a través de cuadros y canciones nocturnas bajo las estrellas, mientras linternas de papel flotaban hacia el horizonte...")
    add_image_safe(doc, vol10_dir / "escena_c4_e2.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c4_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 5
    add_heading_styled(doc, "CAPÍTULO 5: El Renacer bajo los Cerezos", 1)
    add_image_safe(doc, vol10_dir / "escena_c5_e1.jpg", 6.0)
    p_c5 = doc.add_paragraph()
    p_c5.add_run("Cien años después, en las calles pacíficas del Tokio moderno, dos jóvenes sonríen bajo los cerezos en flor. El sol de una nueva era ilumina su camino, cerrando la historia para siempre...")
    add_image_safe(doc, vol10_dir / "escena_c5_e2.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_c5_e3.jpg", 6.0)
    add_image_safe(doc, vol10_dir / "escena_climax.jpg", 6.0)

    # Save to both locations
    out1 = vol10_dir / "libro.docx"
    out2 = vol10_dir_downloads / "libro.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print(f"Generated libro.docx successfully for Volume 10 at {out1}")

if __name__ == "__main__":
    build_vol10_docx()
