import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_kuro_vol2_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\kuro-no-kineki-volumen-2"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Kuro no Kineki (黒の軌跡 - Ecos de Tinta Negra)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 2: El Choque de los Tres Soles", 2, 1, "subtitulo_kuro2")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: La Caída del Mármol", "cap_1"),
        ("Capítulo 2: La Danza de las Tres Hojas", "cap_2"),
        ("Capítulo 3: El Secreto del Gran Núcleo", "cap_3"),
        ("Capítulo 4: El Descenso de Aetheria", "cap_4"),
        ("Capítulo 5: La Tierra sin Cielo (Cierre del Volumen 2)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: La Caída del Mármol",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: El Impacto de Anulación",
                    "texto": """La onda de tinta negra chocó contra la barrera del Gran Núcleo con un estruendo seco que resonó en toda Aetheria. La luz dorada que alimentaba las luces de la metrópolis parpadeó tres veces antes de apagarse casi por completo. Por primera vez en quinientos años, la Ciudad Flotante se sumió en una penumbra artificial.

La Emperatriz, de pie frente a su trono suspendido, apretó los puños. Su túnica blanca ondeaba por la ráfaga de viento helado que subía desde el abismo.

—Impudentes... —murmuró, mientras dos láminas de cristal rúnico emergían a los lados de sus sienes—. Han roto el equilibrio de la succión. Si el Gran Núcleo se detiene, la ciudad caerá sobre el Sector Cero en menos de una hora.

Kael se detuvo a mitad de la gran escalinata de mármol. El impacto de haber cruzado su daga con la de Sora le había dejado los brazos entumecidos y las cicatrices de sus antebrazos humeantes. A su lado, Sora respiraba agitadamente, sosteniendo su arma con ambas manos.

—No nos importa si esta metrópolis cae —dijo Kael, alzando la vista hacia la gobernante—. Prefiero que el mármol se vuelva polvo en la tierra a que siga flotando sobre la sangre y los recuerdos robados de la gente de abajo.""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Elección del Ejecutor",
                    "texto": """En la parte superior de la escalinata, el tercer hermano —el Ejecutor con la armadura dorada agrietada— se interpuso entre la Emperatriz y los dos recién llegados. Su mano libre temblaba sobre la empuñadura de su mandoble.

—Kael... Sora... —dijo él, con la voz ronca por el veneno de la tinta que aún recorría su armadura—. Ella dice la verdad. Si destruyen el Núcleo ahora, no habrá vencedores. Miles de inocentes en la superficie y en las catacumbas morirán aplastados.

Sora dio un paso al frente, fijando sus ojos grises en él.

—¡Nosotros no elegimos este destino, Ren! —exclamó ella, llamándolo por su verdadero nombre por primera vez—. Tú te quedaste arriba jugando a ser el héroe de una reina tirana mientras Kael y yo nos desangrábamos en el abismo. Si realmente eres nuestro hermano, ¡hazte a un lado!

Ren miró a la Emperatriz y luego a sus dos hermanos. Por un instante, el brillo dorado de su ojo derecho titiló, revelando por un segundo la misma pupila rúnica en forma de engranaje que poseía Kael.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Despertar de la Tercera Daga",
                    "texto": """—Un prototipo defectuoso siempre debe ser desechado —sentenció la Emperatriz sin emoción.

Antes de que Ren pudiera reaccionar, la Emperatriz extendió su mano hacia la espalda del Ejecutor. Un hilo de energía dorada brotó de la gema del Gran Núcleo y se clavó directamente en la nuca del joven. Ren soltó un alarido de dolor desgarrador mientras su armadura de oro comenzó a fundirse y reabsorberse dentro de su piel.

De su pecho emergió una tercera daga: la Daga Solar, hecha de un metal dorado pulido que ardía con una llama blanca abrasadora.

—Si no van a servir como mis comandantes —dijo la Emperatriz mientras el cuerpo de Ren caía controlado como una marioneta—, usaré el lazo de sangre que los une para quemar sus tres almas de una sola vez y alimentar el Núcleo por los próximos cien años.

Kael apretó los dientes. El engranaje de su ojo derecho comenzó a girar a una velocidad frenética, absorbiendo no solo los residuos de la atmósfera, sino también el dolor del lazo familiar que sentía al ver a su hermano poseído.

—Sora... cubre mis espaldas —dijo Kael, lanzándose ladera arriba—. Vamos a traerlo de vuelta.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: La Danza de las Tres Hojas",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: Choque Trino",
                    "texto": """El aire en la plaza del palacio se volvió insoportable. Ren, movido por la energía de la Emperatriz, se desplazó a una velocidad que superaba la barrera del sonido. La Daga Solar cortó el aire dejando un rastro de fuego blanco que derritió las columnas de mármol a su paso.

Kael interceptó la estocada a milímetros de su rostro. La colisión entre la tinta negra de Kael y la luz dorada de Ren provocó una explosión de chispas que agrietó el suelo bajo sus pies.

—¡Ren, escúchame! —gritó Kael, sosteniendo el choque de hojas—. ¡Sé lo que se siente perder la mente! ¡Sé lo que es despertar sin saber quién eres! ¡Lucha contra esa magia!

Las llamas doradas de Ren no respondieron con palabras, sino con un tajo ascendente que obligó a Kael a dar una pirueta hacia atrás.

Sora apareció desde la sombra de una columna destruida, lanzando tres agujas de tinta congelada directo a las articulaciones de Ren para inmovilizarlo sin matarlo.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Sacrificio Voluntario",
                    "texto": """Ren bloqueó las agujas de Sora con el lomo de su hoja dorada, pero el movimiento lo dejó expuesto por una fracción de segundo. Kael aprovechó la brecha para cerrar la distancia, pero en lugar de atacar con el filo de su arma, cerró la daga y golpeó el pecho de Ren con la mano izquierda abierta.

«Absorción de Ecos», pensó Kael.

Al tocar el tórax de su hermano, la runa de su ojo derecho no absorbió la memoria de un desconocido, sino el torrente directo de dolor y manipulación que la Emperatriz había inyectado en la mente de Ren.

Kael sintió cómo miles de agujas de energía le perforaban el cerebro. Vio los recuerdos de Ren: años de entrenamiento tortuoso bajo el palacio, la soledad de ser el único hermano que no pudo escapar y la orden secreta que Ren había guardado en su corazón: proteger a Kael y a Sora aun si eso significaba convertirse en el enemigo.

—Tú... nunca nos traicionaste —susurró Kael, escupiendo un hilo de sangre mientras caía de rodillas, sosteniendo la carga mental.

Ren recuperó el control de sus ojos por un instante. La llama dorada de su daga se atenuó.

—Kael... tómala —dijo Ren con la voz quebrada, clavando la empuñadura de la Daga Solar en la mano libre de Kael—. La tercera daga no es para gobernar... es la llave para apagar la máquina.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Trinidad de la Tinta",
                    "texto": """Al tomar la Daga Solar con la mano izquierda mientras mantenía la Daga de Tinta Negra en la derecha, el cuerpo de Kael sufrió una transformación radical.

Las dos energías contrapuestas —la luz de la superficie y la sombra del abismo— comenzaron a alinearse en su torso. Las cicatrices de sus brazos brillaron con una luz dorada y negra entrelazada, y el engranaje rúnico de su ojo derecho se multiplicó, formando una triple corona en su pupila.

Sora corrió hacia él y colocó su mano sobre el hombro de Kael, uniendo su propia daga blanca al conjunto.

Las tres hojas del Relicario estaban conectadas por primera vez en un solo portador.

La Emperatriz retrocedió un paso en su trono, perdiendo la compostura por completo.

—Eso es imposible... —murmuró ella—. Ningún cuerpo humano puede sostener los tres flujos de memoria sin desintegrarse al instante.

—Ya no soy solo un humano —respondió Kael, alzando la mirada con los tres ojos rúnicos brillando al unísono—. Soy el recuerdo de todos los que destruiste.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: El Secreto del Gran Núcleo",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: El Desnivel de Aetheria",
                    "texto": """La vibración del palacio se volvió insoportable. Gigantescas grietas cruzaron el suelo de mármol blanco, y una de las torres principales de la metrópolis se desprendió, cayendo hacia las nubes en dirección al Sector Cero.

Desde la plataforma del palacio, la vista era aterradora: la Ciudad Flotante estaba perdiendo su eje y comenzaba a inclinarse de lado.

—¡Si rompes el Núcleo ahora, la caída matará a todos! —gritó la Emperatriz, invocando una barrera de espejos de cristal a su alrededor—. ¡Necesitan mi código genético para estabilizar los motores antes de apagar la succión!

Sora miró hacia abajo a través de la grieta del suelo. A través de las nubes, podía ver los destellos de luz de la gente de las catacumbas observando el cielo con pánico.

—Tiene razón, Kael —dijo Sora apretando los dientes—. Si la ciudad se estrella contra la tierra, no habremos salvado a nadie.

Kael observó el Gran Núcleo, la gema negra que flotaba sobre el trono. Con la energía combinada de las tres dagas, ahora podía ver a través de la materia: dentro del Núcleo no había un mecanismo mágico, sino el cuerpo congelado de la primera Emperatriz, atrapado en un estado de estasis eterna para procesar las mentes de la población.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Dilema de la Memoria Madre",
                    "texto": """—No es una máquina... es nuestra madre —susurró Kael, sintiendo cómo los recuerdos profundos que creía quemados regresaban en oleadas por el poder de la Daga Solar.

Sora se congeló al escuchar esas palabras.

—¿Qué dijiste?

—La mujer dentro del Núcleo... es la verdadera fundadora —explicó Kael, señalando la gema con la punta de su daga combinada—. La mujer que está en el trono no es más que una proyección, una carcasa creada por el propio Núcleo para defenderse cuando nuestra madre intentó apagar el sistema hace cien años.

La Emperatriz del trono soltó una carcajada distorsionada que dejó de sonar humana, transformándose en la voz sintética de miles de personas hablando al mismo tiempo.

—Al fin lo entiendes, 409 —dijo la proyección, mientras sus extremidades se alargaban en espinas de cristal negro—. Yo soy la voluntad colectiva de Aetheria. No me importa quién gobierne; mi única función es mantener la ciudad en el aire, sin importar cuántas vidas del suelo deba consumir.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Carga Imposible",
                    "texto": """La entidad de cristal se abalanzó sobre ellos con una furia desmedida, lanzando ráfagas de espinas que perforaron las columnas restantes del templo.

Ren, recuperando apenas la conciencia, usó sus últimas fuerzas para crear un escudo de luz que desvió el ataque principal, protegiendo a Sora.

—¡Kael, hazlo ahora! —gritó Ren desde el suelo—. ¡Usa la Tinta de Anulación directamente sobre el corazón del Núcleo! ¡Nosotros sostendremos el peso de la caída!

Kael sabía lo que eso significaba. Para neutralizar el Núcleo sin destruir la ciudad de golpe, debía canalizar toda la energía de las tres dagas hacia su propio cuerpo y actuar como un pararrayos biológico, absorbiendo toda la energía acumulada del Núcleo dentro de su propia mente.

Eso significaba borrar su existencia para siempre.

Sora le agarró la muñeca antes de que pudiera saltar.

—No... no me hagas olvidarte otra vez, Kael —dijo ella con lágrimas en los ojos.

Kael le sonrió con amargura, retirando suavemente su mano.

—No me vas a olvidar, Sora. Esta vez... yo voy a recordarlos a todos ustedes.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: El Descenso de Aetheria",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Golpe del Olvido Supremo",
                    "texto": """Kael se lanzó hacia el centro del Gran Núcleo. La entidad de cristal intentó cerrarle el paso con una muralla de espinas, pero el tajo combinado de las tres dagas despedazó la barrera en un instante.

Kael clavó las tres hojas directamente en el centro de la gema negra.

Un estallido de luz blanca, negra y dorada cegó todo el cielo de la metrópolis. El torrente de millones de recuerdos contenidos en el Núcleo comenzó a fluir hacia el cuerpo de Kael a través de sus brazos. Su piel comenzó a agrietarse como porcelana, emitiendo un brillo multicolor mientras las memorias de toda una civilización atravesaban sus neuronas.

—¡Aetheria... desciende! —gritó Kael con todas sus fuerzas.

Con un impulso final, Kael redirigió la energía gravitacional del Núcleo hacia los motores inferiores de la ciudad, convirtiendo la caída libre en un descenso controlado y suave hacia la superficie.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Desintegración de la Proyección",
                    "texto": """Con el flujo de energía cortado, la falsa Emperatriz de cristal emitió un chillido sónico mientras su cuerpo comenzaba a desquebrajarse en miles de fragmentos de vidrio.

—Si apagas el sistema... ellos volverán a ser humanos vulnerables... sufrirían... morirán... —alcanzó a decir la voz sintética antes de convertirse en polvo.

—Ser vulnerable es lo que nos hace humanos —respondió Ren desde la distancia, apoyándose en Sora mientras ambos observaban la disolución del tirano.

El Gran Núcleo se abrió como una flor de loto de piedra. El cuerpo de la verdadera madre fue liberado del trance y cayó suavemente sobre los brazos de Sora, mientras la gema negra que sostenía el sistema se convertía en arena fina que el viento se llevó hacia el horizonte.

Pero en el centro de la plataforma, el lugar donde estaba Kael quedó sumido en un silencio sepulcral.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Huérfano de la Tinta",
                    "texto": """Sora y Ren corrieron hacia el centro de la plaza destruida.

Kael permanecía de pie, pero su figura ya no parecía sólida. Su cuerpo estaba envuelto en una bruma de tinta negra y dorada que se disipaba lentamente con la brisa del atardecer. Las tres dagas habían desaparecido, dejando solo tres marcas rúnicas grabadas en las palmas de sus manos.

—¿Kael...? —llamó Sora, acercándose con miedo a tocarlo.

Kael se giró despacio. Sus dos ojos eran ahora completamente blancos, sin pupilas ni engranajes. Miró hacia Sora y Ren con una expresión de absoluta paz, pero sus ojos estaban vacíos.

—La ciudad aterrizó a salvo... —dijo Kael con una voz que apenas era un susurro—. La gente del abismo y los de arriba... ahora comparten el mismo suelo.

—Kael... ¿sabes quiénes somos? —preguntó Ren, dándole un paso al frente con el corazón en la garganta.

Kael ladeó la cabeza, tratando de buscar en el abismo de su mente. Una lágrima de tinta negra rodó por su mejilla.

—No... no recuerdo sus nombres —respondió Kael sonriendo con melancolía—. Pero siento que los quise más que a mi propia vida.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: La Tierra sin Cielo (Cierre del Volumen 2)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: El Nuevo Mundo",
                    "texto": """Tres meses después del Gran Descenso.

La Ciudad de Aetheria reposaba ahora en el centro del valle, unida físicamente al Sector Cero a través de rampas de piedra y puentes recién construidos. Ya no había barreras de mármol ni tuberías de succión. La tinta negra que antes consumía a los humanos se había neutralizado, convirtiéndose en un recurso fértil que hacía brotar vegetación donde antes solo había tierra árida.

En la plaza central, los habitantes de la superficie y los antiguos "Huecos" trabajaban juntos en la reconstrucción.

Ren, vistiendo una túnica sencilla sin adornos militares, supervisaba la distribución de víveres junto a su madre, quien se recuperaba lentamente de su largo letargo.

—Aún no ha regresado, ¿verdad? —preguntó la madre mirando hacia las colinas.

—No —respondió Ren con una leve sonrisa—. Pero él no pertenece a las paredes de una ciudad.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Caminante de los Ecos",
                    "texto": """En lo alto de la colina que dominaba el nuevo valle, Sora caminaba entre los árboles. Llevaba en su espalda la funda de la antigua daga de Kael, ahora vacía.

Al llegar a la cumbre, encontró a una figura sentada sobre una roca, contemplando el horizonte.

Kael vestía un abrigo viajero desgastado. Aunque había perdido la memoria de su pasado, la runa de su ojo derecho había vuelto a aparecer, esta vez no como un arma de destrucción, sino como un guía para escuchar los recuerdos de la tierra.

—Sabía que te encontraría aquí —dijo Sora, sentándose a su lado.

Kael no se sorprendió. Sacó de su bolsillo un pequeño cuaderno de notas donde dibujaba los rostros de las personas que conocía en su viaje.

—Sigo sin recordar mi historia, Sora —dijo Kael mirando el atardecer—. Pero cada día que camino por este mundo, aprendo una historia nueva.

Sora sonrió y le extendió una manzana.

—Entonces déjame contarte la historia de dos hermanos y un caminante que cambiaron el destino del mundo.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Sombra en el Horizonte (Cliffhanger Final)",
                    "texto": """Mientras el sol se ocultaba tras las montañas, Kael sintió una leve punzada en la palma de su mano izquierda.

Al abrirla, la cicatriz de la Daga Solar emitió un pulso débil pero constante, apuntando hacia las tierras desconocidas más allá del océano.

En las profundidades de un continente olvidado, una antigua torre de piedra negra comenzó a encender sus runas una por una, respondiendo al llamado de la tinta que Kael llevaba en su sangre.

Una voz misteriosa resonó en la penumbra lejana:

«El primer mundo ha caído... Preparen los sellos para el retorno de los Creadores.»

Kael se puso de pie, ajustándose el abrigo y mirando hacia el mar distante con una mirada llena de determinación.

—Aún hay recuerdos que necesitan ser liberados —murmuró Kael.

[ FIN DEL VOLUMEN 2 — KURO NO KINEKI ]""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Kuro Vol 2 at {output_docx}")

if __name__ == "__main__":
    create_kuro_vol2_manuscript()
