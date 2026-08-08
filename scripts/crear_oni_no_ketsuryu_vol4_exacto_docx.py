import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00) # Dark Red
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slate
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add XML Bookmark
    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_oni_vol4_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-4"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 4: El Distrito de los Espejos y la Mariposa de la Sombra", 2, 1, "subtitulo_vol4")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: La Finca del Fuego", "cap_1"),
        ("Capítulo 2: La Ciudad que Nunca Duerme", "cap_2"),
        ("Capítulo 3: El Doble Estirpe", "cap_3"),
        ("Capítulo 4: La Purificación de la Sangre", "cap_4"),
        ("Capítulo 5: Las Cenizas de la Misericordia (Clímax del Volumen 4)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: La Finca del Fuego",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: Las Memorias del Primer Guerrero de la Hermandad",
                    "texto": """Ren llegó a la Finca del Fuego al atardecer. La residencia, antes llena de vida, permanecía en un silencio sepulcral. El padre de Kenshin, un antiguo Maestro Celestial retirado y consumido por la bebida, lo recibió en el patio central.

Al ver la guarda con forma de llama en la katana negra de Ren y la marca carmesí en su mejilla, el anciano dejó caer su copa de sake.

—Esa marca... —murmuró el anciano con voz temblorosa—. No es una simple cicatriz de venas demoníacas. Es la Marca del Sol, el sello de los primeros espadachines que casi destruyen al Rey Oni hace mil años.

El anciano lo guio hasta la biblioteca secreta de la finca y le entregó un volumen antiguo con páginas de papel arroz desgastado: Los Diarios del Sol.""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Danza del Sol",
                    "texto": """Ren pasó dos días sin comer ni dormir dentro de la biblioteca, estudiando los grabados del diario. Las ilustraciones no mostraban posturas de esgrima complejas, sino una secuencia de doce movimientos de danza ritual que los herreros antiguos realizaban desde el amanecer hasta el anochecer para pedir la bendición del fuego.

Al cerrar los ojos, Ren recordó a su propio padre ejecutando esa misma danza frente al horno de la herrería cada año nuevo.

—No era solo un ritual... —comprendió Ren, poniéndose de pie con la katana en la mano—. La Estilo del Fuego Carmesí de Herrería que me enseñó mi padre era la versión disfrazada de la Danza del Sol.

Al ejecutar el primer movimiento dentro del dojo, una ola de calor puro envolvió la habitación, volviendo la hoja negra de su katana de un color dorado incandescente.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Petición de la Mariposa",
                    "texto": """Al salir de la finca, una mariposa de luz morada revoloteó sobre la cabeza de Ren. Detrás de él reapareció la chica silenciosa de la Selección Final —la joven de la máscara de zorro, cuyo nombre era Aoi, la nueva discípula del Maestro Celestial del Ingesta y Veneno.

—El Gremio requiere tu presencia en el Distrito de las Luces —dijo Aoi con voz suave pero firme—. Un demonio de nivel Quinto Estirpe de Sangre ha convertido el distrito de entretenimientos en su terreno de cacería. Los cazadores enviados anteriormente han desaparecido sin dejar rastro.

Ren miró la caja de madera en su espalda, donde Miyuki descansaba.

—Nos ponemos en marcha —respondió Ren—. Esta vez no dejaremos que ningún inocente caiga.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: La Ciudad que Nunca Duerme",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: El Lujo de las Sombras",
                    "texto": """El Distrito de las Luces era una metrópolis nocturna deslumbrante. Gigantescas linternas de seda roja, casas de te con cortinas de bambú y puentes de madera sobre canales de agua iluminaban la noche. Cientos de personas caminaban festejando con música de samisen.

Pero para los ojos rúnicos de Ren y el olfato desarrollado de Aoi, la ciudad olía a un veneno dulce y denso.

—El demonio no ataca en los callejones —explicó Aoi, ajustando las agujas de veneno de glicina en sus mangas—. Vive dentro de una de las casas de entretenimientos más prestigiosas como una matrona de alta sociedad.

Para infiltrarse sin llamar la atención de los guardias, Ren, Aoi y el muchacho del haori amarillo (que los alcanzó en el camino) tuvieron que disfrazarse con kimonos tradicionales para ingresar a la casa principal.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Matrona del Espejo",
                    "texto": """Dentro de la casa Yoshiwara, la matrona principal —una mujer de belleza deslumbrante llamada Daki— descansaba sobre un diván de seda. Llevaba una faja obi de seda de diez metros de largo decorada con patrones de flores de cerezo.

Sin embargo, detrás de su reflejo en los espejos de bronce de la habitación, su verdadera forma se revelaba: una criatura de piel pálida con venas verdes y la marca del Quinto Estirpe de Sangre en ambos ojos.

—Guerreros de la Hermandad disfrazados... qué falta de respeto al arte —siseó Daki con voz melodiosa.

Con un movimiento de sus dedos, la faja obi de seda cobró vida propia, transformándose en láminas de acero flexible que atravesaron las paredes de madera del dojo.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Trampa de la Seda",
                    "texto": """La faja obi se dividió en seis cintas independientes que atacaron desde todos los ángulos.

Aoi se desplazó con una agilidad impresionante, utilizando el Estilo del Veneno para cortar las cintas con sus dagas cortas llenas de veneno de glicina. Sin embargo, al cortar la seda, las cintas no sangraban: liberaban a las personas que Daki había atrapado y almacenado vivas dentro del tejido.

—¡Están atrapadas dentro de la tela! —gritó Aoi—. ¡Si cortamos a ciegas, mataremos a los rehenes!

Ren reaccionó al instante. Usó la guarda de llama de Kenshin para desviar los cortes de seda sin usar el filo de su katana, mientras protegía a las mujeres liberadas.

De pronto, la pared trasera del edificio colapsó cuando la caja de madera de Miyuki saltó a la acción.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: El Doble Estirpe",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: La Fuerza de Miyuki",
                    "texto": """Miyuki se lanzó directo contra Daki. Con una patada impregnada de llamas púrpuras, destruyó la faja obi principal y arrojó a la matrona demonio a través del techo hacia los tejados del distrito.

En los tejados, bajo la luz de la luna llena, la batalla se volvió abierta.

Daki gritó de rabia al ver su rostro arañado por las garras de Miyuki. La herida en su mejilla comenzó a arder por la sangre de la joven demonio.

—¡Monstruo asqueroso! ¡Arruinaste mi rostro! —chilló Daki, mientras su cuerpo comenzaba a cambiar de forma.

Ren saltó al tejado al lado de su hermana, ejecutando el primer movimiento de la Danza del Sol. Su katana se envolvió en una llamarada de fuego dorado brillante que decapitó a Daki de un solo tajo impecable.

La cabeza de la matrona rodó por las tejas de madera.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Segundo Hermano",
                    "texto": """Pero la cabeza de Daki no se convirtió en cenizas. Comenzó a llorar ruidosamente como una niña pequeña.

—¡Hermano! ¡Me cortaron la cabeza! ¡Ayúdame! —gritó Daki hacia la sombra de su propio cuerpo decapitado.

De la espalda del cuerpo de Daki emergió una figura esquelética y aterradora: un hombre encorvado de piel verdosa que sostenía dos hoces de hueso negro bañadas en veneno mortal. Era Gyutaro, el verdadero núcleo del Quinto Estirpe de Sangre.

—Nadie hace llorar a mi hermana pequeña... —gruñó Gyutaro con una voz profunda y rasposa—. ¿Crees que por tener esa marca en la cara eres especial, cazador?

Gyutaro cruzó sus hoces de hueso, desatando una ráfaga de hoces de sangre venenosa que destruyó tres manzanas del distrito de un solo golpe.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Regla de las Dos Cabezas",
                    "texto": """Aoi y el muchacho del haori amarillo llegaron al tejado para reunirse con Ren.

—¡Tienen que cortar la cabeza de los dos hermanos al mismo tiempo! —advirtió Aoi, analizando la estructura del arte demoníaco—. Si solo decapitan a uno, ninguno morirá.

Gyutaro se movió a una velocidad que superaba la percepción humana. Clavó una de sus hoces en el hombro de Ren, inyectando un veneno letal que comenzó a volver negras las venas de su cuello.

Ren cayó de rodillas sobre las tejas, sintiendo cómo el veneno paralizaba su sistema nervioso.

—Te quedan cinco minutos de vida, muchacho —se burló Gyutaro, alzando la hoz para el golpe final.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: La Purificación de la Sangre",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Veneno Frenado",
                    "texto": """Antes de que la hoz de Gyutaro alcanzara el cuello de Ren, Miyuki se interpuso.

La joven agarró la hoja de hueso con las manos desnudas. Su propia sangre demoníaca brotó de las palmas y quemó el veneno de Gyutaro con sus llamas púrpuras. Luego, Miyuki tocó la herida del hombro de Ren: sus llamas purificaron el veneno del cuerpo de su hermano en cuestión de segundos, restaurando su movilidad.

—Miyuki... gracias —dijo Ren, poniéndose de pie con renovada energía.

El chico del haori amarillo, dormido en su estado de trance del trueno, se encargó de mantener a Daki aislada en los tejados vecinos, evitando que los dos hermanos demonio pudieran reunirse.

Aoi aplicó su veneno de glicina en las hoces de Gyutaro para ralentizar su regeneración.

El contraataque de los cazadores había comenzado.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Danza del Fuego Solar",
                    "texto": """Ren ajustó la postura de sus pies sobre las tejas de madera. Inhaló aire según la secuencia de los Diarios del Sol, coordinando su ritmo cardíaco con los doce movimientos rituales de la herrería.

Su katana negra no emitió fuego común: la hoja se tiñó de una luz blanca y dorada tan intensa que iluminó todo el Distrito de las Luces como si fuera pleno mediodía.

—Danza del Sol... Cuarta Postura: arco del Sol Poniente.

Ren se convirtió en un fénix de luz dorada. Esquivó las ráfagas de hoces de sangre de Gyutaro, cortando los brazos del demonio uno por uno mientras avanzaba hacia su cuello.

Al mismo tiempo, en el tejado vecino, el chico del haori amarillo ejecutó su Destello del Trueno para cortar la cabeza de Daki por segunda vez.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Decapitamiento Simultáneo",
                    "texto": """Las dos katanas alcanzaron sus objetivos al mismo milisegundo.

La hoja dorada de Ren atravesó el cuello de Gyutaro, mientras las katanas de trueno cortaban el cuello de Daki. El impacto doble generó una explosión de luz solar que arrasó la estructura principal de la casa de entretenimientos, reduciéndola a escombros de madera y cenizas.

Las cabezas de los dos hermanos demonio rodaron juntas sobre la calle de piedra.

Esta vez, sus cuerpos comenzaron a desintegrarse en cenizas rojas que el viento de la noche dispersó sobre los canales de agua.

Por primera vez en más de cien años, un Estirpe de Sangre de los seis superiores había sido destruido por el Gremio Cuervo.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: Las Cenizas de la Misericordia (Clímax del Volumen 4)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: Las Almas del Pasado",
                    "texto": """Antes de que las cabezas de Gyutaro y Daki se convirtieran en polvo, sus espíritus permanecieron flotando en la penumbra de la calle abandonada.

Daki, reducida a la forma de la niña humana que solía ser antes de convertirse en demonio, lloraba desconsoladamente, culpando a su hermano por haberlos llevado a la pobreza en su vida humana.

—¡Te odio! ¡Ojalá nunca hubieras sido mi hermano! —chilló la niña.

Gyutaro, con el corazón roto, bajó la mirada.

—Tiene razón... si no hubiera nacido yo, ella habría tenido una vida próspera —susurró el demonio, listo para caminar solo hacia el infierno.

Ren se acercó despacio a las cenizas y se arrodilló entre ellos.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Abrazo en el Fuego",
                    "texto": """—No mientas —dijo Ren con amabilidad—. Los hermanos no se odian. Sé lo que es querer proteger a tu hermana a cualquier costo. Ella te ama, solo tiene miedo.

Daki se congeló al escuchar las palabras de Ren. Corrió hacia Gyutaro y se abrazó a su cuello, llorando sobre su pecho.

—¡Lo siento, hermano! ¡No me dejes sola! ¡Iré contigo a donde sea, incluso al infierno! —sollozó la niña.

Gyutaro la cargó en su espalda, tal como lo hacía cuando eran humanos, y ambos caminaron juntos hacia las llamas del juicio espiritual, sonriendo en paz por primera vez en siglos.

Miyuki observó la escena desde la sombra, tomando la mano de Ren con ternura.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Ira de Kageyama (Cierre del Tomo 4)",
                    "texto": """En las profundidades del Castillo Infinito —un palacio flotante fuera de la dimensión humana—, la noticia de la caída del Quinto Estirpe de Sangre sacudió la corte demoníaca.

Kageyama, el Rey Oni, en su forma de hombre elegante con ojos de gato carmesí, destruyó su laboratorio de alquimia de un solo golpe de rabia.

Alrededor de su trono, los restantes cuatro Estirpes de Sangre Superiores se arrodillaron en pánico.

—Los cazadores han despertado la Marca del Sol... —dijo Kageyama con una voz helada que hizo temblar la dimensión del castillo—. No habrá más cacerías individuales. Desplieguen a los tres Demonios del Abismo. Destruyan la Aldea de los Herreros y maten al chico de la marca.

[ CONTINUARÁ EN EL VOLUMEN 5 ]""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Vol 4 at {output_docx}")

if __name__ == "__main__":
    create_oni_vol4_manuscript()
