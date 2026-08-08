# -*- coding: utf-8 -*-
import sys
import os
import io
import json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

vol8_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-8")
vol8_dir_downloads = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-8")

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(180, 20, 20)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(120, 15, 15)
    return p

def add_image_safe(doc, img_path, width_inches=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

def build_vol8_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    r_title = p_title.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n\nVOLUMEN 8: LA NOCHE DE LOS NOVENTA MINUTOS")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(160, 20, 20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Obra Original por Nicolás Noguera\nEdición Ilustrada de Alta Definición (Dark Fantasy Anime)")
    r_sub.font.name = "Georgia"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Portada image
    add_image_safe(doc, vol8_dir / "portada.jpg", 5.0)

    doc.add_page_break()

    # Synopsis / Resumen
    add_heading_styled(doc, "SINOPSIS DEL VOLUMEN 8", 1)
    p_syn = doc.add_paragraph()
    p_syn.paragraph_format.line_spacing = 1.15
    p_syn.paragraph_format.space_after = Pt(10)
    p_syn.add_run("En las profundidades del Castillo Infinito, los cazadores enfrentan las pruebas más sangrientas contra los Tres Estirpes de Sangre. Ren descubre el Estado de Anulación de Intención contra el compás marcial de Rikudo; Aoi desata la estrategia del veneno de glicina contra Kagura; y los cazadores unen sus Hojas Rojas en la heroica caída de Kurogane. Al romper la fortaleza, Kageyama emerge en los tejados e inicia la cuenta regresiva final de noventa minutos hasta el amanecer.")

    add_image_safe(doc, vol8_dir / "banner.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 1
    add_heading_styled(doc, "CAPÍTULO 1: El Espejo de las Sombras", 1)
    add_image_safe(doc, vol8_dir / "escena_1.jpg", 6.0)
    p_c1 = doc.add_paragraph()
    p_c1.add_run("El compás de nieve azul iluminó el tatami. Rikudo avanzó lanzando ráfagas de puñetazos de energía, buscando el espíritu de lucha del joven cazador...")
    add_image_safe(doc, vol8_dir / "escena_c1_e1.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c1_e2.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c1_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 2
    add_heading_styled(doc, "CAPÍTULO 2: La Redención del Guerrero", 1)
    add_image_safe(doc, vol8_dir / "escena_c2_e1.jpg", 6.0)
    p_c2 = doc.add_paragraph()
    p_c2.add_run("Imágenes del pasado florecieron entre la niebla del combate. El guerrero demonio recordó su nombre humano, Hakuji, y la promesa hecha bajo los fuegos artificiales de verano...")
    add_image_safe(doc, vol8_dir / "escena_c2_e2.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c2_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 3
    add_heading_styled(doc, "CAPÍTULO 3: El Veneno de la Luna Llena", 1)
    add_image_safe(doc, vol8_dir / "escena_c3_e1.jpg", 6.0)
    p_c3 = doc.add_paragraph()
    p_c3.add_run("En la cámara de loto, Kagura sucumbió al veneno concentrado mientras el corte mariposa cercenaba el hielo. La derrota de las Lunas preparó la carga hacia Kurogane...")
    add_image_safe(doc, vol8_dir / "escena_c3_e2.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c3_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 4
    add_heading_styled(doc, "CAPÍTULO 4: Las Seis Ojos del Abismo", 1)
    add_image_safe(doc, vol8_dir / "escena_c4_e1.jpg", 6.0)
    p_c4 = doc.add_paragraph()
    p_c4.add_run("Las hojas solares chocaron al unísono, volviéndose incandescentes como fuego blanco. Cuatro katanas atravesaron el torso de Kurogane poniendo fin a la era del Primer Samurai...")
    add_image_safe(doc, vol8_dir / "escena_c4_e2.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c4_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 5
    add_heading_styled(doc, "CAPÍTULO 5: La Noche de los Noventa Minutos", 1)
    add_image_safe(doc, vol8_dir / "escena_c5_e1.jpg", 6.0)
    p_c5 = doc.add_paragraph()
    p_c5.add_run("El capullo de carne estalló en la plaza pública. Kageyama emergió con latigazos de materia oscura. Los Maestros Celestiales supervivientes formaron línea en la calle: comenzaba la batalla final de 90 minutos hasta el amanecer...")
    add_image_safe(doc, vol8_dir / "escena_c5_e2.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_c5_e3.jpg", 6.0)
    add_image_safe(doc, vol8_dir / "escena_climax.jpg", 6.0)

    # Save to both locations
    out1 = vol8_dir / "libro.docx"
    out2 = vol8_dir_downloads / "libro.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print(f"Generated libro.docx successfully for Volume 8 at {out1}")

if __name__ == "__main__":
    build_vol8_docx()
