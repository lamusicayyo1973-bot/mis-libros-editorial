import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

def create_manga_vol3_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("KURO NO KINEKI\n(黒の軌跡 - Ecos de Tinta Negra)\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(168, 85, 247)
    
    run_sub = title_p.add_run("Volumen 3: El Despertar de los Creadores\n\n[ GRAN FINAL DE LA TRILOGÍA ]\n\nMANGA / LIGHT NOVEL • FANTASÍA OSCURA\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(234, 88, 12)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    img_cover_path = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\kuro-no-kineki-volumen-3\escena_1.jpg"
    if os.path.exists(img_cover_path):
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_cover_path, width=Inches(5))
        
    doc.add_page_break()
    
    # TOC Header
    p_toc_head = doc.add_paragraph()
    add_bookmark(p_toc_head, "TOC", 0)
    run_toc_h = p_toc_head.add_run("Tabla de Contenidos")
    run_toc_h.font.name = "Arial"
    run_toc_h.font.size = Pt(20)
    run_toc_h.font.bold = True
    run_toc_h.font.color.rgb = RGBColor(15, 23, 42)
    end_bookmark(p_toc_head, 0)
    
    index_text = """
• Capítulo 1: El Continente de Plata
  - Escena 1: El Viaje en el Océano de Tinta
  - Escena 2: La Desembocadura del Abismo
  - Escena 3: La Costa de los Gigantes

• Capítulo 2: Los Archivos del Olvido
  - Escena 1: El Interior de la Torre
  - Escena 2: La Verdad del Primer Mundo
  - Escena 3: El Ataque de los Invasores

• Capítulo 3: La Batalla de la Tinta de Plata
  - Escena 1: El Poder del Escriba
  - Escena 2: La Fusión de los Tres Mundos
  - Escena 3: El Portal de los Reyes

• Capítulo 4: La Capital del Olvido
  - Escena 1: El Desembarco en el Continente Rojo
  - Escena 2: El Rey sin Nombre
  - Escena 3: El Gran Borrador

• Capítulo 5: El Lienzo Blanco (Gran Final de la Trilogía)
  - Escena 1: La Reunión de los Mundos
  - Escena 2: El Libro Abierto
  - Escena 3: La Última Página
"""
    p_toc = doc.add_paragraph(index_text)
    p_toc.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    chapters_content = [
        ("Capítulo 1: El Continente de Plata", [
            ("Escena 1: El Viaje en el Océano de Tinta", """
Seis meses después de la caída de Aetheria. Un barco mercante de madera reforzada navegaba por el Mar del Olvido, un océano cuyas aguas no eran de agua salada, sino de una tinta líquida plateada que reflejaba las estrellas como si fuera un espejo perfecto.

En la proa del barco, Kael observaba el horizonte con su abrigo viajero movido por el viento marino. Su ojo derecho brillaba con un patrón rúnico más complejo, capaz de leer las corrientes de memoria que fluían en el mar.

A su lado, Sora afinaba las cuerdas de un instrumento antiguo mientras Kael revisaba los mapas de navegación sobre una mesa de cubierta.

—El mapa antiguo dice que este continente fue sellado hace mil años —dijo Kael, señalando una masa de tierra rodeada de tormentas—. Nadie que haya cruzado la bruma plateada ha regresado para contarlo.

—No necesitamos un mapa completo —respondió Kael, tocando la cicatriz de su mano izquierda, que latía en dirección al centro del continente—. La energía que despertó en la torre nos está llamando. Si los "Creadores" fueron los que diseñaron las tres dagas originales, este es el lugar donde empezó nuestra maldición.
"""),
            ("Escena 2: La Desembocadura del Abismo", """
Una ola gigantesca de tinta plateada se levantó frente a la embarcación. De las aguas emergió una criatura colosal: un dragón marino hecho de cristal negro y runas doradas que emitía un rugido sónico capaz de agrietar el casco del barco.

—¡Es un Guardián de la Niebla! —gritó Sora, desenvainando la daga que había restaurado con su propio poder.

Kael reaccionó al instante, invocando un escudo de luz dorada desde sus palmas para proteger el mástil principal del barco.

Kael no esperó. Saltó hacia la cabeza del monstruo. Esta vez no necesitó una daga física: la energía de los tres elementos (fuego negro, fuego blanco y luz dorada) brotó directamente de las marcas rúnicas grabadas en sus palmas.

Con un solo movimiento cruzado en el aire, Kael trazó un símbolo de anulación que congeló al dragón de cristal en mitad del ataque, convirtiéndolo en una estatua de sal plateada que se desmoronó sobre el mar.

—Tus poderes no han dejado de evolucionar desde que absorbiste el Núcleo... —dijo Kael, mirando la facilidad con la que Kael había controlado la magia.

—No es mi poder —dijo Kael cayendo de pie en la cubierta—. Es el recuerdo de la magia original.
"""),
            ("Escena 3: La Costa de los Gigantes", """
El barco encalló suavemente en la playa de arena negra del continente olvidado. Frente a ellos no había bosques ni ciudades, sino las ruinas de estatuas de doscientos metros de altura que representaban a seres con cuatro ojos y túnicas ceremoniales: los verdaderos Creadores.

Al fondo, alzándose sobre una cordillera de montañas afiladas, se erigía la Torre de la Primera Tinta, emitiendo un pulso de luz carmesí que teñía el cielo de la tarde.

—Llegamos —dijo Sora, ajustando las vendas de sus brazos—. Siento una presencia dentro de esa torre... pero no es un autómata ni una proyección. Es un alma viva.

Kael dio el primer paso sobre la arena negra, sintiendo cómo el suelo bajo sus botas vibraba con una frecuencia familiar.

—Esa alma... es quien nos dio la vida a los tres.
""")
        ]),
        ("Capítulo 2: Los Archivos del Olvido", [
            ("Escena 1: El Interior de la Torre", """
El grupo se adentró en la estructura de la torre. El interior no estaba hecho de piedra ni de metal, sino de estanterías infinitas hechas de cristal líquido donde flotaban millones de orbes de memoria en suspensión.

Cada orbe guardaba la historia completa de un mundo que había existido antes que Aetheria.

En el centro del gran salón, flotando sobre un pedestal de runas rojas, se encontraba un ser anciano con túnicas ceremoniales rotas. Carecía de rostro humano; en su lugar, tres máscaras flotantes de oro giran a su alrededor, representando la Alegría, el Dolor y la Indiferencia.

—Los tres fragmentos han regresado a casa —habló el ser. Su voz no resonaba en el aire, sino en las mentes de los tres hermanos a la vez—. Kael... el recipiente de la tinta. Sora... la custodia de la forma. Kael... el guardián de la luz.

—¿Quién eres? —exigió Kael, poniéndose en guardia.

—Soy el Primer Escriba —respondió la entidad—. El último superviviente de los Creadores. Y quienes ustedes llaman 'humanos' no son más que los bocetos de nuestra última obra.
"""),
            ("Escena 2: La Verdad del Primer Mundo", """
El Primer Escriba levantó la mano y los orbes de memoria del salón comenzaron a girar a gran velocidad, proyectando una ilusión tridimensional en mitad de la habitación.

Kael, Sora y Kael vieron el origen de todo:
Mil años atrás, la humanidad no vivía dividida entre la Superficie y el Abismo. Vivían en una civilización próspera, pero destruida por una guerra civil catastróficamente violenta. Para evitar la extinción total, los Creadores decidieron diseñar el Sistema de Tinta: un mecanismo para borrar la capacidad de odiar de la mente humana, al costo de sacrificar los recuerdos personales.

—Creamos el Relicario y dividimos las tres dagas entre los descendientes de nuestra propia sangre para equilibrar el sistema —explicó el Escriba—. Pero Aetheria fue una prueba fallida. La Emperatriz se corrompió con el poder. Y ahora que habéis apagado el Núcleo, la barrera que protegía este mundo del resto del universo se ha roto.

—¿Qué quieres decir con 'el resto del universo'? —preguntó Sora con una punzada de pánico.

El Escriba señaló hacia el techo de la torre, que se volvió transparente, mostrando el cielo nocturno donde miles de puntos rojos comenzaban a brillar como estrellas caídas.

—Otras civilizaciones que mantuvieron sus recuerdos intactos... vienen a reclamar la Tinta Original.
"""),
            ("Escena 3: El Ataque de los Invasores", """
Antes de que Kael pudiera procesar la revelación, el techo de la torre fue atravesado por tres lanzas de energía carmesí. Un grupo de figuras con armaduras orgánicas y alas de cristal oscuro descendió desde el cielo, destruyendo las estanterías de recuerdos.

Eran los Devoradores de Ecos, guerreros del continente exterior que buscaban absorber la Tinta Original de Kael.

—¡Entrega el contenedor! —gritó el líder de los invasores, desenvainando una guadaña de fuego rojo.

Kael y Sora reaccionaron de inmediato. Kael desató una ráfaga de luz dorada que desvió la guadaña, mientras Sora trazó barreras de tinta blanca para proteger al Escriba y los archivos.

Kael apretó los puños. Las runas de sus palmas volvieron a encenderse, pero esta vez la tinta negra no salía de su cuerpo... la tinta del suelo y del mar plateado comenzó a responder a su llamado.

—No voy a permitir que nadie vuelva a robar los recuerdos de este mundo —sentenció Kael.
""")
        ]),
        ("Capítulo 3: La Batalla de la Tinta de Plata", [
            ("Escena 1: El Poder del Escriba", """
El combate en la torre se volvió un caos de energía. Los Devoradores de Ecos se movían con una coordinación perfecta, atacando en ráfagas que desgastaban la defensa de Kael y Sora.

El Primer Escriba, flotando en el centro del salón, comenzó a cantar en un idioma antiguo. Las tres máscaras doradas que flotaban sobre su cabeza se unieron en un solo rostro de piedra.

—Los bocetos no deben ser destruidos antes de que el trabajo esté terminado —dijo el Escriba.

Con un gesto de sus dedos, miles de orbes de memoria flotantes se convirtieron en agujas de luz que atravesaron a los invasores, sellando sus mentes y congelándolos en estatuas de piedra. Pero el esfuerzo fue demasiado para el anciano: su cuerpo de luz comenzó a agrietarse y disolverse.

—Escriba... —Kael corrió hacia el pedestal para sostenerlo.

—Mi tiempo ha terminado, Kael —susurró la entidad, tocando la frente del joven con su mano de piedra—. Te he transmitido la Llave del Autor. Ahora tú eres el único que puede reescribir las reglas de este mundo.
"""),
            ("Escena 2: La Fusión de los Tres Mundos", """
Al recibir la Llave del Autor, la mente de Kael se expandió más allá de los límites humanos. Por primera vez desde que despertó en la fosa del Capítulo 1, Kael no solo recuperó la totalidad de sus propios recuerdos, sino que vio el mapa completo del planeta: tres continentes flotantes aislados entre sí, cada uno sufriendo bajo sistemas de poder destructivos.

Sora y Kael se acercaron a él, sintiendo la inmensa presión mágica que emanaba de su hermano.

—¿Recuerdas todo, Kael? —preguntó Sora con la voz temblorosa.

Kael abrió los ojos. La pupila de su ojo derecho ya no tenía una sola runa, sino una galaxia entera de engranajes dorados, blancos y negros girando en perfecta armonía.

—Lo recuerdo todo, Sora. Recuerdos de nuestra infancia, el día que hicimos el pacto... y el motivo por el cual la Emperatriz nos separó —dijo Kael mirando a sus dos hermanos—. Pero más importante: ahora sé cómo salvar a los otros continentes.

Kael sonrió y levantó su espada.

—Entonces no hay tiempo que perder. La flota de los Devoradores viene en camino.
"""),
            ("Escena 3: El Portal de los Reyes", """
Desde el techo destruido de la torre, el grupo observó la armada enemiga: cientos de naves orgánicas flotando en la estratósfera, preparándose para bombardear la costa de la playa de arena negra.

Kael no mostró pavor. Extendió ambos brazos hacia el cielo.

La Llave del Autor reaccionó. El mar plateado a sus espaldas se levantó en tres columnas gigantescas de Tinta de Anulación que alcanzaron las nubes, formando un portal intercontinental de dimensiones bíblicas.

—No vamos a defender este continente —dijo Kael, mirando a Sora y Kael—. Vamos a llevar la batalla a la capital de los Devoradores.

Sora y Kael se colocaron a los lados de Kael. Los tres hermanos cruzaron sus manos en el centro del portal, desatando una luz que iluminó todo el océano plateado.
""")
        ]),
        ("Capítulo 4: La Capital del Olvido", [
            ("Escena 1: El Desembarco en el Continente Rojo", """
El portal los transportó directamente sobre la metrópolis enemiga: Ignis-Null, una ciudad de agujas de hierro negro construida dentro de un cráter volcánico. A diferencia de Aetheria, aquí no se robaban recuerdos para flotar, sino que se extraían las emociones de los ciudadanos para alimentar armas de destrucción masiva.

Al pisar la plaza central de la ciudad enemiga, la alarma general resonó con un aullido de sirenas de bronce.

Decenas de legiones de soldados de hierro los rodearon de inmediato.

—Parece que no somos bienvenidos —bromeó Kael, desenfundando su mandoble de luz.

—Déjenme la vanguardia a mí —dijo Sora, abriendo los brazos mientras miles de agujas de tinta blanca flotaban a su alrededor como un ejército de mariposas de cristal.

Kael dio un paso al frente. La Llave del Autor en su mano izquierda se transformó en una pluma de hierro oscuro capaz de alterar la materia del entorno con solo trazar líneas en el aire.
"""),
            ("Escena 2: El Rey sin Nombre", """
Desde el palacio del cráter emergió el gobernante absoluto de Ignis-Null: un ser de tres metros de altura cubierto por una armadura de obsidiana que sostenía el Relicario de la Emoción, el artefacto hermano de las tres dagas de Kael.

—Bocetos de Aetheria... —dijo el Rey sin Nombre con una voz profunda que hizo temblar la tierra—. Han venido al lugar donde se forjó la primera gota de tinta. Aquí sus armas no tienen poder.

El Rey levantó su mano y la energía de las emociones robadas de millones de personas cayó sobre Kael como una ola de gravedad pura.

Kael cayó sobre una rodilla bajo la presión mental, pero en lugar de resistirse con fuerza física, usó la enseñanza del soldado Marcus del Volumen 1: aceptar el dolor ajeno para comprenderlo.

Kael cerró los ojos y tocó el suelo del cráter.

—Las emociones no son un combustible —susurró Kael—. Son lo que nos hace mantenernos en pie cuando todo lo demás se ha perdido.
"""),
            ("Escena 3: El Gran Borrador", """
Kael trazó un único trazo horizontal en el aire con la Pluma del Autor.

La energía de la Tinta de Anulación se extendió en un círculo perfecto por todo el cráter. El efecto fue instantáneo: las armas de los soldados enemigos se convirtieron en vapor inofensivo, y las emociones robadas que alimentaban al Rey fueron liberadas de la armadura de obsidiana, regresando en forma de luces de colores hacia las mentes de la población oprimida.

El Rey de obsidiana cediño sobre sus rodillas, viendo cómo su imperio de control se desmoronaba sin violencia.

—¿Qué... qué clase de magia es esta? —preguntó el rey derrotado.

—No es magia —respondió Kael caminando hacia él—. Es el borrador que limpia el lienzo para que puedan empezar de nuevo.
""")
        ]),
        ("Capítulo 5: El Lienzo Blanco (Gran Final de la Trilogía)", [
            ("Escena 1: La Reunión de los Mundos", """
Un mes después de la caída de Ignis-Null.

Gracias a la Llave del Autor, Kael no destruyó los tres continentes, sino que los unió físicamente en un único e inmenso mundo verde rodeado por el Océano de Plata. Aetheria, el Sector Cero e Ignis-Null ahora compartían el mismo horizonte bajo un cielo sin tiranos ni sistemas de succión de memoria.

En la cima de la antigua Torre del Escriba, los tres hermanos se reunieron por última vez ante el gran libro de la creación.

Kael entregó su mandoble dorado para ser sellado en la piedra. Sora guardó sus dagas en un cofre de bronce que colocaron en la base del monumento.

—¿Qué vas a hacer con la Pluma del Autor, Kael? —preguntó Sora mirando la herramienta en las manos de su hermano.

Kael contempló la pluma por un momento y luego la arrojó hacia el centro del mar plateado, dejando que se disolviera para siempre en el agua.

—Un mundo libre no necesita a nadie que escriba su destino desde las sombras —dijo Kael sonriendo—. A partir de hoy, cada persona escribirá su propia historia.
"""),
            ("Escena 2: El Libro Abierto", """
En la plaza central de la nueva capital unificada, la gente de todos los continentes se había reunido para celebrar el inicio de la era de la Tinta Libre. Los niños jugaban en las calles dibujando con tinta de colores sobre los muros de piedra que antes separaban los mundos.

Kael asumió el liderazgo del consejo civil para coordinar el comercio y la construcción de viviendas. Sora fundó la primera Academia de la Memoria, donde se enseñaba la historia real de la humanidad sin censuras ni manipulaciones.

Y en las afueras de la ciudad, junto a un árbol de hojas doradas, Kael escribía en su cuaderno de viajes.

Ya no era el joven amnésico que despertó en la fosa llena de sangre y bronce; era el hombre que había recordado el valor de cada vida que cruzó en su camino.
"""),
            ("Escena 3: La Última Página", """
Sora se acercó despacio por el sendero y se sentó al lado de Kael en la hierba.

—¿Terminaste de escribir tu libro? —preguntó ella mirándolo con cariño.

Kael cerró el cuaderno de cuero y se lo entregó en las manos. En la portada estaba grabado en letras de luz el título original de su aventura.

Sora abrió la última página. No había dibujos ni mapas; solo una frase escrita con la letra clara de Kael:

"Los recuerdos no nos definen por lo que fuimos en el pasado, sino por lo que elegimos proteger hoy."

Kael se puso de pie, ajustándose el abrigo viajero, y miró hacia el horizonte infinito donde el sol comenzaba a ponerse sobre el nuevo mundo reunificado.

Sora sonrió, guardó el cuaderno en su bolso y se puso de pie a su lado, lista para acompañarlo en el viaje que recién comenzaba.

[ FIN DE KURO NO KINEKI — VOLUMEN 3 ]
[ FIN DE LA TRILOGÍA ]
""")
        ])
    ]
    
    b_id = 1
    for chap_title, scenes in chapters_content:
        p_head = doc.add_paragraph()
        add_bookmark(p_head, f"Chapter_{b_id}", b_id)
        
        run_h = p_head.add_run(chap_title)
        run_h.font.name = "Arial"
        run_h.font.size = Pt(18)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(168, 85, 247)
        end_bookmark(p_head, b_id)
        b_id += 1
        
        for sc_title, sc_text in scenes:
            p_sc = doc.add_paragraph()
            add_bookmark(p_sc, f"Scene_{b_id}", b_id)
            run_sc = p_sc.add_run(sc_title)
            run_sc.font.name = "Arial"
            run_sc.font.size = Pt(14)
            run_sc.font.bold = True
            run_sc.font.color.rgb = RGBColor(234, 88, 12)
            end_bookmark(p_sc, b_id)
            b_id += 1
            
            p_body = doc.add_paragraph(sc_text.strip())
            p_body.style.font.size = Pt(11)
            
            if "Lienzo Blanco" in chap_title and "Última Página" in sc_title:
                img_desc_path = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\kuro-no-kineki-volumen-3\escena_climax.jpg"
                if os.path.exists(img_desc_path):
                    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_desc_path, width=Inches(5.5))
            
    target_path = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\kuro-no-kineki-volumen-3\libro.docx"
    doc.save(target_path)
    print(f"Manga Vol 3 Word docx created: {target_path}")

if __name__ == "__main__":
    create_manga_vol3_docx()
