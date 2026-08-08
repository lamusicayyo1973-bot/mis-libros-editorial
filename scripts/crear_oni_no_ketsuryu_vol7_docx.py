# -*- coding: utf-8 -*-
import sys
import os
import io
import json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

vol7_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-7")
vol7_dir_downloads = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-7")

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(180, 20, 20)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(120, 15, 15)
    return p

def add_image_safe(doc, img_path, width_inches=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

def build_vol7_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    r_title = p_title.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n\nVOLUMEN 7: EL ASEDIO AL CASTILLO INFINITO")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(160, 20, 20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Obra Original por Nicolás Noguera\nEdición Ilustrada de Alta Definición (Dark Fantasy Anime)")
    r_sub.font.name = "Georgia"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Portada image
    add_image_safe(doc, vol7_dir / "portada.jpg", 5.0)

    doc.add_page_break()

    # Synopsis / Resumen
    add_heading_styled(doc, "SINOPSIS DEL VOLUMEN 7", 1)
    p_syn = doc.add_paragraph()
    p_syn.paragraph_format.line_spacing = 1.15
    p_syn.paragraph_format.space_after = Pt(10)
    p_syn.add_run("Atrapados en el laberinto distorsionado del Castillo Infinito, los cazadores son divididos para ser destruidos individualmente. Aoi enfrenta al Segundo Estirpe de Sangre (Kagura) en una batalla de venenos concentrados, mientras Ren, Genba, Kiri y Kazuma encaran la prueba más extrema ante Kurogane, el Primer Estirpe de Sangre de seis ojos. El colapso del palacio llevará el conflicto directo al amanecer.")

    add_image_safe(doc, vol7_dir / "banner.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 1
    add_heading_styled(doc, "CAPÍTULO 1: El Laberinto Flotante", 1)
    add_image_safe(doc, vol7_dir / "escena_1.jpg", 6.0)
    p_c1 = doc.add_paragraph()
    p_c1.add_run("Las estructuras de madera crujían invertidas bajo la gravedad distorsionada del Castillo Infinito. Los cazadores caían por pozos de escaleras sin fin mientras el sonido de la biwa resonaba en la fortaleza...")
    add_image_safe(doc, vol7_dir / "escena_c1_e1.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c1_e2.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c1_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 2
    add_heading_styled(doc, "CAPÍTULO 2: La Venganza de las Flores", 1)
    add_image_safe(doc, vol7_dir / "escena_c2_e1.jpg", 6.0)
    p_c2 = doc.add_paragraph()
    p_c2.add_run("En la cámara de loto helado, la batalla contra Kagura alcanzó su punto de no retorno. El veneno de glicina concentrado en la sangre de las cazadoras preparó el golpe definitivo...")
    add_image_safe(doc, vol7_dir / "escena_c2_e2.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c2_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 3
    add_heading_styled(doc, "CAPÍTULO 3: El Hielo Descompuesto", 1)
    add_image_safe(doc, vol7_dir / "escena_c3_e1.jpg", 6.0)
    p_c3 = doc.add_paragraph()
    p_c3.add_run("Genba y Kazuma unieron fuerzas frente a la cámara del Primer Estirpe. Kurogane desenfundó su espada de carne viva, liberando un mar de ráfagas en forma de luna creciente...")
    add_image_safe(doc, vol7_dir / "escena_c3_e2.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c3_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 4
    add_heading_styled(doc, "CAPÍTULO 4: La Sala de las Seis Lunas", 1)
    add_image_safe(doc, vol7_dir / "escena_c4_e1.jpg", 6.0)
    p_c4 = doc.add_paragraph()
    p_c4.add_run("La bola de picos de hierro de Genba chocó contra la hoja carmesí de Kurogane. La combinación del Fuego Solar y el Viento forzó al demonio a desatar su transformación final...")
    add_image_safe(doc, vol7_dir / "escena_c4_e2.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c4_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 5
    add_heading_styled(doc, "CAPÍTULO 5: La Caída del Primer Samurai", 1)
    add_image_safe(doc, vol7_dir / "escena_c5_e1.jpg", 6.0)
    p_c5 = doc.add_paragraph()
    p_c5.add_run("Con tres katanas al rojo vivo atravesando su torso, el Primer Estirpe vio desintegrar su cuerpo. Toda la fortaleza comenzó a elevarse hacia la superficie mientras aparecía la primera luz del amanecer...")
    add_image_safe(doc, vol7_dir / "escena_c5_e2.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_c5_e3.jpg", 6.0)
    add_image_safe(doc, vol7_dir / "escena_climax.jpg", 6.0)

    # Save to both locations
    out1 = vol7_dir / "libro.docx"
    out2 = vol7_dir_downloads / "libro.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print(f"Generated libro.docx successfully for Volume 7 at {out1}")

if __name__ == "__main__":
    build_vol7_docx()
