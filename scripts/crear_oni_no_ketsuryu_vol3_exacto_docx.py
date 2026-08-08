import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00) # Dark Red
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slate
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add XML Bookmark
    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_oni_vol3_manuscript():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-3"
    
    # 1. PORTADA
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    # 2. TITULO PRINCIPAL
    add_heading_with_bookmark(doc, "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 3: El Tren de las Sombras", 2, 1, "subtitulo_vol3")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # 3. TABLA DE CONTENIDOS (TOC KDP COMPATIBLE)
    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: La Carne del Tren", "cap_1"),
        ("Capítulo 2: El Cuello de Hierro", "cap_2"),
        ("Capítulo 3: La Llegada del Tercer Lunar", "cap_3"),
        ("Capítulo 4: La Última Llama", "cap_4"),
        ("Capítulo 5: La Herencia del Fuego (Clímax del Volumen 3)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    # ESTRUCTURA DE TEXTO CON IMÁGENES EXACTAS POR ESCENA
    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: La Carne del Tren",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: El Despertar del Pilar",
                    "texto": """El calor dentro del vagón se volvió sofocante. Las venas de carne demoníaca que cubrían los asientos se retorcían con violencia, intentando envolver a los pasajeros dormidos.

De pronto, un estallido de fuego dorado rasgó la niebla.

Kenshin, el Pilar del Fuego, rompió la ilusión del sueño. Su haori con patrones de llamas se agitó mientras desenvainaba su katana de tono anaranjado brillante. La temperatura del aire subió diez grados en un segundo.

—¡Joven Hagane! —retumbó la voz de Kenshin con un entusiasmo contagioso—. ¡Excelente trabajo al cortar tu propio cuello en el sueño para despertar! Tu resolución es digna de un verdadero cazador.

—¡El demonio es la locomotora entera! —explicó Ren, esquivando las tentáculos de carne que brotaban del techo—. ¡Si cortamos los vagones, no servirá de nada! ¡Su cuello debe estar cerca de la caldera de carbón!

—¡Entendido! —sonrió Kenshin, apoyando la mano en la empuñadura—. Tú y la joven demonio protejan los ocho vagones de pasajeros. ¡Yo me abriré paso hasta el núcleo del tren!""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Defensa de los Pasajeros",
                    "texto": """Kenshin se desplazó a una velocidad sobrehumana, dejando una estela de fuego dorado a lo largo del pasillo que incineró la carne demoníaca de los primeros cuatro vagones.

Mientras tanto, en la parte trasera del tren, cientos de bocas humanas brotaron de las paredes metálicas para devorar a los civiles dormidos.

La caja de madera de Ren se abrió de golpe. Miyuki saltó al pasillo con los ojos encendidos en una luz violeta. Con una agilidad pasmosa, la joven usó sus garras para cortar la masa de carne, mientras su sangre ardía en llamas púrpuras que sellaban las heridas del metal para que no volvieran a crecer.

Ren luchaba espalda con espalda junto a su hermana. Cada tajo de su katana negra, impulsado por la Respiración de Sangre, trazaba arcos de fuego rojo que mantenían a salvo a los pasajeros.

—¡No dejaré que nadie más pierda a su familia esta noche! —gritó Ren, aumentando el ritmo de su corazón.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Frente de la Locomotora",
                    "texto": """Ren corrió por el techo del tren bajo la luz de la luna llena, esquivando las ráfagas de viento y los ojos de cera que brotaban del metal.

Al llegar al primer vagón, se encontró con el joven asustadizo de la Selección Final —el muchacho de la doble katana corta que llevaba un haori amarillo con bordados de rayos—. Estaba paralizado de terror sobre la caldera, temblando mientras el vapor hirviente salía por las tuberías.

—¡No puedo hacerlo! ¡Voy a morir! ¡Es un Lunar Rojo! —gritaba el joven del haori amarillo entre lágrimas.

De pronto, un rayo de luz blanca atravesó el aire. El muchacho se quedó dormido de pie por el pánico absoluto. En ese estado de trance, su postura cambió por completo: adoptó una posición de esgrima perfecta y una corriente de aire electrificado comenzó a sonar alrededor de sus espaldas.

Era la Respiración del Trueno Inverso.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: El Cuello de Hierro",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: Velocidad del Rayo",
                    "texto": """El chico del haori amarillo se desplazó en un milisegundo.

—Respiración del Trueno... Primera Postura: Chispazo del Destello.

Un destello de luz amarilla cortó la noche. Las decenas de brazos de carne que intentaban sujetar a Ren fueron rebanadas instantáneamente por las dos katanas cortas del joven dormido.

Ren aprovechó la brecha abierta por su compañero. Saltó hacia la escotilla de la caldera de carbón, donde la masa de carne demoníaca se concentraba en un núcleo espeso que latía con el pulso del Sexto Lunar Rojo.

—¡Ahí está el cuello! —gritó Ren.

Desde la masa de carne, la cara del demonio pálido emergió, mostrando las bocas en sus palmas de las manos.

—¡Es inútil, cazadores de poca monta! —siseó el Sexto Lunar—. ¡Incluso si cortan este punto, mi masa corporal regenerará la caldera en un segundo!""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Fusión de los Estilos",
                    "texto": """Ren inhaló aire hasta el fondo de sus pulmones. Las venas negras de su rostro se extendieron hasta la sien derecha, brillando con una luz carmesí brillante.

—¡No si lo hacemos al mismo tiempo! —respondió Ren.

El joven del haori amarillo, aun dormido, se posicionó al lado izquierdo de la caldera, cargando sus katanas con electricidad blanca. Ren se colocó a la derecha, envolviendo su katana negra en llamas rojas de fuego de herrería.

—¡Respiración del Trueno: Destello Doble!
—¡Respiración de Sangre: Tajo del Horno Encendido!

Los dos ataques impactaron el núcleo de carne de la caldera exactamente al mismo tiempo. La combinación del calor del fuego y la velocidad de la electricidad destrozó la estructura de masa demoníaca sin darle tiempo a regenerarse.""",
                    "img": "banner.jpg"
                },
                {
                    "nombre": "Escena 3: El Descarrilamiento",
                    "texto": """El grito agudo del Sexto Lunar Rojo resonó en toda la montaña mientras la masa orgánica del tren comenzaba a disolverse en cenizas.

Sin la carne demoníaca para sostener la estructura, los frenos de la locomotora colapsaron. El tren de vapor de miles de toneladas comenzó a salirse de las vías a más de ochenta kilómetros por hora, volcando sobre el campo abierto.

Ren usó su cuerpo como escudo para proteger al muchacho del haori amarillo y a la caja de Miyuki mientras el metal se retorcía sobre la hierba.

Kenshin, desde el interior, usó su Respiración del Fuego para crear una burbuja de aire caliente que amortiguó el impacto de todos los vagones de pasajeros, salvando la vida de los doscientos civiles a bordo.

El vapor se dispersó sobre el campo bajo la luz del amanecer. El peligro parecía haber pasado.""",
                    "img": "escena_1.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: La Llegada del Tercer Lunar",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: La Presencia Absoluta",
                    "texto": """La victoria se esfumó en un segundo.

El aire sobre el campo de hierba se volvió tan denso y pesado que Ren cayó sobre sus rodillas, incapaz de respirar. El suelo bajo sus pies comenzó a agrietarse sin razón aparente.

Desde las sombras de los árboles de la linde del bosque, un hombre de torso desnudo caminó despacio hacia el lugar del accidente. Su piel era pálida como la nieve, estaba cubierta por tatuajes rúnicos de color azul marino y sus ojos brillaban con un tono dorado donde se leía claramente la marca del Tercer Lunar Rojo.

Su sola presencia emitía un aura de presión física que aplastaba la hierba a su alrededor.

—Una técnica de fuego bastante pulida... —dijo el Tercer Lunar, mirando a Kenshin con una sonrisa llena de anticipación—. Hacía más de cincuenta años que no encontraba a un Pilar de la Respiración del Fuego.""",
                    "img": "escena_climax.jpg"
                },
                {
                    "nombre": "Escena 2: La Propuesta del Demonio",
                    "texto": """Kenshin dio un paso al frente, interponiéndose entre el Tercer Lunar y los jóvenes heridos. Aunque estaba agotado por mantener a salvo a los pasajeros, su postura con la katana de fuego era impecable.

—Soy Kenshin, Pilar del Fuego del Gremio Cuervo —dijo con voz serena.

—Soy Rikudo —respondió el demonio, ajustando sus puños en una posición de artes marciales ancestrales—. Tienes un espíritu de combate supremo, Kenshin. Tu cuerpo humano alcanzará su límite en unos años; te volverás viejo, débil y morirás. Conviértete en un demonio. Con el poder de la sangre de nuestro Señor, podrás entrenar tu técnica por toda la eternidad.

Kenshin no dudó ni un segundo. Sonrió con amabilidad y negó con la cabeza.

—Envejecer y morir es lo que le da valor y belleza a la vida humana. Precisamente porque somos efímeros, nuestro esfuerzo es sagrado. No me convertiré en un monstruo.""",
                    "img": "banner.jpg"
                },
                {
                    "nombre": "Escena 3: El Choque de los Titanes",
                    "texto": """Sin decir una palabra más, Rikudo se lanzó al ataque.

El choque entre los puños de energía del demonio y la katana de fuego de Kenshin levantó una onda de choque que destruyó los restos del primer vagón de tren. La velocidad del combate era tan absurda que Ren solo podía ver destellos de luz dorada y azul cruzando el campo.

—¡Arte Demoníaco: Aguja de la Muerte! —gritó Rikudo, desatando una ráfaga de cientos de golpes invisibles.

—¡Respiración del Fuego: Quinta Postura: Tigre de las Llamas! —respondió Kenshin, creando un felino gigante de fuego dorado que devoró la ráfaga del enemigo.

El suelo tembló mientras la batalla alcanzaba un nivel de destrucción destructivo.""",
                    "img": "escena_climax.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: La Última Llama",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Límite del Cuerpo Humano",
                    "texto": """A pesar de la maestría de Kenshin, la regeneración del Tercer Lunar era instantánea. Cada corte hecho por la katana de fuego se cerraba en milisegundos, mientras que las heridas en el cuerpo de Kenshin comenzaban a acumularse.

El Pilar del Fuego sangraba por la frente y tenía dos costillas rotas, pero su mirada dorada no mostraba una sola gota de miedo.

Ren intentó levantarse para ayudarlo, agarrando su katana negra, pero Rikudo ni siquiera se giró; usó la presión de su aura para lanzar a Ren contra el suelo.

—¡No te metas, chico! —gritó Kenshin sin perder la sonrisa—. ¡Mantén la posición y protege a tu hermana! ¡Este es el deber de un Pilar!

Kenshin inhaló una cantidad monumental de aire, haciendo crujir sus pulmones mientras el fuego de su espada se volvía de un color blanco incandescente.""",
                    "img": "escena_1.jpg"
                },
                {
                    "nombre": "Escena 2: La Novena Postura",
                    "texto": """—Respiración del Fuego... Novena Postura: Purgatorio.

Kenshin se convirtió en un meteoro de fuego blanco que arrasó la hierba del campo. Se lanzó directo al pecho de Rikudo en una estocada definitiva.

El impacto fue ensordecedor. Rikudo clavó su puño derecho en el torso de Kenshin, pero al mismo tiempo, la katana de fuego de Kenshin atravesó el cuello del Tercer Lunar hasta la mitad.

Rikudo intentó retirar el puño para decapitar al Pilar, pero Kenshin usó sus propios músculos abdominales y su fuerza física sobrehumana para atrapar el brazo del demonio dentro de su cuerpo.

—¡No te moverás de aquí hasta que el sol aparezca! —gritó Kenshin, apretando el cuello de Rikudo con la mano izquierda desnuda.

En el horizonte, los primeros rayos del sol matutino comenzaron a asomar sobre la cumbre de las montañas.""",
                    "img": "escena_climax.jpg"
                },
                {
                    "nombre": "Escena 3: La Huida de la Sombra",
                    "texto": """Al sentir el calor del primer rayo de sol sobre su piel, Rikudo entró en pánico absoluto. La piel de sus hombros comenzó a quemarse y desintegrarse en cenizas.

Con un alarido de desesperación, el Tercer Lunar usó una fuerza bruta desmedida para arrancarse sus propios brazos, liberándose del agarre de Kenshin. Se arrojó de cabeza hacia las sombras del bosque espeso antes de que la luz solar lo alcanzara por completo.

Ren, lleno de una ira ciega, tomó su katana negra y se la arrojó con todas sus fuerzas por la espalda.

La hoja de Ren atravesó el pecho de Rikudo justo cuando el demonio desaparecía en la penumbra del bosque.

—¡Cobarde! —gritó Ren con lágrimas en los ojos, corriendo hacia el bosque—. ¡No huyas! ¡El señor Kenshin no huyó! ¡Él ganó esta batalla porque protegía a los demás! ¡Tú solo eres una bestia miedosa que le teme al sol!""",
                    "img": "banner.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: La Herencia del Fuego (Clímax del Volumen 3)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: Las Últimas Palabras del Pilar",
                    "texto": """El sol matutino bañó el campo de hierba.

Kenshin permanecía de pie, apoyado en la empuñadura de su katana de fuego clavada en la tierra. Su figura erguida proyectaba una sombra larga sobre el suelo.

Ren y Miyuki corrieron a su lado, cayendo de rodillas. Ren intentaba presionar las heridas del Pilar con sus manos sangrantes, pero Kenshin le colocó suavemente la mano sobre la cabeza para detenerlo.

—Ya es suficiente, joven Hagane... —dijo Kenshin con voz débil pero cálida—. Mi tiempo ha llegado.

—Señor Kenshin... lo siento... no pude ayudarlo... —sollozó Ren, apoyando la frente en la hierba.

—No te disculpes —sonrió Kenshin, mirando al cielo—. Salvaste a los pasajeros de los vagones. Tu hermana luchó con corazón humano. Confío en ustedes dos. Vayan a la finca de mi familia... lean los escritos antiguos de los anteriores Pilares del Fuego... ahí encontrarán la pista sobre la Respiración del Sol.""",
                    "img": "escena_1.jpg"
                },
                {
                    "nombre": "Escena 2: La Visión de la Madre",
                    "texto": """Kenshin alzó la mirada hacia la luz del sol. En mitad de la bruma dorada de la mañana, vio la figura de su difunta madre sonriéndole con ternura.

—¿Lo hice bien, madre? —susurró Kenshin—. ¿Cumplí con el deber con el que nací?

La visión de su madre asintió con la cabeza, extendiéndole los brazos.

—Lo hiciste maravillosamente, mi querido hijo.

Kenshin cerró los ojos con una sonrisa radiante en el rostro. Su corazón dejó de latir, pero su cuerpo permaneció erguido como una estatua inamovible en medio del campo de flores.

Ren y Miyuki inclinaron la cabeza en un silencio profundo, honrando al hombre que había dado su vida para proteger la de ellos.""",
                    "img": "escena_climax.jpg"
                },
                {
                    "nombre": "Escena 3: La Decisión del Guerrero (Cierre del Tomo 3)",
                    "texto": """Tres días después.

En el cementerio del Gremio Cuervo, los cuervos mensajeros llevaron la noticia de la caída del Pilar del Fuego a todos los rincones del país. Los restantes ocho Pilares recibieron el mensaje en silencio, apretando sus armas con una nueva determinación.

Ren caminaba por el sendero hacia la Finca del Fuego, llevando la guarda de la katana con forma de llama de Kenshin enganchada en su propia espada negra.

Las venas de su rostro habían cambiado: ya no eran líneas desordenadas, sino que habían tomado la forma de una llama carmesí sobre su mejilla derecha.

—No dejaré que tu sacrificio sea en vano, Kenshin —dijo Ren mirando al cielo azul—. Voy a dominar la Respiración del Sol y voy a cortar la cabeza de Kageyama con mis propias manos.

[ CONTINUARÁ EN EL VOLUMEN 4 ]""",
                    "img": "banner.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    # Guardar manuscrito final
    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx at {output_docx}")

if __name__ == "__main__":
    create_oni_vol3_manuscript()
