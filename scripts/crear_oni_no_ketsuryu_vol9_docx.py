# -*- coding: utf-8 -*-
import sys
import os
import io
import json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

vol9_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-9")
vol9_dir_downloads = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-9")

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(180, 20, 20)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(120, 15, 15)
    return p

def add_image_safe(doc, img_path, width_inches=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

def build_vol9_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    r_title = p_title.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n\nVOLUMEN 9: LA BATALLA DEL AMANECER")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(160, 20, 20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Obra Original por Nicolás Noguera\nEdición Ilustrada de Alta Definición (Dark Fantasy Anime)")
    r_sub.font.name = "Georgia"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Portada image
    add_image_safe(doc, vol9_dir / "portada.jpg", 5.0)

    doc.add_page_break()

    # Synopsis / Resumen
    add_heading_styled(doc, "SINOPSIS DEL VOLUMEN 9", 1)
    p_syn = doc.add_paragraph()
    p_syn.paragraph_format.line_spacing = 1.15
    p_syn.paragraph_format.space_after = Pt(10)
    p_syn.add_run("El Castillo Infinito colapsa hacia la superficie de la capital imperial. Kageyama emerge en su forma final con látigos de carne y bocas aberrantes. Mientras los Maestros Celestiales luchan contra la parálisis del veneno y la droga de envejecimiento de Sumire cobra efecto, Ren conecta las doce posturas del Estilo de Dominio Solar en la Decimotercera Postura para fijar a la masa gigante del Rey Demonio hasta el primer rayo de luz del sol.")

    add_image_safe(doc, vol9_dir / "banner.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 1
    add_heading_styled(doc, "CAPÍTULO 1: La Emergencia a la Ciudad", 1)
    add_image_safe(doc, vol9_dir / "escena_1.jpg", 6.0)
    p_c1 = doc.add_paragraph()
    p_c1.add_run("Los escombros de la fortaleza cayeron sobre la plaza principal. Kageyama emergió en el centro del cráter, agitando látigos espinados de materia oscura que cortaban estructuras de piedra a velocidad hipersónica...")
    add_image_safe(doc, vol9_dir / "escena_c1_e1.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c1_e2.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c1_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 2
    add_heading_styled(doc, "CAPÍTULO 2: El Veneno en la Sangre", 1)
    add_image_safe(doc, vol9_dir / "escena_c2_e1.jpg", 6.0)
    p_c2 = doc.add_paragraph()
    p_c2.add_run("La toxina creada por la Dra. Sumire comenzó a envejecer aceleradamente las células del Rey Demonio. Genba y Kazuma mantuvieron la presión en las calles bajo la noche estrellada...")
    add_image_safe(doc, vol9_dir / "escena_c2_e2.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c2_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 3
    add_heading_styled(doc, "CAPÍTULO 3: La Secuencia del Sol", 1)
    add_image_safe(doc, vol9_dir / "escena_c3_e1.jpg", 6.0)
    p_c3 = doc.add_paragraph()
    p_c3.add_run("Ren encadenó las doce posturas solares en un dragón de fuego continuo. El reloj de la plaza central marcaba los últimos quince minutos antes del amanecer...")
    add_image_safe(doc, vol9_dir / "escena_c3_e2.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c3_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 4
    add_heading_styled(doc, "CAPÍTULO 4: La Luz en el Horizonte", 1)
    add_image_safe(doc, vol9_dir / "escena_c4_e1.jpg", 6.0)
    p_c4 = doc.add_paragraph()
    p_c4.add_run("Kageyama se expandió en una masa monstruosa tratando de enterrarse en el suelo. Los guerreros de la Hermandad tiraron de las cadenas de hierro para mantenerlo expuesto...")
    add_image_safe(doc, vol9_dir / "escena_c4_e2.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c4_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 5
    add_heading_styled(doc, "CAPÍTULO 5: La Estocada del Sol", 1)
    add_image_safe(doc, vol9_dir / "escena_c5_e1.jpg", 6.0)
    p_c5 = doc.add_paragraph()
    p_c5.add_run("Los primeros rayos dorados del sol tocaron la cima de los edificios. La masa demoníaca se desintegró en chispas de luz, trayendo la victoria definitiva...")
    add_image_safe(doc, vol9_dir / "escena_c5_e2.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_c5_e3.jpg", 6.0)
    add_image_safe(doc, vol9_dir / "escena_climax.jpg", 6.0)

    # Save to both locations
    out1 = vol9_dir / "libro.docx"
    out2 = vol9_dir_downloads / "libro.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print(f"Generated libro.docx successfully for Volume 9 at {out1}")

if __name__ == "__main__":
    build_vol9_docx()
