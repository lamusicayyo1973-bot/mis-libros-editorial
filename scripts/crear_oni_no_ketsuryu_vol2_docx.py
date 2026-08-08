import os
import glob
import shutil
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

src_dir = r"C:\Users\nicol\.gemini\antigravity\brain\6adf8ce5-9839-4292-a8a1-57beed4c3827"
target_root = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-2"
target_sys = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\oni-no-ketsuryu-volumen-2"

os.makedirs(target_root, exist_ok=True)
os.makedirs(target_sys, exist_ok=True)

mapping = {
    'oni_vol2_portada*.jpg': 'portada.jpg',
    'oni_vol2_thumbnail*.jpg': 'thumbnail.jpg',
    'oni_vol2_banner*.jpg': 'banner.jpg'
}

for pattern, target_name in mapping.items():
    matches = glob.glob(os.path.join(src_dir, pattern))
    if matches:
        matches.sort(key=os.path.getmtime, reverse=True)
        latest = matches[0]
        shutil.copy2(latest, os.path.join(target_root, target_name))
        shutil.copy2(latest, os.path.join(target_sys, target_name))
        print(f"Copied {target_name}")

# Duplicate banner as escena placeholders if specific scene images hit quota limit
banner_p = os.path.join(target_root, "banner.jpg")
if os.path.exists(banner_p):
    if not os.path.exists(os.path.join(target_root, "escena_1.jpg")):
        shutil.copy2(banner_p, os.path.join(target_root, "escena_1.jpg"))
        shutil.copy2(banner_p, os.path.join(target_sys, "escena_1.jpg"))
    if not os.path.exists(os.path.join(target_root, "escena_climax.jpg")):
        shutil.copy2(banner_p, os.path.join(target_root, "escena_climax.jpg"))
        shutil.copy2(banner_p, os.path.join(target_sys, "escena_climax.jpg"))

def create_oni_vol2_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(220, 38, 38)
    
    run_sub = title_p.add_run("Volumen 2: El Examen de la Montaña Sombría\n\nMANGA / LIGHT NOVEL • FANTASÍA OSCURA & ACCIÓN SENGOKU\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(185, 28, 28)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    img_cover_path = os.path.join(target_root, "portada.jpg")
    if os.path.exists(img_cover_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_cover_path, width=Inches(4.5))
        
    doc.add_page_break()
    
    # TOC Header
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = toc_p.add_run("TABLA DE CONTENIDOS\n")
    r_toc.font.name = "Arial"
    r_toc.font.size = Pt(20)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(220, 38, 38)
    
    chapters_toc = [
        ("Capítulo 1", "La Cuvo-Espada de la Noche", "chap1"),
        ("Capítulo 2", "El Pueblo de los Faroles", "chap2"),
        ("Capítulo 3", "La Danza de la Tinta Roja", "chap3"),
        ("Capítulo 4", "El Tren de las Sombras", "chap4"),
        ("Capítulo 5", "La Trampa de los Sueños (Clímax del Volumen 2)", "chap5")
    ]
    
    for c_num, c_title, c_bkm in chapters_toc:
        p_t = doc.add_paragraph()
        r_num = p_t.add_run(f"• {c_num}: ")
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(185, 28, 28)
        r_name = p_t.add_run(c_title)
        r_name.font.italic = True
        
    doc.add_page_break()
    
    # Full Content Text
    content = [
        {
            "id": "chap1",
            "number": "Capítulo 1",
            "title": "La Cuvo-Espada de la Noche",
            "scenes": [
                {
                    "num": "Escena 1: El Regreso de la Selección",
                    "text": """El amanecer iluminó los picos del monte Fujikane. De los más de sesenta aspirantes que ingresaron a la prueba, solo cuatro permanecieron de pie en la salida rodeada de glicinas púrpuras: Ren, el joven de las dos katanas cortas, la chica silenciosa de la máscara de zorro y un muchacho asustadizo que no paraba de temblar.

Las dos niñas con rostros de porcelana los recibieron con una reverencia formal.

—Felicitaciones a los sobrevivientes —dijeron al unísono—. Ahora son oficialmente Guerreros de la Hermandad de Rango Mizunoto, el escalafón inicial del Gremio Cuervo.

Dos asistentes del gremio se acercaron portando una mesa de madera. Sobre ella descansaban varios bloques de mineral Tamahagane-Sombra, el metal especial que absorbe la luz del sol.

—Elijan el bloque que moldeará su katana definitiva —indicó una de las niñas—. Cada hoja responderá al color del alma de su portador.

Ren extendió su mano y eligió un bloque oscuro que emitía un calor sutil, idéntico al mineral de la antigua herrería de su padre."""
                },
                {
                    "num": "Escena 2: El Cuervo Mensajero",
                    "text": """Un graznido agudo resonó sobre el cielo de la mañana. Un cuervo de plumaje negro azabache, portando una pequeña placa de bronce en el pecho, descendió y se posó sobre el hombro derecho de Ren.

—¡Ren Hagane! ¡Ren Hagane! —graznó la ave con una voz metálica e inesperadamente humana—. ¡Tu primera misión te aguarda! ¡Misión de reconocimiento en el Pueblo de los Faroles! ¡Jóvenes están desapareciendo cada noche en la penumbra!

Ren ajustó la caja de madera de su espalda, donde Miyuki descansaba del sol de la mañana.

—Entendido —dijo Ren, apretando el mango de su katana provisoria—. Nos ponemos en marcha inmediatamente."""
                },
                {
                    "num": "Escena 3: La Entrega del Acero Negro",
                    "img": "escena_1.jpg",
                    "text": """Antes de partir al Pueblo de los Faroles, un herrero excéntrico del gremio —cubierto con un sombrero de paja y una máscara de Hyottoko— alcanzó a Ren en el camino de la montaña.

Llevaba un envoltorio de tela roja. Al desplegarlo, reveló la nueva katana de Ren: la hoja era de un negro azabache profundo, pero a lo largo del filo corría una fina línea carmesí que brillaba al contacto con la respiración de Ren.

—Es una Hoja de Tinta y Fuego —explicó el herrero, señalando el metal—. Absorbió el fragmento de cristal de tu antigua arma rota. Si tu espíritu flaquea, la hoja te consumirá. Si tu espíritu es fuerte, cortará el acero de los demonios más antiguos.

Ren desenvainó la espada. El peso era perfecto; se sentía como una extensión natural de su propio brazo."""
                }
            ]
        },
        {
            "id": "chap2",
            "number": "Capítulo 2",
            "title": "El Pueblo de los Faroles",
            "scenes": [
                {
                    "num": "Escena 1: Las Sombras del Mercado",
                    "text": """El Pueblo de los Faroles era una villa próspera conocida por sus festivales nocturnos. Cientos de faroles de papel rojo y amarillo iluminaban las calles de piedra mientras la gente caminaba vistiendo yukatas de colores.

Pero tras la fachada festiva, Ren sintió el olor fétido a azufre y sangre rancia flotando en el aire.

—Las desapariciones ocurren siempre cuando la última vela del festival se apaga —murmuró Ren, caminando por un callejón oscuro alejado de la multitud.

De pronto, la caja de madera en su espalda comenzó a golpear suavemente. Miyuki estaba despierta y alertando a su hermano sobre una amenaza inminente desde los tejados."""
                },
                {
                    "num": "Escena 2: El Demonio de la Cera",
                    "text": """Desde la sombra de una vivienda, la pared de madera se derritió como si fuera cera caliente. Una criatura de tres metros hecha de sustancia viscosa y brillante emergió, atrapando a un joven poblador con sus tentáculos de cera hirviente.

—¡Un nuevo cazador de carne tierna! —siseó el demonio, mientras su rostro sin ojos se abría revelando hileras de dientes de cristal.

Ren no dudó. Desenvainó su katana negra y se lanzó al ataque. La cera del demonio intentó envolver su hoja, pero al contacto con la Estilo de Esgrima de Sangre, el calor de la katana evaporó la sustancia al instante.

Sin embargo, la criatura se regeneraba a una velocidad pasmosa: cada corte hecho por Ren se cerraba en una fracción de segundo.

—¡Las espadas comunes no pueden cortar lo que se derrite! —se burló el demonio, atrapando el tobillo de Ren."""
                },
                {
                    "num": "Escena 3: El Despertar de Miyuki",
                    "text": """Antes de que el demonio pudiera aplastar a Ren contra el suelo, la tapa de la caja de madera voló por los aires.

Miyuki saltó al combate. Con el bambú aún en la boca, dio una patada voladora con una fuerza descomunal que destrozó el torso de cera del demonio, arrojándolo diez metros hacia atrás. La sangre de la joven, al caer sobre la cera, comenzó a arder en llamas púrpuras que impidieron la regeneración de la monstruosidad.

—¡Sangre de Demonio Explosiva! —comprendió Ren al ver la habilidad de su hermana.

Miyuki miró a Ren con sus ojos de pupila rasgada, dándole una señal con la cabeza.

Hermano y hermana estaban luchando juntos por primera vez en perfecta sincronía."""
                }
            ]
        },
        {
            "id": "chap3",
            "number": "Capítulo 3",
            "title": "La Danza de la Tinta Roja",
            "scenes": [
                {
                    "num": "Escena 1: La Purificación del Fuego",
                    "text": """Aprovechando que las llamas púrpuras de Miyuki congelaron la regeneración del monstruo, Ren inhaló una bocanada profunda de aire, llenando sus pulmones hasta el límite.

Las venas negras de su rostro brillaron con intensidad.

—Estilo de la Sangre Negra... Tercera Postura: Espiral del Horno Encendido.

Ren giró sobre su propio eje a una velocidad cegadora. Su katana trazó una espiral de fuego rojo y sombras que cortó al demonio en tres partes principales, alcanzando el núcleo de su cuello.

La criatura emitió un alarido sónico antes de disolverse en cenizas de fuego rojo que el viento de la noche dispersó sobre los tejados."""
                },
                {
                    "num": "Escena 2: Los Recuerdos de la Cera",
                    "text": """Antes de que las cenizas del demonio desaparecieran por completo, la Estilo de Esgrima de Sangre de Ren le permitió percibir los últimos pensamientos de la criatura.

Vio la visión de un anciano fabricante de velas de la era Edo, abandonado por su familia en la pobreza y tentado por un hombre elegante que vestía un sombrero blanco de época y trajo la maldición sobre su cuerpo.

—Kageyama... —murmuró Ren, escuchando el nombre del Rey Oni por primera vez en la mente de la víctima.

Miyuki se acercó despacio a las cenizas y colocó su mano sobre el lugar donde cayó el demonio, demostrando compasión por el alma que alguna vez fue humana.

Ren guardó su katana en la funda. La primera misión había concluido, pero la verdadera cacería recién comenzaba."""
                },
                {
                    "num": "Escena 3: La Sombra en la Niebla",
                    "text": """Desde lo alto de la torre del reloj del pueblo, una figura envuelta en una túnica oscura con bordados de telaraña observaba la escena.

Llevaba un biwa (instrumento musical japonés) en sus manos y la pupila de su ojo izquierdo estaba marcada con el número del Sexto Estirpe de Sangre.

—El chico de las venas negras y la niña que no consume carne... —susurró la entidad, tocando una cuerda del biwa—. El señor Estará muy complacido de saber que la estirpe sobrevivió.

Con un sonido seco de cuerda (¡TONG!), la figura desapareció en el aire, dejando solo un rastro de humo negro."""
                }
            ]
        },
        {
            "id": "chap4",
            "number": "Capítulo 4",
            "title": "El Tren de las Sombras",
            "scenes": [
                {
                    "num": "Escena 1: El Cuervo de la Segunda Misión",
                    "text": """Al amanecer del día siguiente, el cuervo mensajero de Ren descendió agitado sobre la barandilla de la posada.

—¡Nuevas órdenes! ¡Nuevas órdenes! —gritó la ave—. ¡Abordar el Tren de la Noche en la estación central! ¡Más de cuarenta pasajeros han desaparecido dentro de los vagones! ¡Un Maestro Celestial del Gremio ya está a bordo!

Ren miró a Miyuki, quien descansaba dentro de la caja de madera.

—Un Maestro Celestial... —dijo Ren—. Los guerreros más poderosos del Gremio Cuervo. Esta es nuestra oportunidad de aprender cómo derrotar a los Seis Lunares."""
                },
                {
                    "num": "Escena 2: La Estación del Vapor",
                    "text": """La estación de trenes de la ciudad estaba cubierta por una densa nube de vapor blanco producido por la gigantesca locomotora de hierro negro.

Ren caminaba por el andén entre la multitud, impresionado por la tecnología de la máquina de vapor. Al subir al vagón de pasajeros, el aire dentro del tren se sentía anormalmente pesado, como si el metal de la locomotora estuviera vivo.

En el asiento central del vagón de primera clase, un hombre de hombros anchos vestía un haori con patrones de llamas doradas. Comía bento de carne a una velocidad pasmosa, gritando "¡Sabroso!" con cada bocado.

Era Kenshin, el Maestro Celestial del Fuego."""
                },
                {
                    "num": "Escena 3: La Bienvenida del Maestro Celestial",
                    "text": """Ren se acercó con cautela al asiento de Kenshin.

—Disculpe... ¿es usted el Maestro Celestial del Fuego? —preguntó Ren inclinándose levemente.

Kenshin se detuvo en seco, clavando sus ojos dorados y brillantes sobre Ren. En un instante, su mirada se volvió seria al notar la caja de madera en la espalda del joven y las marcas de venas en su rostro.

—Joven del acero negro... —dijo Kenshin con un tono de voz retumbante que hizo vibrar las ventanas del vagón—. Siento la presencia de un demonio a tu espalda, y sin embargo, tu espíritu no emite sed de sangre. Siéntate. La noche será larga y el enemigo ya está entre nosotros."""
                }
            ]
        },
        {
            "id": "chap5",
            "number": "Capítulo 5",
            "title": "La Trampa de los Sueños (Clímax del Volumen 2)",
            "scenes": [
                {
                    "num": "Escena 1: El Revisor del Tren",
                    "text": """Un hombre pálido con uniforme de revisor avanzó por el pasillo del vagón, picando los boletos de los pasajeros con una herramienta de metal.

Cuando picó el boleto de Ren y Kenshin, un sonido seco resonó en el aire. Sin darse cuenta, una niebla invisible de origen demoníaco envolvió las mentes de todos los presentes.

Ren sintió que sus párpados se volvían extremadamente pesados. Cayó sentado en su asiento, incapaz de mantener los ojos abiertos.

El boleto no era de papel común: era un arte demoníaco de ilusión que atrapaba a las víctimas dentro de sus recuerdos más profundos."""
                },
                {
                    "num": "Escena 2: El Sueño de la Herrería",
                    "text": """Dentro de su mente, Ren despertó en el taller de su infancia. Su padre estaba vivo frente al horno de la herrería y su hermana Miyuki cantaba felizmente sin los cuernos ni la marca de demonio.

El calor del hogar se sentía tan real y reconfortante que Ren sintió el deseo de quedarse en ese lugar para siempre.

Pero al mirarse las manos en el reflejo del agua, notó la katana de cristal negro en su cintura.

—Esto no es real... —susurró Ren con dolor en el pecho—. Mi familia ya no está. No puedo quedarme en una mentira mientras Miyuki me necesita afuera.

Para romper la ilusión del demonio dentro del sueño, Ren tomó su katana y ejecutó el acto más difícil: cortar su propio cuello dentro de la visión."""
                },
                {
                    "num": "Escena 3: La Fusión de la Locomotora (Cierre del Tomo 2)",
                    "img": "escena_climax.jpg",
                    "text": """Ren abrió los ojos de golpe en el mundo real, jadeando por aire.

A su lado, Kenshin luchaba contra su propia ilusión mientras su cuerpo emitía chispas de fuego dorado. El tren completo comenzó a retorcerse: las paredes de metal y los asientos se transformaron en carne, venas y boca humana.

El demonio no estaba escondido dentro del tren: el demonio se había fusionado con la locomotora entera.

Desde el techo del primer vagón, el Sexto Estirpe de Sangre —un demonio de piel pálida con bocas en las palmas de las manos— reía con locura mientras el tren avanzaba a toda velocidad hacia un barranco destruido.

—¡Despierta, Maestro Celestial del Fuego! —gritó Ren, desenvainando su katana negra envuelta en llamas rojas—. ¡Si el tren se cae, nadie sobrevivirá!

[ CONTINUARÁ EN EL VOLUMEN 3 ]"""
                }
            ]
        }
    ]
    
    bkm_id = 1
    for chap in content:
        p_c = doc.add_paragraph()
        add_bookmark(p_c, chap["id"], bkm_id)
        
        r_cnum = p_c.add_run(f"{chap['number']}\n")
        r_cnum.font.name = "Arial"
        r_cnum.font.size = Pt(14)
        r_cnum.font.bold = True
        r_cnum.font.color.rgb = RGBColor(185, 28, 28)
        
        r_ctitle = p_c.add_run(chap["title"])
        r_ctitle.font.name = "Arial"
        r_ctitle.font.size = Pt(20)
        r_ctitle.font.bold = True
        r_ctitle.font.color.rgb = RGBColor(220, 38, 38)
        
        end_bookmark(p_c, bkm_id)
        bkm_id += 1
        
        for sc in chap["scenes"]:
            p_s = doc.add_paragraph()
            r_snum = p_s.add_run(f"\n◆ {sc['num']}\n")
            r_snum.font.name = "Arial"
            r_snum.font.size = Pt(13)
            r_snum.font.bold = True
            r_snum.font.color.rgb = RGBColor(71, 85, 105)
            
            p_stext = doc.add_paragraph()
            r_stext = p_stext.add_run(sc["text"])
            r_stext.font.name = "Georgia"
            r_stext.font.size = Pt(11)
            
            if "img" in sc:
                img_p = os.path.join(target_root, sc["img"])
                if os.path.exists(img_p):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_p, width=Inches(5))
                    
        doc.add_page_break()
        
    doc_dest_root = os.path.join(target_root, "libro.docx")
    doc_dest_sys = os.path.join(target_sys, "libro.docx")
    doc.save(doc_dest_root)
    shutil.copy2(doc_dest_root, doc_dest_sys)
    print(f"Generated docx at {doc_dest_root}")

create_oni_vol2_docx()

# Generate ficha_producto.json
ficha_data = {
    "titulo": "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre) - Volumen 2: El Examen de la Montaña Sombría",
    "autor": "Nicolás Noguera",
    "precio": 20.00,
    "moneda": "USD",
    "genero": "Manga / Light Novel / Fantasía Oscura / Acción Sengoku",
    "headline": "La entrega de la Katana de Tinta y Fuego. El Tren demoníaco. Y la primera alianza con un Maestro Celestial del Gremio.",
    "descripcion": "Tras graduarse en la Selección Final, Ren recibe su arma definitiva: una katana de acero azabache con un filo de fuego carmesí. Junto a su hermana Miyuki, es enviado a investigar las misteriosas desapariciones en el Pueblo de los Faroles y a abordar el Tren de la Noche, donde el Sexto Estirpe de Sangre ha transformado la locomotora entera en un monstruo viviente. Una batalla a velocidad extrema por la supervivencia de docenas de pasajeros.",
    "beneficios": [
        "Manuscrito oficial ilustrado completo en formato .docx listo para eReaders y Amazon KDP.",
        "Ilustraciones de alta definición en estética anime dark fantasy.",
        "Continuación oficial inmediata del Volumen 1 de la saga de Nicolás Noguera."
    ],
    "capitulos": [
        "Capítulo 1: La Cuvo-Espada de la Noche",
        "Capítulo 2: El Pueblo de los Faroles",
        "Capítulo 3: La Danza de la Tinta Roja",
        "Capítulo 4: El Tren de las Sombras",
        "Capítulo 5: La Trampa de los Sueños"
    ]
}

with open(os.path.join(target_root, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)
with open(os.path.join(target_sys, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)

# Generate index.html landing page for Oni no Ketsuryu Vol 2
html_content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Oni no Ketsuryū Vol 2 - Nicolás Noguera | Tienda Oficial</title>
    <style>
        :root {
            --bg-color: #0b0f19;
            --card-bg: #111827;
            --accent: #dc2626;
            --accent-hover: #b91c1c;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
        }
        body {
            font-family: 'Segoe UI', system-ui, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 0;
            line-height: 1.6;
        }
        .header-banner {
            width: 100%;
            max-height: 400px;
            object-fit: cover;
            border-bottom: 3px solid var(--accent);
        }
        .container {
            max-width: 1000px;
            margin: 0 auto;
            padding: 40px 20px;
        }
        .product-grid {
            display: grid;
            grid-template-columns: 1fr 2fr;
            gap: 40px;
            margin-bottom: 50px;
        }
        @media (max-width: 768px) {
            .product-grid { grid-template-columns: 1fr; }
        }
        .cover-img {
            width: 100%;
            border-radius: 12px;
            box-shadow: 0 10px 30px rgba(220, 38, 38, 0.3);
            border: 1px solid rgba(220, 38, 38, 0.3);
        }
        .badge {
            background: var(--accent);
            color: white;
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            display: inline-block;
            margin-bottom: 15px;
        }
        h1 { font-size: 2.2rem; margin: 10px 0; color: #fff; }
        .author { color: var(--accent); font-weight: 600; font-size: 1.1rem; margin-bottom: 20px; }
        .price-tag {
            font-size: 2rem;
            font-weight: 800;
            color: #fff;
            margin: 20px 0;
        }
        .price-tag span { font-size: 1rem; color: var(--text-muted); }
        .buy-btn {
            background: linear-gradient(135deg, #dc2626, #991b1b);
            color: white;
            text-decoration: none;
            padding: 16px 36px;
            border-radius: 8px;
            font-weight: bold;
            font-size: 1.2rem;
            display: inline-block;
            box-shadow: 0 6px 20px rgba(220, 38, 38, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
        }
        .buy-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(220, 38, 38, 0.6);
        }
        .section-title {
            font-size: 1.5rem;
            border-left: 4px solid var(--accent);
            padding-left: 12px;
            margin: 40px 0 20px;
            color: #fff;
        }
        .chapters-list {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 25px;
            border: 1px solid rgba(255,255,255,0.05);
        }
        .chapters-list li {
            margin-bottom: 12px;
            color: var(--text-main);
        }
        .gallery-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin-top: 20px;
        }
        .gallery-grid img {
            width: 100%;
            border-radius: 8px;
            border: 1px solid rgba(255,255,255,0.1);
        }
    </style>
</head>
<body>
    <img src="banner.jpg" alt="Banner Oni no Ketsuryu Vol 2" class="header-banner">
    
    <div class="container">
        <div class="product-grid">
            <div>
                <img src="portada.jpg" alt="Portada Oni no Ketsuryu Vol 2" class="cover-img">
            </div>
            <div>
                <span class="badge">MANGA / LIGHT NOVEL</span>
                <h1>Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)</h1>
                <div class="author">Volumen 2: El Examen de la Montaña Sombría • Por Nicolás Noguera</div>
                <p>La entrega de la Katana de Tinta y Fuego. El Tren demoníaco. Y la primera alianza con un Maestro Celestial del Gremio.</p>
                <p>Tras graduarse en la Selección Final, Ren recibe su arma definitiva: una katana de acero azabache con un filo de fuego carmesí. Junto a su hermana Miyuki, es enviado a abordar el Tren de la Noche, donde el Sexto Estirpe de Sangre ha transformado la locomotora entera en un monstruo viviente.</p>
                
                <div class="price-tag">$20.00 <span>USD</span></div>
                <a href="#" class="buy-btn">COMPRAR AHORA ($20 USD)</a>
            </div>
        </div>

        <h2 class="section-title">Contenido del Volumen 2</h2>
        <div class="chapters-list">
            <ul>
                <li><strong>Capítulo 1:</strong> La Cuvo-Espada de la Noche (Escenas 1-3)</li>
                <li><strong>Capítulo 2:</strong> El Pueblo de los Faroles (Escenas 1-3)</li>
                <li><strong>Capítulo 3:</strong> La Danza de la Tinta Roja (Escenas 1-3)</li>
                <li><strong>Capítulo 4:</strong> El Tren de las Sombras (Escenas 1-3)</li>
                <li><strong>Capítulo 5:</strong> La Trampa de los Sueños (Clímax del Volumen 2 - Escenas 1-3)</li>
            </ul>
        </div>

        <h2 class="section-title">Ilustraciones Interiores Destacadas</h2>
        <div class="gallery-grid">
            <img src="escena_1.jpg" alt="La Katana de Tinta y Fuego">
            <img src="escena_climax.jpg" alt="El Tren Demoníaco en el Clímax">
        </div>
    </div>
</body>
</html>
"""

with open(os.path.join(target_root, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)
with open(os.path.join(target_sys, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)

print("Generated HTML landing pages for Vol 2 successfully")
