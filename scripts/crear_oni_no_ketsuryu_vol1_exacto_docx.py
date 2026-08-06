import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_oni_vol1_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-1"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 1: La Noche de las Hojas Rotas", 2, 1, "subtitulo_vol1")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: El Olor a Nieve y Sangre", "cap_1"),
        ("Capítulo 2: La Hermana de las Sombras", "cap_2"),
        ("Capítulo 3: La Prueba del Cazador", "cap_3"),
        ("Capítulo 4: El Arte de la Hoja Maldita", "cap_4"),
        ("Capítulo 5: La Selección Final", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: El Olor a Nieve y Sangre",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: La Herrería de la Montaña",
                    "texto": """La nieve caía en copos pesados sobre el monte Kurodake. Dentro de la herrería de la familia Hagane, el fuego del horno brillaba con un rojo intenso. Ren, un joven de quince años con cabello negro despeinado y ropas de trabajo manchadas de hollín, martillaba con ritmo constante un lingote de acero Tamahagane.

Su padre, el maestro herrero, observaba en silencio desde la esquina del taller mientras su hermana menor, Miyuki, atizaba el fuego cantando una antigua melodía de la aldea.

—El acero no se moldea solo con fuerza, Ren —dijo su padre, ajustando sus anteojos de madera—. Responde a la intención de tu corazón. Si tu mente está agitada, la hoja se romperá al primer impacto.""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Masacre del Viento Helado",
                    "texto": """El silencio de la noche fue destruido por un alarido gutural. La puerta de madera de la herrería saltó hecha pedazos cuando una monstruosa criatura de tres metros de altura y cuatro cuernos retorcidos irrumpió en el taller. Sus ojos amarillos brillaban con un hambre insaciable.

Antes de que Ren pudiera reaccionar, el demonio barrió a su padre con un zarpazo brutal. El impacto arrojó al maestro herrero contra la pared de piedra.

—¡Padre! —gritó Ren, abalanzándose sobre la primera hoja inacabada que encontró en la mesa de trabajo.

Miyuki corrió a proteger a su padre herido, pero el demonio la sujetó del cuello con una garra enorme, alzándola en el aire mientras la nieve entraba violentamente por la brecha del tejado.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Pacto de la Hoja Rota",
                    "texto": """Lleno de una furia ciega, Ren atacó al demonio. Su espada sin afilar impactó contra el cuello de la bestia, pero el metal no pudo atravesar la piel dura como la piedra y se rompió en mil pedazos.

El demonio rio con desprecio y arrojó a Miyuki contra el suelo sangrante. Un hilo de sangre demoníaca cayó directamente sobre la boca abierta de la niña.

En ese instante de desesperación absoluta, la hoja rota de cristal negro que descansaba en el altar de la herrería comenzó a brillar con una luz carmesí intensa. Líneas de fuego negro se extendieron por el rostro de Ren, marcando sus mejillas con símbolos rúnicos ancestrales.

—No tocarás a mi hermana... —susurró Ren, mientras la energía de la hoja rota envolvía su mano derecha.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: La Hermana de las Sombras",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: Las Venas de la Maldición",
                    "texto": """El demonio cayó decapitado en la nieve mientras las cenizas de su cuerpo eran consumidas por el fuego rojo de la espada de Ren.

Ren corrió hacia Miyuki, pero la niña comenzó a retorcerse de dolor. La sangre demoníaca que había ingresado en su cuerpo estaba transformando su anatomía: pequeñas protuberancias con forma de cuernos brotaron de su frente y sus pupilas se volvieron verticales como las de una bestia.

Sin embargo, a diferencia de los otros monstruos, Miyuki no atacó a Ren. Lagrimas cayeron de sus ojos rasgados mientras luchaba por mantener el control de su mente humana.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Silencio del Bambú",
                    "texto": """Para evitar que su hermana fuera dominada por la sed de sangre o lastimara a inocentes, Ren colocó un trozo de bambú sellado con cuerda sobre la boca de Miyuki.

La niña asintió con la cabeza en un gesto de absoluta confianza hacia su hermano.

Ren cubrió a su hermana con una manta gruesa para protegerla de la luz del sol matutino y preparó un cesto de mimbre reforzado para llevarla en su espalda.

El calor de la herrería se había extinguido. La travesía por la supervivencia comenzaba en mitad del invierno helado.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Filo del Cuervo",
                    "texto": """Antes de salir del paso de la montaña, una sombra descendió de la copa de los pinos nevados.

Un cazador alto, vistiendo un haori azul con patrones de olas y una máscara de Tengu cuervo que ocultaba su rostro, bloqueó el camino con una katana de filo rojo carmesí.

—Una humana convertida en demonio debe ser ejecutada inmediatamente —dijo el espadachín de la máscara de Tengu con voz fría y distante.

Ren desenvainó su espada negra rota y se colocó frente a su hermana en posición de combate.

—¡Primero tendrás que pasar por encima de mi cadáver! —gritó Ren.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: La Prueba del Cazador",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: La Apuesta sobre la Nieve",
                    "texto": """El cazador del haori azul atacó con una velocidad que Ren no pudo ni registrar. En un pestañeo, el filo rojo estuvo a milímetros del cuello de Ren.

Pero antes de que la espada golpeara, Miyuki saltó de la cesta y se interpuso entre la espada y su hermano, extendiendo los brazos para protegerlo.

El espadachín de la máscara de Tengu detuvo su ataque a milímetros de la niña, sorprendido al ver a un demonio protegiendo a un ser humano en lugar de devorarlo.

—Interesante... —murmuró el cazador enfundando su katana—. Si quieres protegerla y devolverle su humanidad, busca al Maestro del Monte Sagrado en la cima del Fujikane.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Ascenso a la Montaña Sombría",
                    "texto": """Ren caminó durante tres días y tres noches con la cesta de madera a sus espaldas, soportando tempestades de nieve y terrenos rocosos.

El aire en la cima de la montaña era tan delgado que a Ren le costaba respirar a cada paso. Sus manos estaban cubiertas de llagas por el frío y la falta de alimentos, pero no se detuvo ni un segundo.

Miyuki se mantenía tranquila dentro de la cesta, infundiéndole fuerza a su hermano con su presencia silenciosa.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Cabaña del Maestro",
                    "texto": """Al llegar a la cumbre del Fujikane, una pequeña casa de madera rodeada de flores de glicina púrpura emergió entre la niebla.

En el porche sentado en posición de loto, un anciano guerrero retirado con cicaprices en el rostro los esperaba.

Ren cayó de rodillas en la nieve y se inclinó profundamente.

—¡Por favor, enséñeme a luchar para proteger a mi hermana y cortar a los demonios! —suplicó Ren.

El anciano lo miró en silencio antes de responder:

—La montaña será tu primer maestro. Si sobrevives a sus trampas, moldearé tu acero.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: El Arte de la Hoja Maldita",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Entrenamiento del Acero",
                    "texto": """Durante dos largos años, Ren fue sometido a un entrenamiento inhumano en el bosque del Monte Fujikane.

Tenía que esquivar trampa de troncos y cuchillas escondidas con los ojos vendados, aprendiendo a percibir el peligro a través del olor del aire y el flujo de la energía.

Sus brazos y torso se cubrieron de marcas de entrenamiento y su resistencia física alcanzó niveles extraordinarios.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Prueba del Peñasco",
                    "texto": """Al finalizar el segundo año, el maestro lo llevó frente a una roca gigantesca de tres metros de diámetro en mitad del claro del bosque.

—Para ingresar a la Selección Final del Gremio Cuervo, debes partir esta roca por la mitad usando una katana común —dijo el maestro.

Ren meditó frente al peñasco durante días. Al recordar el martilleo de la herrería de su padre y coordinar su respiración con los latidos de su corazón, desató un tajo perfecto envuelto en llamas rojas que dividió la mole de piedra por la mitad.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Partida hacia el Fujikane",
                    "texto": """El maestro contempló el peñasco dividido con orgullo. Le entregó a Ren un nuevo haori negro con patrones carmesí y una funda especial para su katana.

—Estás listo, Ren —dijo el maestro—. Ve a la Selección Final en el monte rodeado de glicinas. Trae gloria a la memoria de tu familia.

Ren cargó la cesto de Miyuki en su espalda y emprendió el camino hacia la prueba definitiva.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: La Selección Final",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: El Monte de las Glicinas",
                    "texto": """El Monte Fujikane estaba rodeado por un cinturón inmenso de árboles de glicina púrpura que florecían todo el año. Su aroma mantenía a los demonios prisioneros dentro del perímetro de la montaña.

Más de sesenta jóvenes aspirantes de todo el país estaban reunidos frente a la gran puerta de madera.

Dos niñas con rostros de porcelana explicaron la regla: sobrevivir siete noches dentro de la montaña infestada de demonios.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Demonio de las Manos",
                    "texto": """En la tercera noche de la prueba, la tierra tembló.

Un demonio grotesco y gigantesco, cubierto por decenas de brazos deformes que brotaban de su cuello y torso, emergió de las sombras del bosque. La criatura había devorado a docenas de aspirantes en las selecciones anteriores.

—¡Otro aprendiz con una máscara del maestro! —siseó la criatura con furia, extendiendo sus brazos gigantes para aplastar a Ren.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Postura Suprema",
                    "texto": """Ren esquivó las decenas de brazos con agilidad felina. Enfiló su katana negra y concentró toda la Respiración de Sangre en un solo punto.

—Respiración de Sangre: Primera Postura: Tajo del Horno Encendido.

Ren giró en un tornado de fuego rojo brillante que amputó los brazos del monstruo y le cercenó el cuello de un solo golpe magistral.

El demonio gigante se disolvió en cenizas bajo los primeros rayos del sol matutino entre las glicinas púrpuras.

Ren había sobrevivido a la Selección Final.""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Vol 1 at {output_docx}")

if __name__ == "__main__":
    create_oni_vol1_manuscript()
