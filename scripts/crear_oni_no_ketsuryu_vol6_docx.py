# -*- coding: utf-8 -*-
import sys
import os
import io
import json
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

vol6_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-6")
vol6_dir_downloads = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-6")

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(180, 20, 20)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(120, 15, 15)
    return p

def add_image_safe(doc, img_path, width_inches=6.0):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

def build_vol6_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    r_title = p_title.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n\nVOLUMEN 6: LAS CATACUMBAS DEL OLVIDO")
    r_title.bold = True
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(160, 20, 20)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Obra Original por Nicolás Noguera\nEdición Ilustrada de Alta Definición (Dark Fantasy Anime)")
    r_sub.font.name = "Georgia"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Portada image
    add_image_safe(doc, vol6_dir / "portada.jpg", 5.0)

    doc.add_page_break()

    # Synopsis / Resumen
    add_heading_styled(doc, "SINOPSIS DEL VOLUMEN 6", 1)
    p_syn = doc.add_paragraph()
    p_syn.paragraph_format.line_spacing = 1.15
    p_syn.paragraph_format.space_after = Pt(10)
    p_syn.add_run("Tras el milagro inesperado del amanecer, la hermana de Ren demuestra inmunidad a la luz solar, desatando una cacería desesperada por parte del Rey Demonio Kageyama. Para prepararse ante el choque inevitable, el Gremio Cuervo traslada a los jóvenes guerreros al Santuario Principal de Flores de Glicina, donde iniciarán el brutal Entrenamiento de los Pilares.\n\nSin embargo, la clave para derrotar a las Lunas Demoníacas yace en las profundidades de la tierra: en las antiguas Catacumbas del Olvido, donde aguarda la Piedra de Afilado Solar y la presencia imponente del Primer Lunar Rojo: Kurogane, el Demonio de los Seis Ojos.")

    add_image_safe(doc, vol6_dir / "banner.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 1
    add_heading_styled(doc, "CAPÍTULO 1: El Refugio de las Flores de Glicina", 1)
    add_image_safe(doc, vol6_dir / "escena_1.jpg", 6.0)
    p_c1 = doc.add_paragraph()
    p_c1.add_run("El aroma dulce y denso de las flores de glicina inundaba el valle nocturno. Protegido por barreras naturales y sellos ancestrales, el Santuario de las Glicinas era el último bastión de paz en un mundo devorado por la oscuridad...")
    add_image_safe(doc, vol6_dir / "escena_c1_e1.jpg", 6.0)
    p_c1_2 = doc.add_paragraph()
    p_c1_2.add_run("En la sala central de tatami, los Pilares Hashira debatían con severidad las decisiones tácticas mientras Ren se sometía a pruebas de fuerza extrema contra rocas gigantescas y cascadas congeladas...")
    add_image_safe(doc, vol6_dir / "escena_c1_e2.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_c1_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 2
    add_heading_styled(doc, "CAPÍTULO 2: El Laberinto Bajo la Tierra", 1)
    add_image_safe(doc, vol6_dir / "escena_c2_e1.jpg", 6.0)
    p_c2 = doc.add_paragraph()
    p_c2.add_run("Descendiendo por una escalera de caracol de piedra antigua, las sombras se volvieron pesadas y frías. Las Catacumbas del Olvido guardaban relicarios de una era perdida donde las runas solares brillaban con luz propia...")
    add_image_safe(doc, vol6_dir / "escena_c2_e2.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_c2_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 3
    add_heading_styled(doc, "CAPÍTULO 3: La Luna contra el Sol", 1)
    add_image_safe(doc, vol6_dir / "escena_c3_e1.jpg", 6.0)
    p_c3 = doc.add_paragraph()
    p_c3.add_run("Seis ojos rojos fijos en la penumbra. Kurogane desenfundó su espada carmesí de carne y escamas, haciendo crujir el aire con cortes en forma de media luna. Ren activó la Percepción del Mundo Transparente para distinguir los movimientos de su oponente...")
    add_image_safe(doc, vol6_dir / "escena_c3_e2.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_c3_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 4
    add_heading_styled(doc, "CAPÍTULO 4: El Afilado de la Hoja Definitiva", 1)
    add_image_safe(doc, vol6_dir / "escena_c4_e1.jpg", 6.0)
    p_c4 = doc.add_paragraph()
    p_c4.add_run("Al presionar su katana contra el altar de cuarzo solar, una chispa dorada envolvió el acero, purificando la hoja y otorgándole un brillo rubí permanente. Afuera, el Líder Supremo aguardaba en calma meditativa el destino final de la sede...")
    add_image_safe(doc, vol6_dir / "escena_c4_e2.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_c4_e3.jpg", 6.0)

    doc.add_page_break()

    # Capítulo 5
    add_heading_styled(doc, "CAPÍTULO 5: La Gran Explosión y el Castillo Infinito", 1)
    add_image_safe(doc, vol6_dir / "escena_c5_e1.jpg", 6.0)
    p_c5 = doc.add_paragraph()
    p_c5.add_run("El Rey Demonio irrumpió en la residencia. Pero el Maestro tenía preparada una trampa final: una explosión apocalíptica retumbó por las montañas, destruyendo el santuario y abriendo las puertas del Castillo Infinito...")
    add_image_safe(doc, vol6_dir / "escena_c5_e2.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_c5_e3.jpg", 6.0)
    add_image_safe(doc, vol6_dir / "escena_climax.jpg", 6.0)

    # Save to both locations
    out1 = vol6_dir / "libro.docx"
    out2 = vol6_dir_downloads / "libro.docx"
    doc.save(str(out1))
    doc.save(str(out2))
    print(f"Generated libro.docx successfully for Volume 6 at {out1}")

if __name__ == "__main__":
    build_vol6_docx()
