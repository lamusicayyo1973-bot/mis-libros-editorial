import os
import glob
import shutil
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

src_dir = r"C:\Users\nicol\.gemini\antigravity\brain\6adf8ce5-9839-4292-a8a1-57beed4c3827"
target_root = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-3"
target_sys = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\oni-no-ketsuryu-volumen-3"

os.makedirs(target_root, exist_ok=True)
os.makedirs(target_sys, exist_ok=True)

matches = glob.glob(os.path.join(src_dir, 'oni_vol3_portada*.jpg'))
if matches:
    matches.sort(key=os.path.getmtime, reverse=True)
    latest_portada = matches[0]
    shutil.copy2(latest_portada, os.path.join(target_root, "portada.jpg"))
    shutil.copy2(latest_portada, os.path.join(target_sys, "portada.jpg"))

# Duplicate portada as placeholders for thumbnail, banner, and scene illustrations
portada_p = os.path.join(target_root, "portada.jpg")
if os.path.exists(portada_p):
    for fn in ["thumbnail.jpg", "banner.jpg", "escena_1.jpg", "escena_climax.jpg"]:
        shutil.copy2(portada_p, os.path.join(target_root, fn))
        shutil.copy2(portada_p, os.path.join(target_sys, fn))

def create_oni_vol3_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(220, 38, 38)
    
    run_sub = title_p.add_run("Volumen 3: El Tren de las Sombras\n\nMANGA / LIGHT NOVEL • FANTASÍA OSCURA & ACCIÓN SENGOKU\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(185, 28, 28)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    img_cover_path = os.path.join(target_root, "portada.jpg")
    if os.path.exists(img_cover_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_cover_path, width=Inches(4.5))
        
    doc.add_page_break()
    
    # TOC Header
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = toc_p.add_run("TABLA DE CONTENIDOS\n")
    r_toc.font.name = "Arial"
    r_toc.font.size = Pt(20)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(220, 38, 38)
    
    chapters_toc = [
        ("Capítulo 1", "La Carne del Tren", "chap1"),
        ("Capítulo 2", "El Cuello de Hierro", "chap2"),
        ("Capítulo 3", "La Llegada del Tercer Estirpe", "chap3"),
        ("Capítulo 4", "La Última Llama", "chap4"),
        ("Capítulo 5", "La Herencia del Fuego (Clímax del Volumen 3)", "chap5")
    ]
    
    for c_num, c_title, c_bkm in chapters_toc:
        p_t = doc.add_paragraph()
        r_num = p_t.add_run(f"• {c_num}: ")
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(185, 28, 28)
        r_name = p_t.add_run(c_title)
        r_name.font.italic = True
        
    doc.add_page_break()
    
    # Full Content Text
    content = [
        {
            "id": "chap1",
            "number": "Capítulo 1",
            "title": "La Carne del Tren",
            "scenes": [
                {
                    "num": "Escena 1: El Despertar del Maestro Celestial",
                    "text": """El calor dentro del vagón se volvió sofocante. Las venas de carne demoníaca que cubrían los asientos se retorcían con violencia, intentando envolver a los pasajeros dormidos.

De pronto, un estallido de fuego dorado rasgó la niebla.

Kenshin, el Maestro Celestial del Fuego, rompió la ilusión del sueño. Su haori con patrones de llamas se agitó mientras desenvainaba su katana de tono anaranjado brillante. La temperatura del aire subió diez grados en un segundo.

—¡Joven Hagane! —retumbó la voz de Kenshin con un entusiasmo contagioso—. ¡Excelente trabajo al cortar tu propio cuello en el sueño para despertar! Tu resolución es digna de un verdadero cazador.

—¡El demonio es la locomotora entera! —explicó Ren, esquivando las tentáculos de carne que brotaban del techo—. ¡Si cortamos los vagones, no servirá de nada! ¡Su cuello debe estar cerca de la caldera de carbón!

—¡Entendido! —sonrió Kenshin, apoyando la mano en la empuñadura—. Tú y la joven demonio protejan los ocho vagones de pasajeros. ¡Yo me abriré paso hasta el núcleo del tren!"""
                },
                {
                    "num": "Escena 2: La Defensa de los Pasajeros",
                    "text": """Kenshin se desplazó a una velocidad sobrehumana, dejando una estela de fuego dorado a lo largo del pasillo que incineró la carne demoníaca de los primeros cuatro vagones.

Mientras tanto, en la parte trasera del tren, cientos de bocas humanas brotaron de las paredes metálicas para devorar a los civiles dormidos.

La caja de madera de Ren se abrió de golpe. Miyuki saltó al pasillo con los ojos encendidos en una luz violeta. Con una agilidad pasmosa, la joven usó sus garras para cortar la masa de carne, mientras su sangre ardía en llamas púrpuras que sellaban las heridas del metal para que no volvieran a crecer.

Ren luchaba espalda con espalda junto a su hermana. Cada tajo de su katana negra, impulsado por la Estilo de Esgrima de Sangre, trazaba arcos de fuego rojo que mantenían a salvo a los pasajeros.

—¡No dejaré que nadie más pierda a su familia esta noche! —gritó Ren, aumentando el ritmo de su corazón."""
                },
                {
                    "num": "Escena 3: El Frente de la Locomotora",
                    "img": "escena_1.jpg",
                    "text": """Ren corrió por el techo del tren bajo la luz de la luna llena, esquivando las ráfagas de viento y los ojos de cera que brotaban del metal.

Al llegar al primer vagón, se encontró con el joven asustadizo de la Selección Final —el muchacho de la doble katana corta que llevaba un haori amarillo con bordados de rayos—. Estaba paralizado de terror sobre la caldera, temblando mientras el vapor hirviente salía por las tuberías.

—¡No puedo hacerlo! ¡Voy a morir! ¡Es un Estirpe de Sangre! —gritaba el joven del haori amarillo entre lágrimas.

De pronto, un rayo de luz blanca atravesó el aire. El muchacho se quedó dormido de pie por el pánico absoluto. En ese estado de trance, su postura cambió por completo: adoptó una posición de esgrima perfecta y una corriente de aire electrificado comenzó a sonar alrededor de sus espaldas.

Era la Estilo de Esgrima del Trueno Inverso."""
                }
            ]
        },
        {
            "id": "chap2",
            "number": "Capítulo 2",
            "title": "El Cuello de Hierro",
            "scenes": [
                {
                    "num": "Escena 1: Velocidad del Rayo",
                    "text": """El chico del haori amarillo se desplazó en un milisegundo.

—Estilo de Esgrima del Trueno... Primera Postura: Chispazo del Destello.

Un destello de luz amarilla cortó la noche. Las decenas de brazos de carne que intentaban sujetar a Ren fueron rebanadas instantáneamente por las dos katanas cortas del joven dormido.

Ren aprovechó la brecha abierta por su compañero. Saltó hacia la escotilla de la caldera de carbón, donde la masa de carne demoníaca se concentraba en un núcleo espeso que latía con el pulso del Sexto Estirpe de Sangre.

—¡Ahí está el cuello! —gritó Ren.

Desde la masa de carne, la cara del demonio pálido emergió, mostrando las bocas en sus palmas de las manos.

—¡Es inútil, cazadores de poca monta! —siseó el Sexto Estirpe—. ¡Incluso si cortan este punto, mi masa corporal regenerará la caldera en un segundo!"""
                },
                {
                    "num": "Escena 2: La Fusión de los Estilos",
                    "text": """Ren inhaló aire hasta el fondo de sus pulmones. Las venas negras de su rostro se extendieron hasta la sien derecha, brillando con una luz carmesí brillante.

—¡No si lo hacemos al mismo tiempo! —respondió Ren.

El joven del haori amarillo, aun dormido, se posicionó al lado izquierdo de la caldera, cargando sus katanas con electricidad blanca. Ren se colocó a la derecha, envolviendo su katana negra en llamas rojas de fuego de herrería.

—¡Estilo de Esgrima del Trueno: Destello Doble!
—¡Estilo de Esgrima de Sangre: Tajo del Horno Encendido!

Los dos ataques impactaron el núcleo de carne de la caldera exactamente al mismo tiempo. La combinación del calor del fuego y la velocidad de la electricidad destrozó la estructura de masa demoníaca sin darle tiempo a regenerarse."""
                },
                {
                    "num": "Escena 3: El Descarrilamiento",
                    "text": """El grito agudo del Sexto Estirpe de Sangre resonó en toda la montaña mientras la masa orgánica del tren comenzaba a disolverse en cenizas.

Sin la carne demoníaca para sostener la estructura, los frenos de la locomotora colapsaron. El tren de vapor de miles de toneladas comenzó a salirse de las vías a más de ochenta kilómetros por hora, volcando sobre el campo abierto.

Ren usó su cuerpo como escudo para proteger al muchacho del haori amarillo y a la caja de Miyuki mientras el metal se retorcía sobre la hierba.

Kenshin, desde el interior, usó su Estilo del Fuego Carmesí para crear una burbuja de aire caliente que amortiguó el impacto de todos los vagones de pasajeros, salvando la vida de los doscientos civiles a bordo.

El vapor se dispersó sobre el campo bajo la luz del amanecer. El peligro parecía haber pasado."""
                }
            ]
        },
        {
            "id": "chap3",
            "number": "Capítulo 3",
            "title": "La Llegada del Tercer Estirpe",
            "scenes": [
                {
                    "num": "Escena 1: La Presencia Absoluta",
                    "text": """La victoria se esfumó en un segundo.

El aire sobre el campo de hierba se volvió tan denso y pesado que Ren cayó sobre sus rodillas, incapaz de respirar. El suelo bajo sus pies comenzó a agrietarse sin razón aparente.

Desde las sombras de los árboles de la linde del bosque, un hombre de torso desnudo caminó despacio hacia el lugar del accidente. Su piel era pálida como la nieve, estaba cubierta por tatuajes rúnicos de color azul marino y sus ojos brillaban con un tono dorado donde se leía claramente la marca del Tercer Estirpe de Sangre.

Su sola presencia emitía un aura de presión física que aplastaba la hierba a su alrededor.

—Una técnica de fuego bastante pulida... —dijo el Tercer Estirpe, mirando a Kenshin con una sonrisa llena de anticipación—. Hacía más de cincuenta años que no encontraba a un Maestro Celestial de la Estilo del Fuego Carmesí."""
                },
                {
                    "num": "Escena 2: La Propuesta del Demonio",
                    "text": """Kenshin dio un paso al frente, interponiéndose entre el Tercer Estirpe y los jóvenes heridos. Aunque estaba agotado por mantener a salvo a los pasajeros, su postura con la katana de fuego era impecable.

—Soy Kenshin, Maestro Celestial del Fuego del Gremio Cuervo —dijo con voz serena.

—Soy Rikudo —respondió el demonio, ajustando sus puños en una posición de artes marciales ancestrales—. Tienes un espíritu de combate supremo, Kenshin. Tu cuerpo humano alcanzará su límite en unos años; te volverás viejo, débil y morirás. Conviértete en un demonio. Con el poder de la sangre de nuestro Señor, podrás entrenar tu técnica por toda la eternidad.

Kenshin no dudó ni un segundo. Sonrió con amabilidad y negó con la cabeza.

—Envejecer y morir es lo que le da valor y belleza a la vida humana. Precisamente porque somos efímeros, nuestro esfuerzo es sagrado. No me convertiré en un monstruo."""
                },
                {
                    "num": "Escena 3: El Choque de los Titanes",
                    "text": """Sin decir una palabra más, Rikudo se lanzó al ataque.

El choque entre los puños de energía del demonio y la katana de fuego de Kenshin levantó una onda de choque que destruyó los restos del primer vagón de tren. La velocidad del combate era tan absurda que Ren solo podía ver destellos de luz dorada y azul cruzando el campo.

—¡Arte Demoníaco: Aguja de la Muerte! —gritó Rikudo, desatando una ráfaga de cientos de golpes invisibles.

—¡Estilo del Fuego Carmesí: Quinta Postura: Tigre de las Llamas! —respondió Kenshin, creando un felino gigante de fuego dorado que devoró la ráfaga del enemigo.

El suelo tembló mientras la batalla alcanzaba un nivel de destrucción destructivo."""
                }
            ]
        },
        {
            "id": "chap4",
            "number": "Capítulo 4",
            "title": "La Última Llama",
            "scenes": [
                {
                    "num": "Escena 1: El Límite del Cuerpos Humano",
                    "text": """A pesar de la maestría de Kenshin, la regeneración del Tercer Estirpe era instantánea. Cada corte hecho por la katana de fuego se cerraba en milisegundos, mientras que las heridas en el cuerpo de Kenshin comenzaban a acumularse.

El Maestro Celestial del Fuego sangraba por la frente y tenía dos costillas rotas, pero su mirada dorada no mostraba una sola gota de miedo.

Ren intentó levantarse para ayudarlo, agarrando su katana negra, pero Rikudo ni siquiera se giró; usó la presión de su aura para lanzar a Ren contra el suelo.

—¡No te metas, chico! —gritó Kenshin sin perder la sonrisa—. ¡Mantén la posición y protege a tu hermana! ¡Este es el deber de un Maestro Celestial!

Kenshin inhaló una cantidad monumental de aire, haciendo crujir sus pulmones mientras el fuego de su espada se volvía de un color blanco incandescente."""
                },
                {
                    "num": "Escena 2: La Novena Postura",
                    "text": """—Estilo del Fuego Carmesí... Novena Postura: Purgatorio.

Kenshin se convirtió en un meteoro de fuego blanco que arrasó la hierba del campo. Se lanzó directo al pecho de Rikudo en una estocada definitiva.

El impacto fue ensordecedor. Rikudo clavó su puño derecho en el torso de Kenshin, pero al mismo tiempo, la katana de fuego de Kenshin atravesó el cuello del Tercer Estirpe hasta la mitad.

Rikudo intentó retirar el puño para decapitar al Maestro Celestial, pero Kenshin usó sus propios músculos abdominales y su fuerza física sobrehumana para atrapar el brazo del demonio dentro de su cuerpo.

—¡No te moverás de aquí hasta que el sol aparezca! —gritó Kenshin, apretando el cuello de Rikudo con la mano izquierda desnuda.

En el horizonte, los primeros rayos del sol matutino comenzaron a asomar sobre la cumbre de las montañas."""
                },
                {
                    "num": "Escena 3: La Huida de la Sombra",
                    "text": """Al sentir el calor del primer rayo de sol sobre su piel, Rikudo entró en pánico absoluto. La piel de sus hombros comenzó a quemarse y desintegrarse en cenizas.

Con un alarido de desesperación, el Tercer Estirpe usó una fuerza bruta desmedida para arrancarse sus propios brazos, liberándose del agarre de Kenshin. Se arrojó de cabeza hacia las sombras del bosque espeso antes de que la luz solar lo alcanzara por completo.

Ren, lleno de una ira ciega, tomó su katana negra y se la arrojó con todas sus fuerzas por la espalda.

La hoja de Ren atravesó el pecho de Rikudo justo cuando el demonio desaparecía en la penumbra del bosque.

—¡Cobarde! —gritó Ren con lágrimas en los ojos, corriendo hacia el bosque—. ¡No huyas! ¡El señor Kenshin no huyó! ¡Él ganó esta batalla porque protegía a los demás! ¡Tú solo eres una bestia miedosa que le teme al sol!"""
                }
            ]
        },
        {
            "id": "chap5",
            "number": "Capítulo 5",
            "title": "La Herencia del Fuego (Clímax del Volumen 3)",
            "scenes": [
                {
                    "num": "Escena 1: Las Últimas Palabras del Maestro Celestial",
                    "text": """El sol matutino bañó el campo de hierba.

Kenshin permanecía de pie, apoyado en la empuñadura de su katana de fuego clavada en la tierra. Su figura erguida proyectaba una sombra larga sobre el suelo.

Ren y Miyuki corrieron a su lado, cayendo de rodillas. Ren intentaba presionar las heridas del Maestro Celestial con sus manos sangrantes, pero Kenshin le colocó suavemente la mano sobre la cabeza para detenerlo.

—Ya es suficiente, joven Hagane... —dijo Kenshin con voz débil pero cálida—. Mi tiempo ha llegado.

—Señor Kenshin... lo siento... no pude ayudarlo... —sollozó Ren, apoyando la frente en la hierba.

—No te disculpes —sonrió Kenshin, mirando al cielo—. Salvaste a los pasajeros de los vagones. Tu hermana luchó con corazón humano. Confío en ustedes dos. Vayan a la finca de mi familia... lean los escritos antiguos de los anteriores Maestros Celestiales del Fuego... ahí encontrarán la pista sobre la Estilo de Dominio Solar."""
                },
                {
                    "num": "Escena 2: La Visión de la Madre",
                    "text": """Kenshin alzó la mirada hacia la luz del sol. En mitad de la bruma dorada de la mañana, vio la figura de su difunta madre sonriéndole con ternura.

—¿Lo hice bien, madre? —susurró Kenshin—. ¿Cumplí con el deber con el que nací?

La visión de su madre asintió con la cabeza, extendiéndole los brazos.

—Lo hiciste maravillosamente, mi querido hijo.

Kenshin cerró los ojos con una sonrisa radiante en el rostro. Su corazón dejó de latir, pero su cuerpo permaneció erguido como una estatua inamovible en medio del campo de flores.

Ren y Miyuki inclinaron la cabeza en un silencio profundo, honrando al hombre que había dado su vida para proteger la de ellos."""
                },
                {
                    "num": "Escena 3: La Decisión del Guerrero (Cierre del Tomo 3)",
                    "img": "escena_climax.jpg",
                    "text": """Tres días después.

En el cementerio del Gremio Cuervo, las cuervos mensajeros llevaron la noticia de la caída del Maestro Celestial del Fuego a todos los rincones del país. Los restantes ocho Maestros Celestiales recibieron el mensaje en silencio, apretando sus armas con una nueva determinación.

Ren caminaba por el sendero hacia la Finca del Fuego, llevando la guarda de la katana con forma de llama de Kenshin enganchada en su propia espada negra.

Las venas de su rostro habían cambiado: ya no eran líneas desordenadas, sino que habían tomado la forma de una llama carmesí sobre su mejilla derecha.

—No dejaré que tu sacrificio sea en vano, Kenshin —dijo Ren mirando al cielo azul—. Voy a dominar la Estilo de Dominio Solar y voy a cortar la cabeza de Kageyama con mis propias manos.

[ CONTINUARÁ EN EL VOLUMEN 4 ]"""
                }
            ]
        }
    ]
    
    bkm_id = 1
    for chap in content:
        p_c = doc.add_paragraph()
        add_bookmark(p_c, chap["id"], bkm_id)
        
        r_cnum = p_c.add_run(f"{chap['number']}\n")
        r_cnum.font.name = "Arial"
        r_cnum.font.size = Pt(14)
        r_cnum.font.bold = True
        r_cnum.font.color.rgb = RGBColor(185, 28, 28)
        
        r_ctitle = p_c.add_run(chap["title"])
        r_ctitle.font.name = "Arial"
        r_ctitle.font.size = Pt(20)
        r_ctitle.font.bold = True
        r_ctitle.font.color.rgb = RGBColor(220, 38, 38)
        
        end_bookmark(p_c, bkm_id)
        bkm_id += 1
        
        for sc in chap["scenes"]:
            p_s = doc.add_paragraph()
            r_snum = p_s.add_run(f"\n◆ {sc['num']}\n")
            r_snum.font.name = "Arial"
            r_snum.font.size = Pt(13)
            r_snum.font.bold = True
            r_snum.font.color.rgb = RGBColor(71, 85, 105)
            
            p_stext = doc.add_paragraph()
            r_stext = p_stext.add_run(sc["text"])
            r_stext.font.name = "Georgia"
            r_stext.font.size = Pt(11)
            
            if "img" in sc:
                img_p = os.path.join(target_root, sc["img"])
                if os.path.exists(img_p):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_p, width=Inches(5))
                    
        doc.add_page_break()
        
    doc_dest_root = os.path.join(target_root, "libro.docx")
    doc_dest_sys = os.path.join(target_sys, "libro.docx")
    doc.save(doc_dest_root)
    shutil.copy2(doc_dest_root, doc_dest_sys)
    print(f"Generated docx at {doc_dest_root}")

create_oni_vol3_docx()

# Generate ficha_producto.json
ficha_data = {
    "titulo": "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre) - Volumen 3: El Tren de las Sombras",
    "autor": "Nicolás Noguera",
    "precio": 20.00,
    "moneda": "USD",
    "genero": "Manga / Light Novel / Fantasía Oscura / Acción Sengoku",
    "headline": "La batalla en el Tren Demoníaco. El choque del Maestro Celestial del Fuego contra el Tercer Estirpe. Y el legado sagrado de la Estilo de Dominio Solar.",
    "descripcion": "El Tren de las Sombras se ha convertido en una masa viva de carne demoníaca. Ren, Miyuki y el joven del haori de rayos combinan sus estilos para cortar el núcleo de la caldera, pero tras el descarrilamiento, el aterrador Rikudo —Tercer Estirpe de Sangre— emerge del bosque. Kenshin, el Maestro Celestial del Fuego, desata su Novena Postura Purgatorio en una batalla legendaria hasta las últimas consecuencias bajo los primeros rayos del sol.",
    "beneficios": [
        "Manuscrito oficial ilustrado completo en formato .docx listo para eReaders y Amazon KDP.",
        "Ilustraciones de alta calidad en estética anime/manga dark fantasy.",
        "Gran clímax del arco del Tren Demoníaco en la saga de Nicolás Noguera."
    ],
    "capitulos": [
        "Capítulo 1: La Carne del Tren",
        "Capítulo 2: El Cuello de Hierro",
        "Capítulo 3: La Llegada del Tercer Estirpe",
        "Capítulo 4: La Última Llama",
        "Capítulo 5: La Herencia del Fuego"
    ]
}

with open(os.path.join(target_root, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)
with open(os.path.join(target_sys, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)

# Generate index.html landing page for Oni no Ketsuryu Vol 3
html_content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Oni no Ketsuryū Vol 3 - Nicolás Noguera | Tienda Oficial</title>
    <style>
        :root {
            --bg-color: #0b0f19;
            --card-bg: #111827;
            --accent: #dc2626;
            --accent-hover: #b91c1c;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
        }
        body {
            font-family: 'Segoe UI', system-ui, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 0;
            line-height: 1.6;
        }
        .header-banner {
            width: 100%;
            max-height: 400px;
            object-fit: cover;
            border-bottom: 3px solid var(--accent);
        }
        .container {
            max-width: 1000px;
            margin: 0 auto;
            padding: 40px 20px;
        }
        .product-grid {
            display: grid;
            grid-template-columns: 1fr 2fr;
            gap: 40px;
            margin-bottom: 50px;
        }
        @media (max-width: 768px) {
            .product-grid { grid-template-columns: 1fr; }
        }
        .cover-img {
            width: 100%;
            border-radius: 12px;
            box-shadow: 0 10px 30px rgba(220, 38, 38, 0.3);
            border: 1px solid rgba(220, 38, 38, 0.3);
        }
        .badge {
            background: var(--accent);
            color: white;
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            display: inline-block;
            margin-bottom: 15px;
        }
        h1 { font-size: 2.2rem; margin: 10px 0; color: #fff; }
        .author { color: var(--accent); font-weight: 600; font-size: 1.1rem; margin-bottom: 20px; }
        .price-tag {
            font-size: 2rem;
            font-weight: 800;
            color: #fff;
            margin: 20px 0;
        }
        .price-tag span { font-size: 1rem; color: var(--text-muted); }
        .buy-btn {
            background: linear-gradient(135deg, #dc2626, #991b1b);
            color: white;
            text-decoration: none;
            padding: 16px 36px;
            border-radius: 8px;
            font-weight: bold;
            font-size: 1.2rem;
            display: inline-block;
            box-shadow: 0 6px 20px rgba(220, 38, 38, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
        }
        .buy-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(220, 38, 38, 0.6);
        }
        .section-title {
            font-size: 1.5rem;
            border-left: 4px solid var(--accent);
            padding-left: 12px;
            margin: 40px 0 20px;
            color: #fff;
        }
        .chapters-list {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 25px;
            border: 1px solid rgba(255,255,255,0.05);
        }
        .chapters-list li {
            margin-bottom: 12px;
            color: var(--text-main);
        }
        .gallery-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin-top: 20px;
        }
        .gallery-grid img {
            width: 100%;
            border-radius: 8px;
            border: 1px solid rgba(255,255,255,0.1);
        }
    </style>
</head>
<body>
    <img src="banner.jpg" alt="Banner Oni no Ketsuryu Vol 3" class="header-banner">
    
    <div class="container">
        <div class="product-grid">
            <div>
                <img src="portada.jpg" alt="Portada Oni no Ketsuryu Vol 3" class="cover-img">
            </div>
            <div>
                <span class="badge">MANGA / LIGHT NOVEL</span>
                <h1>Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)</h1>
                <div class="author">Volumen 3: El Tren de las Sombras • Por Nicolás Noguera</div>
                <p>La batalla en el Tren Demoníaco. El choque del Maestro Celestial del Fuego contra el Tercer Estirpe. Y el legado sagrado de la Estilo de Dominio Solar.</p>
                <p>El Tren de las Sombras se ha convertido en una masa viva de carne demoníaca. Ren, Miyuki y el joven del haori de rayos combinan sus estilos para cortar el núcleo de la caldera, pero tras el descarrilamiento, el aterrador Rikudo —Tercer Estirpe de Sangre— emerge del bosque.</p>
                
                <div class="price-tag">$20.00 <span>USD</span></div>
                <a href="#" class="buy-btn">COMPRAR AHORA ($20 USD)</a>
            </div>
        </div>

        <h2 class="section-title">Contenido del Volumen 3</h2>
        <div class="chapters-list">
            <ul>
                <li><strong>Capítulo 1:</strong> La Carne del Tren (Escenas 1-3)</li>
                <li><strong>Capítulo 2:</strong> El Cuello de Hierro (Escenas 1-3)</li>
                <li><strong>Capítulo 3:</strong> La Llegada del Tercer Estirpe (Escenas 1-3)</li>
                <li><strong>Capítulo 4:</strong> La Última Llama (Escenas 1-3)</li>
                <li><strong>Capítulo 5:</strong> La Herencia del Fuego (Clímax del Volumen 3 - Escenas 1-3)</li>
            </ul>
        </div>

        <h2 class="section-title">Ilustraciones Interiores Destacadas</h2>
        <div class="gallery-grid">
            <img src="escena_1.jpg" alt="La Estilo de Esgrima del Trueno en el Tren">
            <img src="escena_climax.jpg" alt="La Herencia del Fuego">
        </div>
    </div>
</body>
</html>
"""

with open(os.path.join(target_root, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)
with open(os.path.join(target_sys, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)

print("Generated HTML landing pages for Vol 3 successfully")
