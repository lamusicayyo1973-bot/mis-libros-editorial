import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

def create_ebook_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("DE CERO A NEGOCIO CON IA\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    run_sub = title_p.add_run("Cómo lanzar y ejecutar tu emprendimiento en 90 días\n\nGUÍA PRÁCTICA • EMPRENDIMIENTO\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(234, 88, 12)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_page_break()
    
    # TOC Header
    p_toc_head = doc.add_paragraph()
    add_bookmark(p_toc_head, "TOC", 0)
    run_toc_h = p_toc_head.add_run("Tabla de Contenidos")
    run_toc_h.font.name = "Arial"
    run_toc_h.font.size = Pt(20)
    run_toc_h.font.bold = True
    run_toc_h.font.color.rgb = RGBColor(15, 23, 42)
    end_bookmark(p_toc_head, 0)
    
    index_text = """
• Introducción: Por qué ahora sí es tu momento
• Cómo usar este libro

PARTE 1: LA IDEA Y LA VALIDACIÓN
• Capítulo 1: Encuentra tu oportunidad de negocio
• Capítulo 2: Valida antes de construir
• Capítulo 3: Define tu oferta irresistible

PARTE 2: CONSTRUCCIÓN CON IA
• Capítulo 4: Tu stack de herramientas de IA
• Capítulo 5: Construye tu producto mínimo viable
• Capítulo 6: Automatiza antes de escalar
• Capítulo 7: Marca y presencia mínima viable

PARTE 3: CONSIGUE TUS PRIMEROS CLIENTES
• Capítulo 8: La estrategia de los primeros 10 clientes
• Capítulo 9: Contenido que atrae clientes
• Capítulo 10: Cierra ventas sin sentirte "vendedor"

PARTE 4: EJECUCIÓN Y CRECIMIENTO
• Capítulo 11: Sistemas para operar sin quemarte
• Capítulo 12: Cuándo y cómo escalar
• Capítulo 13: Tu plan de 90 días
• Cierre: Tu próximo paso
"""
    p_toc = doc.add_paragraph(index_text)
    p_toc.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Content sections
    sections_data = [
        ("Introducción: Por qué ahora sí es tu momento", """
Vamos a ser honestos desde la primera página: probablemente ya pensaste en emprender antes. Y probablemente también encontraste mil razones para no hacerlo: "No tengo capital", "No sé programar", "Alguien ya lo está haciendo mejor que yo", "¿Y si fracaso?".

Te entiendo. Yo también me hice esas preguntas. Pero algo cambió en los últimos años, y no es una moda pasajera: la inteligencia artificial bajó de golpe el costo de entrada a emprender. Lo que antes te tomaba un equipo de cinco personas y seis meses, hoy lo puedes armar tú solo, en semanas, con un presupuesto que cabe en una tarjeta de crédito.

No te voy a vender la fantasía de "hazte rico en 30 días con IA". Eso no existe. Lo que sí te voy a mostrar es un camino realista: cómo encontrar una oportunidad de negocio, validarla sin gastar de más, construir tu primera oferta con herramientas de IA, conseguir tus primeros clientes, y montar sistemas para que el negocio no dependa de que tú estés despierto 16 horas al día.

Mito 1: "Necesito saber programar." No. Vas a usar IA precisamente para no tener que aprender a programar desde cero.
Mito 2: "Necesito mucho capital." La mayoría de los negocios de este libro arrancan con menos de $200 dólares.
Mito 3: "La IA me va a robar la idea." La IA nivela el campo de juego de la ejecución, sí. Pero tu ventaja nunca fue saber usar una herramienta. Tu ventaja es entender un problema real de un grupo real de personas.
"""),
        ("PARTE 1: LA IDEA Y LA VALIDACIÓN", ""),
        ("Capítulo 1: Encuentra tu oportunidad de negocio", """
Aquí es donde la mayoría de la gente se traba. Se sienta a "pensar en una idea" como si la idea perfecta fuera a aparecer de la nada. No funciona así. Las buenas ideas de negocio casi nunca nacen de una lluvia de ideas; nacen de observar problemas reales que ya existen.

Cuatro caminos principales para construir un negocio apalancado en IA:
1. Productiza un servicio: Tomas una habilidad que ya tienes y la conviertes en un servicio empaquetado.
2. Automatización o agente para un nicho: Resuelves una tarea repetitiva específica.
3. Contenido y educación: Enseñas algo apoyándote en IA para producir más rápido.
4. Producto digital ligero: Una herramienta o plantilla que resuelve un problema puntual.

Prompt útil para investigación de mercado:
"Actúa como analista de mercado. Quiero entender los principales problemas que enfrentan [tipo de negocio/audiencia] al momento de [tarea específica]. Dame una lista de posibles dolores ordenados por costo/frustración."

TU EJERCICIO:
Completa la matriz de oportunidades con al menos 10 combinaciones antes de pasar al siguiente capítulo.
"""),
        ("Capítulo 2: Valida antes de construir", """
Voy a confesarte algo: la primera vez que emprendí, construí durante tres meses algo que nadie quería. El error no fue la ejecución, fue saltarme la validación.

Cómo validar en 48 horas:
Paso 1: Escribe 5 preguntas abiertas sobre cómo las personas manejan hoy el problema.
Paso 2: Busca a esas personas donde ya están. Ofrece 15 minutos de conversación.
Paso 3: Usa IA para resumir patrones de respuesta.

Señales de que SÍ hay negocio:
• La gente ya está pagando por algo (aunque sea una mala solución).
• Cuando describes tu idea, la reacción es "¿dónde me anoto?".
"""),
        ("Capítulo 3: Define tu oferta irresistible", """
Una buena idea sin una oferta clara no vende nada.

Estructura de una oferta clara:
1. ¿Cuál es el problema que resuelves?
2. ¿Cómo lo resuelves? (tu método en una línea).
3. ¿Qué resultado específico obtiene el cliente?

Plantilla: "Ayudo a [audiencia específica] a lograr [resultado específico] sin [el dolor principal] a través de [tu método]."
"""),
        ("PARTE 2: CONSTRUCCIÓN CON IA", ""),
        ("Capítulo 4: Tu stack de herramientas de IA", """
No caigas en la parálisis de herramientas. Piensa en tu stack en categorías:
• Texto y redacción: para escribir copy y propuestas.
• Imagen y diseño: para piezas visuales y logotipo.
• Automatización: para conectar herramientas entre sí.
• Atención al cliente: para responder preguntas frecuentes.
• Organización: para gestionar clientes y tareas.

Presupuesto realista: entre $0 y $50 mensuales en herramientas durante el primer mes.
"""),
        ("Capítulo 5: Construye tu producto mínimo viable (MVP)", """
Construye lo mínimo necesario para entregarle valor real a tu primer cliente.

El proceso paso a paso:
1. Define el entregable exacto.
2. Divide el proceso en pasos de principio a fin.
3. Identifica qué pasos puede acelerar la IA.
4. Ejecuta el proceso completo con un caso de prueba.
"""),
        ("Capítulo 6: Automatiza antes de escalar", """
Automatizar demasiado pronto es tan riesgoso como no automatizar nunca.

Qué automatizar primero:
• Confirmación y recordatorios de entregas.
• Seguimiento automático a leads.
• Envío de onboarding a nuevos clientes.

Qué NO automatizar todavía:
• Cualquier interacción donde el cliente decida si comprar o no.
• Procesos que cambian semana a semana.
"""),
        ("Capítulo 7: Marca y presencia mínima viable", """
Necesitas una presencia lo suficientemente creíble para que un cliente confíe en pagarte.

Lo mínimo indispensable:
• Nombre claro y fácil de recordar.
• Página de aterrizaje simple que explique tu oferta.
• Tonos consistentes de comunicación.
"""),
        ("PARTE 3: CONSIGUE TUS PRIMEROS CLIENTES", ""),
        ("Capítulo 8: La estrategia de los primeros 10 clientes", """
Tus primeros 10 clientes casi nunca llegan de publicidad pagada. Llegan de outreach directo.

Método de outreach asistido por IA:
1. Identifica dónde está tu audiencia.
2. Usa IA para investigar el contexto de cada prospecto.
3. Personaliza cada mensaje.
"""),
        ("Capítulo 9: Contenido que atrae clientes", """
El contenido es para que la gente correcta te encuentre confiando en lo que hablas. Usa IA para generar borradores, pero inyecta tus propias historias y ejemplos.
"""),
        ("Capítulo 10: Cierra ventas sin sentirte vendedor", """
Vender bien no es manipular: es ayudar a alguien a tomar una decisión clara.

Guion consultivo:
1. Pregunta primero la situación actual.
2. Refleja el problema con sus palabras.
3. Presenta tu solución conectada a lo que dijeron.
"""),
        ("PARTE 4: EJECUCIÓN Y CRECIMIENTO", ""),
        ("Capítulo 11: Sistemas para operar sin quemarte", """
Documenta procesos con ayuda de IA. Cada vez que hagas una tarea repetible, redacta un procedimiento claro paso a paso.
"""),
        ("Capítulo 12: Cuándo y cómo escalar", """
Señales de que es momento de escalar: estás rechazando clientes por falta de tiempo y tu proceso de entrega es repetible.
"""),
        ("Capítulo 13: Tu plan de 90 días", """
Roadmap ejecutable:
• Días 1–15: Idea y validación.
• Días 16–30: Construcción mínima.
• Días 31–60: Primeros clientes.
• Días 61–90: Sistemas y crecimiento.

Cierre: Tu próximo paso es abrir una hoja en blanco y completar la matriz de oportunidades hoy mismo.
""")
    ]
    
    b_id = 1
    for title, text in sections_data:
        p_head = doc.add_paragraph()
        add_bookmark(p_head, f"Chapter_{b_id}", b_id)
        
        if title.startswith("PARTE"):
            run_h = p_head.add_run(title)
            run_h.font.name = "Arial"
            run_h.font.size = Pt(18)
            run_h.font.bold = True
            run_h.font.color.rgb = RGBColor(234, 88, 12)
        elif title.startswith("Capítulo") or title.startswith("Introducción"):
            run_h = p_head.add_run(title)
            run_h.font.name = "Arial"
            run_h.font.size = Pt(15)
            run_h.font.bold = True
            run_h.font.color.rgb = RGBColor(15, 23, 42)
        else:
            run_h = p_head.add_run(title)
            run_h.font.name = "Arial"
            run_h.font.size = Pt(13)
            run_h.font.bold = True
            
        end_bookmark(p_head, b_id)
        b_id += 1
            
        if text.strip():
            p_body = doc.add_paragraph(text.strip())
            p_body.style.font.size = Pt(11)
            
    target_path = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\de-cero-a-negocio-con-ia\libro.docx"
    doc.save(target_path)
    print(f"File updated with TOC bookmarks: {target_path}")

if __name__ == "__main__":
    create_ebook_docx()
