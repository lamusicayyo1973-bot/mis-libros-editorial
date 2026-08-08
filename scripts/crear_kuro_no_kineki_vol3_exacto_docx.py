import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Dark Blue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_kuro_vol3_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\kuro-no-kineki-volumen-3"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Kuro no Kineki (黒の軌跡 - Ecos de Tinta Negra)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 3: El Despertar de los Creadores", 2, 1, "subtitulo_kuro3")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: El Continente de Plata", "cap_1"),
        ("Capítulo 2: Los Archivos del Olvido", "cap_2"),
        ("Capítulo 3: La Batalla de la Tinta de Plata", "cap_3"),
        ("Capítulo 4: La Capital del Olvido", "cap_4"),
        ("Capítulo 5: El Lienzo Blanco (Gran Final de la Trilogía)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: El Continente de Plata",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: El Viaje en el Océano de Tinta",
                    "texto": """Seis meses después de la caída de Aetheria. Un barco mercante de madera reforzada navegaba por el Mar del Olvido, un océano cuyas aguas no eran de agua salada, sino de una tinta líquida plateada que reflejaba las estrellas como si fuera un espejo perfecto.

En la proa del barco, Kael observaba el horizonte con su abrigo viajero movido por el viento marino. Su ojo derecho brillaba con un patrón rúnico más complejo, capaz de leer las corrientes de memoria que fluían en el mar.

A su lado, Sora afinaba las cuerdas de un instrumento antiguo mientras Kael revisaba los mapas de navegación sobre una mesa de cubierta.

—El mapa antiguo dice que este continente fue sellado hace mil años —dijo Kael, señalando una masa de tierra rodeada de tormentas—. Nadie que haya cruzado la bruma plateada ha regresado para contarlo.

—No necesitamos un mapa completo —respondió Kael, tocando la cicatriz de su mano izquierda, que latía en dirección al centro del continente—. La energía que despertó en la torre nos está llamando. Si los "Creadores" fueron los que diseñaron las tres dagas originales, este es el lugar donde empezó nuestra maldición.""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Desembocadura del Abismo",
                    "texto": """Una ola gigantesca de tinta plateada se levantó frente a la embarcación. De las aguas emergió una criatura colosal: un dragón marino hecho de cristal negro y runas doradas que emitía un rugido sónico capaz de agrietar el casco del barco.

—¡Es un Guardián de la Niebla! —gritó Sora, desenvainando la daga que había restaurado con su propio poder.

Kael reaccionó al instante, invocando un escudo de luz dorada desde sus palmas para proteger el mástil principal del barco.

Kael no esperó. Saltó hacia la cabeza del monstruo. Esta vez no necesitó una daga física: la energía de los tres elementos (fuego negro, fuego blanco y luz dorada) brotó directamente de las marcas rúnicas grabadas en sus palmas.

Con un solo movimiento cruzado en el aire, Kael trazó un símbolo de anulación que congeló al dragón de cristal en mitad del ataque, convirtiéndolo en una estatua de sal plateada que se desmoronó sobre el mar.

—Tus poderes no han dejado de evolucionar desde que absorbiste el Núcleo... —dijo Kael, mirando la facilidad con la que Kael había controlado la magia.

—No es mi poder —dijo Kael cayendo de pie en la cubierta—. Es el recuerdo de la magia original.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Costa de los Gigantes",
                    "texto": """El barco encalló suavemente en la playa de arena negra del continente olvidado. Frente a ellos no había bosques ni ciudades, sino las ruinas de estatuas de doscientos metros de altura que representaban a seres con cuatro ojos y túnicas ceremoniales: los verdaderos Creadores.

Al fondo, alzándose sobre una cordillera de montañas afiladas, se erigía la Torre de la Primera Tinta, emitiendo un pulso de luz carmesí que teñía el cielo de la tarde.

—Llegamos —dijo Sora, ajustando las vendas de sus brazos—. Siento una presencia dentro de esa torre... pero no es un autómata ni una proyección. Es un alma viva.

Kael dio el primer paso sobre la arena negra, sintiendo cómo el suelo bajo sus botas vibraba con una frecuencia familiar.

—Esa alma... es quien nos dio la vida a los tres.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: Los Archivos del Olvido",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: El Interior de la Torre",
                    "texto": """El grupo se adentró en la estructura de la torre. El interior no estaba hecho de piedra ni de metal, sino de estanterías infinitas hechas de cristal líquido donde flotaban millones de orbes de memoria en suspensión.

Cada orbe guardaba la historia completa de un mundo que había existido antes que Aetheria.

En el centro del gran salón, flotando sobre un pedestal de runas rojas, se encontraba un ser anciano con túnicas ceremoniales rotas. Carecía de rostro humano; en su lugar, tres máscaras flotantes de oro giran a su alrededor, representando la Alegría, el Dolor y la Indiferencia.

—Los tres fragmentos han regresado a casa —habló el ser. Su voz no resonaba en el aire, sino en las mentes de los tres hermanos a la vez—. Kael... el recipiente de la tinta. Sora... la custodia de la forma. Kael... el guardián de la luz.

—¿Quién eres? —exigió Kael, poniéndose en guardia.

—Soy el Primer Escriba —respondió la entidad—. El último superviviente de los Creadores. Y quienes ustedes llaman 'humanos' no son más que los bocetos de nuestra última obra.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Verdad del Primer Mundo",
                    "texto": """El Primer Escriba levantó la mano y los orbes de memoria del salón comenzaron a girar a gran velocidad, proyectando una ilusión tridimensional en mitad de la habitación.

Kael, Sora y Kael vieron el origen de todo:
Mil años atrás, la humanidad no vivía dividida entre la Superficie y el Abismo. Vivían en una civilización próspera, pero destruida por una guerra civil catastróficamente violenta. Para evitar la extinción total, los Creadores decidieron diseñar el Sistema de Tinta: un mecanismo para borrar la capacidad de odiar de la mente humana, al costo de sacrificar los recuerdos personales.

—Creamos el Relicario y dividimos las tres dagas entre los descendientes de nuestra propia sangre para equilibrar el sistema —explicó el Escriba—. Pero Aetheria fue una prueba fallida. La Emperatriz se corrompió con el poder. Y ahora que habéis apagado el Núcleo, la barrera que protegía este mundo del resto del universo se ha roto.

—¿Qué quieres decir con 'el resto del universo'? —preguntó Sora con una punzada de pánico.

El Escriba señaló hacia el techo de la torre, que se volvió transparente, mostrando el cielo nocturno donde miles de puntos rojos comenzaban a brillar como estrellas caídas.

—Otras civilizaciones que mantuvieron sus recuerdos intactos... vienen a reclamar la Tinta Original.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Ataque de los Invasores",
                    "texto": """Antes de que Kael pudiera procesar la revelación, el techo de la torre fue atravesado por tres lanzas de energía carmesí. Un grupo de figuras con armaduras orgánicas y alas de cristal oscuro descendió desde el cielo, destruyendo las estanterías de recuerdos.

Eran los Devoradores de Ecos, guerreros del continente exterior que buscaban absorber la Tinta Original de Kael.

—¡Entrega el contenedor! —gritó el líder de los invasores, desenvainando una guadaña de fuego rojo.

Kael y Sora reaccionaron de inmediato. Kael desató una ráfaga de luz dorada que desvió la guadaña, mientras Sora trazó barreras de tinta blanca para proteger al Escriba y los archivos.

Kael apretó los puños. Las runas de sus palmas volvieron a encenderse, pero esta vez la tinta negra no salía de su cuerpo... la tinta del suelo y del mar plateado comenzó a responder a su llamado.

—No voy a permitir que nadie vuelva a robar los recuerdos de este mundo —sentenció Kael.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: La Batalla de la Tinta de Plata",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: El Poder del Escriba",
                    "texto": """El combate en la torre se volvió un caos de energía. Los Devoradores de Ecos se movían con una coordinación perfecta, atacando en ráfagas que desgastaban la defensa de Kael y Sora.

El Primer Escriba, flotando en el centro del salón, comenzó a cantar en un idioma antiguo. Las tres máscaras doradas que flotaban sobre su cabeza se unieron en un solo rostro de piedra.

—Los bocetos no deben ser destruidos antes de que el trabajo esté terminado —dijo el Escriba.

Con un gesto de sus dedos, miles de orbes de memoria flotantes se convirtieron en agujas de luz que atravesaron a los invasores, sellando sus mentes y congelándolos en estatuas de piedra. Pero el esfuerzo fue demasiado para el anciano: su cuerpo de luz comenzó a agrietarse y disolverse.

—Escriba... —Kael corrió hacia el pedestal para sostenerlo.

—Mi tiempo ha terminado, Kael —susurró la entidad, tocando la frente del joven con su mano de piedra—. Te he transmitido la Llave del Autor. Ahora tú eres el único que puede reescribir las reglas de este mundo.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Fusión de los Tres Mundos",
                    "texto": """Al recibir la Llave del Autor, la mente de Kael se expandió más allá de los límites humanos. Por primera vez desde que despertó en la fosa del Capítulo 1, Kael no solo recuperó la totalidad de sus propios recuerdos, sino que vio el mapa completo del planeta: tres continentes flotantes aislados entre sí, cada uno sufriendo bajo sistemas de poder destructivos.

Sora y Kael se acercaron a él, sintiendo la inmensa presión mágica que emanaba de su hermano.

—¿Recuerdas todo, Kael? —preguntó Sora con la voz temblorosa.

Kael abrió los ojos. La pupila de su ojo derecho ya no tenía una sola runa, sino una galaxia entera de engranajes dorados, blancos y negros girando en perfecta armonía.

—Lo recuerdo todo, Sora. Recuerdos de nuestra infancia, el día que hicimos el pacto... y el motivo por el cual la Emperatriz nos separó —dijo Kael mirando a sus dos hermanos—. Pero más importante: ahora sé cómo salvar a los otros continentes.

Kael sonrió y levantó su espada.

—Entonces no hay tiempo que perder. La flota de los Devoradores viene en camino.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Portal de los Reyes",
                    "texto": """Desde el techo destruido de la torre, el grupo observó la armada enemiga: cientos de naves orgánicas flotando en la estratósfera, preparándose para bombardear la costa de la playa de arena negra.

Kael no mostró pavor. Extendió ambos brazos hacia el cielo.

La Llave del Autor reaccionó. El mar plateado a sus espaldas se levantó en tres columnas gigantescas de Tinta de Anulación que alcanzaron las nubes, formando un portal intercontinental de dimensiones bíblicas.

—No vamos a defender este continente —dijo Kael, mirando a Sora y Kael—. Vamos a llevar la batalla a la capital de los Devoradores.

Sora y Kael se colocaron a los lados de Kael. Los tres hermanos cruzaron sus manos en el centro del portal, desatando una luz que iluminó todo el océano plateado.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: La Capital del Olvido",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Desembarco en el Continente Rojo",
                    "texto": """El portal los transportó directamente sobre la metrópolis enemiga: Ignis-Null, una ciudad de agujas de hierro negro construida dentro de un cráter volcánico. A diferencia de Aetheria, aquí no se robaban recuerdos para flotar, sino que se extraían las emociones de los ciudadanos para alimentar armas de destrucción masiva.

Al pisar la plaza central de la ciudad enemiga, la alarma general resonó con un aullido de sirenas de bronce.

Decenas de legiones de soldados de hierro los rodearon de inmediato.

—Parece que no somos bienvenidos —bromeó Kael, desenfundando su mandoble de luz.

—Déjenme la vanguardia a mí —dijo Sora, abriendo los brazos mientras miles de agujas de tinta blanca flotaban a su alrededor como un ejército de mariposas de cristal.

Kael dio un paso al frente. La Llave del Autor en su mano izquierda se transformó en una pluma de hierro oscuro capaz de alterar la materia del entorno con solo trazar líneas en el aire.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Rey sin Nombre",
                    "texto": """Desde el palacio del cráter emergió el gobernante absoluto de Ignis-Null: un ser de tres metros de altura cubierto por una armadura de obsidiana que sostenía el Relicario de la Emoción, el artefacto hermano de las tres dagas de Kael.

—Bocetos de Aetheria... —dijo el Rey sin Nombre con una voz profunda que hizo temblar la tierra—. Han venido al lugar donde se forjó la primera gota de tinta. Aquí sus armas no tienen poder.

El Rey levantó su mano y la energía de las emociones robadas de millones de personas cayó sobre Kael como una ola de gravedad pura.

Kael cayó sobre una rodilla bajo la presión mental, pero en lugar de resistirse con fuerza física, usó la enseñanza del soldado Marcus del Volumen 1: aceptar el dolor ajeno para comprenderlo.

Kael cerró los ojos y tocó el suelo del cráter.

—Las emociones no son un combustible —susurró Kael—. Son lo que nos hace mantenernos en pie cuando todo lo demás se ha perdido.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Gran Borrador",
                    "texto": """Kael trazó un único trazo horizontal en el aire con la Pluma del Autor.

La energía de la Tinta de Anulación se extendió en un círculo perfecto por todo el cráter. El efecto fue instantáneo: las armas de los soldados enemigos se convirtieron en vapor inofensivo, y las emociones robadas que alimentaban al Rey fueron liberadas de la armadura de obsidiana, regresando en forma de luces de colores hacia las mentes de la población oprimida.

El Rey de obsidiana cediño sobre sus rodillas, viendo cómo su imperio de control se desmoronaba sin violencia.

—¿Qué... qué clase de magia es esta? —preguntó el rey derrotado.

—No es magia —respondió Kael caminando hacia él—. Es el borrador que limpia el lienzo para que puedan empezar de nuevo.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: El Lienzo Blanco (Gran Final de la Trilogía)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: La Reunión de los Mundos",
                    "texto": """Un mes después de la caída de Ignis-Null.

Gracias a la Llave del Autor, Kael no destruyó los tres continentes, sino que los unió físicamente en un único e inmenso mundo verde rodeado por el Océano de Plata. Aetheria, el Sector Cero e Ignis-Null ahora compartían el mismo horizonte bajo un cielo sin tiranos ni sistemas de succión de memoria.

En la cima de la antigua Torre del Escriba, los tres hermanos se reunieron por última vez ante el gran libro de la creación.

Kael entregó su mandoble dorado para ser sellado en la piedra. Sora guardó sus dagas en un cofre de bronce que colocaron en la base del monumento.

—¿Qué vas a hacer con la Pluma del Autor, Kael? —preguntó Sora mirando la herramienta en las manos de su hermano.

Kael contempló la pluma por un momento y luego la arrojó hacia el centro del mar plateado, dejando que se disolviera para siempre en el agua.

—Un mundo libre no necesita a nadie que escriba su destino desde las sombras —dijo Kael sonriendo—. A partir de hoy, cada persona escribirá su propia historia.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Libro Abierto",
                    "texto": """En la plaza central de la nueva capital unificada, la gente de todos los continentes se había reunido para celebrar el inicio de la era de la Tinta Libre. Los niños jugaban en las calles dibujando con tinta de colores sobre los muros de piedra que antes separaban los mundos.

Kael asumió el liderazgo del consejo civil para coordinar el comercio y la construcción de viviendas. Sora fundó la primera Academia de la Memoria, donde se enseñaba la historia real de la humanidad sin censuras ni manipulaciones.

Y en las afueras de la ciudad, junto a un árbol de hojas doradas, Kael escribía en su cuaderno de viajes.

Ya no era el joven amnésico que despertó en la fosa llena de sangre y bronce; era el hombre que había recordado el valor de cada vida que cruzó en su camino.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Última Página",
                    "texto": """Sora se acercó despacio por el sendero y se sentó al lado de Kael en la hierba.

—¿Terminaste de escribir tu libro? —preguntó ella mirándolo con cariño.

Kael cerró el cuaderno de cuero y se lo entregó en las manos. En la portada estaba grabado en letras de luz el título original de su aventura.

Sora abrió la última página. No había dibujos ni mapas; solo una frase escrita con la letra clara de Kael:

"Los recuerdos no nos definen por lo que fuimos en el pasado, sino por lo que elegimos proteger hoy."

Kael se puso de pie, ajustándose el abrigo viajero, y miró hacia el horizonte infinito donde el sol comenzaba a ponerse sobre el nuevo mundo reunificado.

Sora sonrió, guardó el cuaderno en su bolso y se puso de pie a su lado, lista para acompañarlo en el viaje que recién comenzaba.

[ FIN DE KURO NO KINEKI — VOLUMEN 3 ]
[ FIN DE LA TRILOGÍA ]""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Kuro Vol 3 at {output_docx}")

if __name__ == "__main__":
    create_kuro_vol3_manuscript()
