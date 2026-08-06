import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading_with_bookmark(doc, text, level, bookmark_id, bookmark_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_elem = p._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    p_elem.insert(0, bm_start)

    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p_elem.append(bm_end)

def create_kuro_vol1_manuscript():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    book_dir = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\kuro-no-kineki-volumen-1"
    
    portada_path = os.path.join(book_dir, "portada.jpg")
    if os.path.exists(portada_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(portada_path, width=Inches(5.0))
        doc.add_page_break()

    add_heading_with_bookmark(doc, "Kuro no Kineki (黒の軌跡 - Ecos de Tinta Negra)", 1, 0, "titulo_principal")
    add_heading_with_bookmark(doc, "Volumen 1: El Precio del Primer Paso", 2, 1, "subtitulo_kuro1")
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Autor: Nicolás Noguera | Formato: Manga / Light Novel Oficial")
    r_meta.font.italic = True
    r_meta.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    add_heading_with_bookmark(doc, "ÍNDICE DE CONTENIDOS", 2, 2, "toc_header")
    
    capitulos_info = [
        ("Capítulo 1: El Precio del Primer Paso", "cap_1"),
        ("Capítulo 2: Lazos de Tinta y Sangre", "cap_2"),
        ("Capítulo 3: La Ciudad de los Huecos", "cap_3"),
        ("Capítulo 4: El Ejecutor de la Superficie", "cap_4"),
        ("Capítulo 5: El Eclipse del Pobre (Clímax del Volumen 1)", "cap_5"),
    ]
    
    for cap_title, bookmark in capitulos_info:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Inches(0.5)
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(f"•  {cap_title}")
        r_toc.font.size = Pt(11)
        r_toc.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    doc.add_page_break()

    CAPITULOS_TEXTO = [
        {
            "titulo": "Capítulo 1: El Precio del Primer Paso",
            "bm_id": 10,
            "bm_name": "cap_1",
            "escenas": [
                {
                    "nombre": "Escena 1: El Despertar en la Fosa",
                    "texto": """El frío del metal fundido fue lo primero que sintió. Cuando Kael abrió los ojos, un dolor punzante le atravesó el cráneo. A su alrededor no había más que ruinas de bronce, columnas caídas y una niebla densa que olía a ozono y sangre.

No recordaba dónde estaba. No recordaba cómo había llegado allí. Peor aún: al intentar buscar en su mente quién era, solo encontró un abismo negro.

En su mano derecha apretaba con fuerza el mango de una daga de hierro oscuro. A sus pies yace el cuerpo de un guardián de piedra, destrozado y disolviéndose en charcos de tinta hirviente. Al mirarse en un reflejo de agua estancada, notó algo extraño: su ojo derecho ya no era humano; en el centro de la pupila brillaba un grabado rúnico en forma de engranaje blanco.

Una inscripción grabada en la pared frente a él parecía arder en luz propia:

"Para recordar quién eres, primero debes olvidar lo que más amabas."
""",
                    "img": "escena_c1_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Cacería del Colector",
                    "texto": """Un retumbo ensordecedor sacudió la caverna. Desde las sombras del techo abovedado descendió una figura grotesca: un gigante de cuatro metros, compuesto por placas de armadura oxidada y vendas sagradas que flotaban como tentáculos. Carecía de rostro; en su lugar, una gran rueda de bronce giraba sobre su cabeza.

—Sujeto Registro 409... —la voz del gigante resonó como dos placas de metal frotándose entre sí—. Tu cuota de recuerdos ha caducado. Entrega la materia prima.

El autómata extendió sus cuatro brazos mecánicos. De sus palmas brotaron lanzas de hierro que salieron disparadas directo al pecho de Kael.

El instinto de Kael actuó antes que su mente. Esquivó la primera lanza por milímetros, rodando sobre el suelo de piedra. Sabía que no tenía la fuerza física para destruir a semejante monstruo, pero una voz dentro de su cabeza le dictó la instrucción como un susurro antiguo:

«Usa la tinta. Paga el costo.»

Kael no lo dudó: clavó la punta de la daga en su propio antebrazo. La hoja absorbió su sangre y se tiñó de una llama negra deslumbrante. En ese instante, en la mente de Kael, el recuerdo de la primera vez que sostuvo una espada se desintegró como papel quemado. Olvidó a su maestro, olvidó la sonrisa de quien le enseñó a luchar. Pero a cambio, un poder descomunal recorrió sus venas.""",
                    "img": "escena_c1_e2.jpg"
                },
                {
                    "nombre": "Escena 3: Fuego Oscuro y Revelación",
                    "texto": """Kael se impulsó hacia adelante. El mundo pareció congelarse. Dejando una estela de fuego negro a su paso, cortó el aire en diagonal. La daga atravesó el torso blindado del Colector como si fuera manteca.

El gigante emitió un chirrido metálico y se desmoronó, convirtiéndose en una lluvia de polvo dorado y bronce.

Entre los restos del autómata, algo brilló. Kael se agachó y recogió un pequeño medallón de plata. Al presionarlo, el mecanismo se abrió revelando un retrato mágico en movimiento: en la imagen aparecía él mismo, unos años más joven y sonriendo, al lado de la mismísima Emperatriz que gobernaba la despiadada Ciudad Flotante de la superficie.

—¿Yo... al lado del tirano? —murmuró Kael, sintiendo que el suelo se le hundía.

De pronto, un aplauso lento e irónico resonó desde el balcón superior de las ruinas.

—Vaya, vaya... Veo que al fin despertaste de tu pequeña amnesia —dijo una voz femenina y serena.

Kael levantó la vista. Sobre la cornisa, bañada por la luz de la luna subterránea, una joven de túnica blanca y cabello plateado lo observaba. En su mano derecha sostenía una daga exactamente idéntica a la de él, envuelta en las mismas llamas oscuras.

—¿Quién eres? —exigió Kael, poniéndose en guardia.

La joven sonrió con amargura y saltó hacia el vacío, aterrizando sin hacer un solo ruido frente a él.

—Mi querido hermano... soy la persona que prometiste matar antes de perder la memoria.""",
                    "img": "escena_c1_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 2: Lazos de Tinta y Sangre",
            "bm_id": 20,
            "bm_name": "cap_2",
            "escenas": [
                {
                    "nombre": "Escena 1: Una Alianza Venenosa",
                    "texto": """Las palabras de Sora cayeron sobre Kael como un balde de agua helada. Su mano, apretada en torno al mango de la daga, comenzó a temblar.

—¿Hermano...? —repitió Kael, intentando escarbar en su mente. Pero el abismo de su memoria no le devolvió más que un dolor punzante tras las sienes—. Mientes. No recuerdo tu rostro. No recuerdo haber hecho ninguna promesa.

Sora dio un paso al frente. La llama negra de su propia daga se extinguió lentamente, transformándose en una fina hebra de humo que se reabsorbió en las cicatrices de sus antebrazos.

—Claro que no me recuerdas, Kael —dijo ella con una sonrisa gélida y melancólica—. Para obtener la fuerza suficiente para bajar hasta esta fosa y romper el primer sello de la Emperatriz, tuviste que quemar casi todo tu pasado. Me olvidaste a mí, olvidaste nuestra infancia... y olvidaste que fuiste tú quien planeó la rebelión en primer lugar.

Kael contempló el medallón en su mano. La imagen mágica de él junto a la Emperatriz seguía moviéndose en un bucle infinito: la gobernante sonreía con frialdad mientras le colocaba una insignia militar en el pecho.

—Si yo era su aliado, ¿por qué estoy atrapado aquí abajo? —preguntó él, alzando la mirada.

—Porque la traicionaste —respondió Sora, señalando el túnel oscuro que se adentraba hacia el corazón de la caverna—. Y porque el ejército de la superficie ya sabe que despertaste. Si no nos movemos ahora, el próximo en bajar no será un autómata... será la Guardia de Hierro.""",
                    "img": "escena_c2_e1.jpg"
                },
                {
                    "nombre": "Escena 2: Las Catacumbas de Bronce",
                    "texto": """Sin darle tiempo a procesar la información, Sora echó a correr hacia las profundidades de la cúpula. Kael dudó un segundo, pero el sonido distante de trompetas de guerra resonando desde la superficie lo obligó a seguirla.

Se adentraron en las Catacumbas de Bronce, una red de túneles industriales construidos siglo atrás. Gigantescas tuberías cruzaban las paredes de piedra, goteando un fluido denso y oscuro que alimentaba la maquinaria de la ciudad.

—Escúchame bien —dijo Sora sin frenar el paso—. Tu daga y la mía son dos mitades del Relicario de Tinta. Cada vez que usas una técnica, la daga consume un fragmento de tu alma. Si sigues luchando a ciegas, te quedarás 'hueco' antes de llegar al segundo nivel.

—¿Qué significa 'hueco'? —preguntó Kael, esquivando un engranaje oxidado que sobresalía del suelo.

—Significa convertirse en un cascarón sin mente. Como las criaturas que vagan en el sector bajo. Humanos que quemaron hasta el último de sus recuerdos y ahora solo buscan carne para llenar el vacío.

De pronto, Sora se detuvo en seco en una intersección de túneles. Delante de ellos, bloqueando el único paso, la penumbra comenzó a retorcerse. Decenas de ojos blancos y desorbitados brillaron en la oscuridad.""",
                    "img": "escena_c2_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Emboscada de los Huecos",
                    "texto": """No eran autómatas de metal. Eran figuras humanas demacradas, con la piel grisácea y cubiertas por jirones de ropa militar. Avanzaban a zancadas irregulares, emitiendo gemidos guturales mientras la tinta negra goteaba de sus bocas y oídos.

—Los Huecos... —susurró Sora, desenvainando nuevamente su arma—. No los ataques en la cabeza. Su núcleo está en el pecho, donde solía estar el corazón.

Un grupo de tres Huecos se abalanzó sobre Kael. El instinto lo llevó a canalizar la magia de su daga, pero al sentir el tirón en su mente —la amenaza inminente de perder otro recuerdo de su infancia—, contuvo el ataque y bloqueó con el metal de la hoja.

El impacto lo arrojó contra la pared de piedra.

—¡Kael, no seas idiota! —gritó Sora, mientras se desplazaba con una agilidad sobrehumana, trazando círculos de fuego blanco en el aire con su daga que descuartizaban a los atacantes—. Si no usas la tinta, te matarán.

—¡Tiene que haber otra forma! —exclamó Kael, levantándose rápidamente.

En medio del caos, Kael notó algo: en el pecho del Hueco más cercano, la tinta negra pulsaba al mismo ritmo que la runa de su propio ojo derecho. Sin usar el fuego de la daga, Kael extendió su mano izquierda y tocó el tórax de la criatura.

Una descarga eléctrica recorrió su brazo. En lugar de quemar un recuerdo propio, la runa de su ojo absorbió la tinta del enemigo. El Hueco colapsó instantáneamente, convertido en ceniza, mientras una visión ajena —el recuerdo de la vida de un soldado desconocido— cruzó velozmente por la mente de Kael.

Sora se congeló en mitad del combate, mirando a su hermano con verdadero terror en los ojos.

—Tú... ¿acabas de devorar su memoria? —murmuró ella, dando un paso atrás.""",
                    "img": "escena_c2_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 3: La Ciudad de los Huecos",
            "bm_id": 30,
            "bm_name": "cap_3",
            "escenas": [
                {
                    "nombre": "Escena 1: El Eco del Soldado",
                    "texto": """El silencio cayó pesado sobre las catacumbas tras la disipación de la horda. Kael cayó de rodillas, sujetándose la cabeza con ambas manos. La mente le ardía. El recuerdo robado no le pertenecía, pero resonaba en sus sienes con una claridad desgarradora: la imagen de una niña pequeña esperando junto a la puerta de una choza de madera, esperando a un padre que jamás regresarías.

—Ese soldado... tenía una familia —susurró Kael con la voz entrecortada—. Su nombre era Marcus.

Sora no se movió de su posición. Su rostro, habitualmente inescrutable, mostraba una mezcla de fascinación y temor. Bajó la daga lentamente, dejando que el humo se disipara.

—Ningún usuario del Relicario de Tinta ha sido capaz de hacer eso —dijo Sora, acercándose con pasos cautelosos—. La magia de tinta es una vía de un solo sentido: entregas tu alma para obtener fuerza. Devorar los ecos de otros debería haber fragmentado tu mente al instante.

—No elegí hacerlo, simplemente ocurrió —respondió Kael, poniéndose de pie con dificultad. La runa de su ojo derecho latía ahora con un tono carmesí tenue antes de volver a su blanco rúnico habitual—. Pero sentí algo más. Este lugar no es solo un laberinto... es una prisión masiva.

Sora miró hacia el fondo del pasaje, donde la arquitectura de piedra comenzaba a dar paso a estructuras más amplias e iluminadas por faroles de gas.

—Bienvenido al Sector Cero, Kael. La ciudad a la que la Superficie le da la espalda.""",
                    "img": "escena_c3_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Mercado de las Sombras",
                    "texto": """Al salir del túnel, la vista se abrió ante una cavidad colosal. Construida en los cimientos de la tierra, La Ciudad de los Huecos se alzaba como una maraña de chozas de chapa, puentes colgantes de madera oxidada y torres hechas con restos de maquinaria victoriana.

Miles de personas caminaban en penumbra por las calles embarradas. Algunos vendían fragmentos de metal; otros permanecían sentados contra los muros, con la mirada completamente perdida en la nada: humanos en distintas etapas del vaciado de memoria.

—Aquí terminan todos los que la Ciudad Flotante ya no necesita —explicó Sora mientras se cubría la cabeza con la capucha de su túnica—. Cuando un obrero o un soldado pierde la memoria por el trabajo en las minas de tinta, lo arrojan por los conductos de desecho.

Kael caminaba observando los rostros de la gente. Sintió una punzada de culpa al comprender que, según Sora, él había formado parte del régimen que administraba este sistema.

—Si yo fui el comandante de la Guardia de la Emperatriz... ¿por qué intenté destruirlo todo?

Sora se detuvo frente a un enorme portón de hierro blindado que conducía a los niveles superiores.

—Porque descubriste lo que hay en el corazón de la Ciudad Flotante —respondió ella en voz baja—. El Gran Núcleo no se alimenta de magia mineral, Kael. Se alimenta de las memorias de toda la población de la superficie. Les roban la historia para mantener la ciudad flotando en el cielo.""",
                    "img": "escena_c3_e2.jpg"
                },
                {
                    "nombre": "Escena 3: La Sombra del Ejecutor",
                    "texto": """Antes de que Kael pudiera responder, el aire del lugar se volvió helado. Los faroles de gas de toda la calle se apagaron en secuencia, dejando la ciudad en una penumbra casi absoluta. Un silencio sepulcral dominó a la multitud, que comenzó a dispersarse en pánico.

En lo alto del portón de hierro, una figura solitaria emergió de las sombras. Llevaba una capa blanca impoluta que contrastaba grotescamente con la mugre del ambiente, y una máscara de porcelana sin rasgos faciales, marcada únicamente por la insignia de un sol negro.

En su mano derecha sostenía un mandoble de dos metros de largo, rodeado de chispas doradas.

—Unidad de Traición 409... y la prófuga Sora —resonó una voz distorsionada y metálica desde detrás de la máscara—. Por orden de la Suprema Emperatriz, el juicio por la purga de la memoria comienza ahora.

Sora desenvainó su daga de inmediato, posicionándose frente a Kael.

—El primer Ejecutor de la Guardia de Hierro... —murmuró ella, apretando los dientes—. Nos encontraron más rápido de lo que pensaba.

Kael dio un paso al frente, colocándose al lado de su hermana. Apretó la empuñadura de su daga oscura, sintiendo cómo la memoria absorbida del soldado Marcus le otorgaba una nueva templanza para la batalla.

—Ya no huyo más —dijo Kael, mientras el engranaje de su ojo derecho comenzaba a girar.""",
                    "img": "escena_c3_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 4: El Ejecutor de la Superficie",
            "bm_id": 40,
            "bm_name": "cap_4",
            "escenas": [
                {
                    "nombre": "Escena 1: El Choque de la Luz y la Tinta",
                    "texto": """El Ejecutor no esperó. Con un movimiento casi invisible, se impulsó desde lo alto del portón de hierro. La gravedad pareció no afectarle; descendió como un meteoro dorado directo hacia Kael y Sora. El impacto de su mandoble contra el suelo agrietó la piedra y levantó una onda de choque que destruyó las estructuras de madera más cercanas.

—¡Separémonos! —gritó Sora, rodando hacia la izquierda mientras trazaba tres arcos de tinta blanca en el aire para frenar al enemigo.

Kael saltó hacia la derecha, usando el impulso para flanquear al Ejecutor. La hoja del mandoble dorado emitía una luz dorada cegadora: la Magia de Solarización, el poder exclusivo de los guardias de élite de la Emperatriz.

—Tu traición le costó cara a la corte, Kael —dijo el Ejecutor mientras desviaba los ataques de Sora con el lomo de su espada sin despeinarse—. Eras el perro más fiel de Su Majestad. Ver en lo que te has convertido me causa verdadera repugnancia.

El guerrero de blanco giró sobre su eje y lanzó un tajo horizontal. Una ráfaga de luz cortante voló hacia Kael. Sin tiempo para esquivar, Kael levantó su daga y activó la memoria robada del soldado Marcus.

La llama de la daga cambió de un negro puro a un matiz violeta profundo. El impacto de la ráfaga de luz contra el escudo de tinta hizo retroceder a Kael diez metros, haciendo crujir sus huesos, pero la barrera resistió.""",
                    "img": "escena_c4_e1.jpg"
                },
                {
                    "nombre": "Escena 2: La Verdad Detrás de la Máscara",
                    "texto": """—No puedes ganar usando los ecos de los muertos —sentenció el Ejecutor, caminando despacio entre el polvo—. Sus memorias son débiles, fragmentadas. La luz de la Emperatriz es eterna porque se nutre de la voluntad de millones.

Sora reapareció a la espalda del Ejecutor con una velocidad pasmosa. Clavó su daga directo en la juntura del cuello del guerrero. Sin embargo, el metal de su arma rebotó contra una barrera invisible de energía dorada.

—¡Kael, ahora! ¡Su escudo se debilita por un segundo tras bloquear un ataque físico! —advirtió Sora.

Kael no lo pensó. Comprendió que los recuerdos prestados no serían suficientes para atravesar esa defensa; debía arriesgar una parte de sí mismo. Cerró los ojos por un instante y tomó una decisión consciente: quemó el recuerdo del día en que aprendió a leer y escribir.

Una explosión de tinta negra pura brotó del mango de su daga. El fuego oscuro se volvió tan denso que la luz dorada del ambiente comenzó a ser absorbedora por la hoja. Kael se lanzó en un estallido de velocidad y clavó la daga de lleno en la máscara de porcelana del Ejecutor.

La máscara se agrietó y se rompió en mil pedazos.

El mandoble dorado cayó al suelo con un estruendo metálico. El Ejecutor retrocedió tambaleándose, llevándose la mano al rostro al descubierto. Kael se congeló al ver a su enemigo: no era un hombre anciano ni un monstruo. Era un joven de su misma edad, con rasgos faciales aterradoramente parecidos a los de Kael y Sora.""",
                    "img": "escena_c4_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Tercer Lazo",
                    "texto": """El Ejecutor cayó sobre una rodilla, respirando con dificultad mientras la tinta negra de la daga de Kael envenenaba su armadura dorada.

—Hermano... —murmuró el Ejecutor, alzando la mirada con ojos llenos de rabia y dolor—. Te enviaste a ti mismo al abismo... y dejaste que a mí me rehicieran a su imagen y semejanza.

Kael dio un paso atrás, sintiendo que la cabeza le daba vueltas. El recuerdo de las palabras recién quemadas le dejó un vacío extraño en el pensamiento, pero la revelación ante sus ojos era aún más devastadora.

—¿Tres...? —susurró Sora, dejando caer sus brazos con incredulidad—. La Emperatriz no creó dos dagas del Relicario... creó tres.

—Éramos los trillizos de la línea de sucesión —dijo el Ejecutor con una sonrisa amarga, mientras su cuerpo comenzaba a ser envuelto por una luz dorada de emergencia activa por la Superficie para recuperarlo—. Yo soy el que se quedó para servir. Tú eres el que olvidó para destruir. Y ella... es la que manipula a ambos.

Antes de que Kael o Sora pudieran reaccionar, un pilar de luz dorada descendió desde el techo de la cúpula, tragándose al Ejecutor herido y transportándolo de vuelta a la Ciudad Flotante.

En el suelo solo quedó la enorme grieta del combate y el mandoble dorado, cuyo brillo se apagaba lentamente.

Sora miró a Kael con una expresión ilegible.

—El portal de transporte dejó la frecuencia de la Superficie abierta por unos minutos —dijo ella, señalando el residuo de luz rúnica que aún flotaba en el aire—. Si entramos ahora, subiremos directamente al palacio. Pero no habrá vuelta atrás.

Kael apretó la daga de tinta, mirando la luz que conducía al origen de todas sus pesadillas.""",
                    "img": "escena_c4_e3.jpg"
                }
            ]
        },
        {
            "titulo": "Capítulo 5: El Eclipse del Pobre (Clímax del Volumen 1)",
            "bm_id": 50,
            "bm_name": "cap_5",
            "escenas": [
                {
                    "nombre": "Escena 1: El Ascenso a la Ciudad Flotante",
                    "texto": """El residuo de luz rúnica vibraba en el centro del mercado abandonado, emitiendo un zumbido agudo que hacía temblar los cristales de las pocas linternas intactas.

Kael y Sora se alinearon frente al portal. El aire a su alrededor se volvió tan pesado que cada respiración dolía. Sora miró a Kael de reojo, ajustando el agarre sobre su daga.

—En el instante en que crucemos la frecuencia de transporte —advirtió ella—, los sensores del Palacio Real detectarán la presencia de energía de tinta. No habrá sigilo. Nos recibirán en la plataforma de desembarco con todo el ejército.

—Que vengan —respondió Kael. La runa de su ojo derecho ya no brillaba con la luz titilante del principio, sino con un destello blanco fijo e intenso—. Prefiero enfrentarlos sabiendo que la verdad está arriba, y no pudriéndome en el olvido de esta fosa.

Los dos hermanos dieron el paso al frente al mismo tiempo.

La luz dorada los envolvió al instante, descomponiendo sus siluetas en destellos de energía. Durante tres segundos de vértigo puro, la Ciudad de los Huecos, las catacumbas y el hedor a humedad desaparecieron. En su lugar, el paisaje se abrió de golpe ante la luz enceguecedora del sol real.

Aparecieron en la plaza central de la Ciudad Flotante de Aetheria: una metrópolis suspendida sobre las nubes, construida con mármol blanco, agujas de cristal y jardines flotantes. Pero bajo la aparente belleza celestial, Kael pudo ver las enormes tuberías de succión de memoria que se clavaban en el abismo, extrayendo el fluido oscuro de los habitantes del suelo.""",
                    "img": "escena_c5_e1.jpg"
                },
                {
                    "nombre": "Escena 2: El Salón de la Emperatriz",
                    "texto": """Decenas de soldados blindados con armaduras doradas rodeaban la plaza, pero ninguno se atrevió a dar un paso al frente. En el extremo superior de la gran escalinata de mármol, sentado sobre un trono suspendido que flotaba entre dos anillos rúnicos, se encontraba la figura que Kael había visto en el medallón.

La Suprema Emperatriz.

A su derecha, con la armadura agrietada y el rostro al descubierto, permanecía de pie el tercer hermano, el Ejecutor derrotado en las catacumbas.

—Mis tres creaciones al fin reunidas —habló la Emperatriz. Su voz no era violenta; sonaba con una serenidad maternal y aterradora que resonaba directamente dentro de las mentes de los tres—. Kael... pensaste que tu rebelión fue un acto de libre albedrío. Pensaste que olvidar era el costo que pagabas para destruirme.

La Emperatriz levantó su mano derecha. En su palma flotaba la esfera del Gran Núcleo, una gema negra del tamaño de una cabeza humana que latía al ritmo de miles de recuerdos robados.

—Fuiste tú quien diseñó el sistema de memoria, Kael —reveló la Emperatriz con una leve sonrisa—. Creaste las tres dagas para fragmentar tu propia mente porque sabías que el peso de lo que hicimos para mantener esta ciudad en el cielo te destruiría la cordura. No eres el héroe de esta historia... eres el arquitecto de esta prisión.

Kael sintió que la tierra temblaba bajo sus pies. Sora dio un paso atrás, mirando a Kael con horror y conmoción.""",
                    "img": "escena_c5_e2.jpg"
                },
                {
                    "nombre": "Escena 3: El Eclipse de Tinta (Cierre del Tomo 1)",
                    "texto": """Las palabras de la Emperatriz resonaban en la mente de Kael como un eco interminable. Por un segundo, la tentación de caer en la desesperación amenazó con apagar la luz de su ojo rúnico. Pero al mirar sus antebrazos —cubiertos de cicatrices y marcas de fuego negro— y recordar la mirada del soldado Marcus y la gente condenada en el Sector Cero, Kael apretó los dientes.

—Tal vez yo creé esta monstruosidad cuando tenía memoria —dijo Kael, dando un paso al frente en la escalinata—. Pero el Kael que soy hoy... el que no tiene pasado, el que no tiene títulos... es el que va a destruirla.

Kael no esperó. Cruzó su daga de tinta negra con la daga de fuego blanco de Sora.

Al colisionar ambas armas, una reacción en cadena sacudió el palacio. El poder combinado de las dos dagas no quemó un recuerdo individual: desató una onda de choque de Tinta de Anulación que se extendió por toda la plaza, apagando momentáneamente el brillo del Gran Núcleo y agrietando los cimientos de mármol de la Ciudad Flotante.

El cielo se oscureció de golpe cuando una cortina de tinta negra eclipsó el sol sobre Aetheria.

La Emperatriz se puso de pie por primera vez, borrándose la sonrisa de su rostro, mientras Kael y Sora se abalanzaban por la escalinata hacia el trono para la batalla final por el destino de ambos mundos.

[ CONTINUARÁ EN EL VOLUMEN 2 ]""",
                    "img": "escena_c5_e3.jpg"
                }
            ]
        }
    ]

    for cap in CAPITULOS_TEXTO:
        add_heading_with_bookmark(doc, cap["titulo"], 1, cap["bm_id"], cap["bm_name"])
        
        for esc in cap["escenas"]:
            add_heading_with_bookmark(doc, esc["nombre"], 2, cap["bm_id"]+1, f"{cap['bm_name']}_esc")
            
            p_body = doc.add_paragraph()
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.space_after = Pt(8)
            r_body = p_body.add_run(esc["texto"])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(11)
            
            img_filename = esc["img"]
            img_path = os.path.join(book_dir, img_filename)
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(14)
                r_img = p_img.add_run()
                r_img.add_picture(img_path, width=Inches(5.2))
                
        doc.add_page_break()

    output_docx = os.path.join(book_dir, "libro.docx")
    doc.save(output_docx)
    print(f"Generated EXACT PROMPTS docx for Kuro Vol 1 at {output_docx}")

if __name__ == "__main__":
    create_kuro_vol1_manuscript()
