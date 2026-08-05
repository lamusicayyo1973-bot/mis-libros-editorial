import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

def create_ebook_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("EL ALGORITMO PERSONAL\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(14, 165, 233)
    
    run_sub = title_p.add_run("Rediseñá tus hábitos, dominá tu enfoque y ejecutá con claridad\n\nGUÍA PRÁCTICA • DESARROLLO PERSONAL Y PRODUCTIVIDAD\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(234, 88, 12)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_page_break()
    
    # TOC Header
    p_toc_head = doc.add_paragraph()
    add_bookmark(p_toc_head, "TOC", 0)
    run_toc_h = p_toc_head.add_run("Tabla de Contenidos")
    run_toc_h.font.name = "Arial"
    run_toc_h.font.size = Pt(20)
    run_toc_h.font.bold = True
    run_toc_h.font.color.rgb = RGBColor(15, 23, 42)
    end_bookmark(p_toc_head, 0)
    
    index_text = """
• Capítulo 1: La Trampa de la Motivación y el Poder del Sistema
• Capítulo 2: Auditar el Entorno: Cómo eliminar la fricción invisible
• Capítulo 3: La Regla del Micro-Aumento: La ciencia de la consistencia
• Capítulo 4: Gestión de la Energía: Enfoque profundo en un mundo hiperconectado
• Capítulo 5: El Manual de Ejecución: Tu protocolo diario de alto rendimiento
"""
    p_toc = doc.add_paragraph(index_text)
    p_toc.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    sections_data = [
        ("Capítulo 1: La Trampa de la Motivación y el Poder del Sistema", """
1. El mito de la fuerza de voluntad
Casi todas las personas que fracasan en cambiar una conducta cometen el mismo error inicial: confían en la motivación.

La motivación es una emoción, y como cualquier emoción, es volátil, reactiva y efímera. Esperar a "sentir ganas" para entrenar, estudiar, trabajar en tu proyecto personal o comer saludable es una estrategia destinada al fracaso. La fuerza de voluntad no es una virtud moral; es un recurso finito que se agota a lo largo del día con cada decisión que tomás.

Los superrealizadores no tienen más fuerza de voluntad que el resto; simplemente han diseñado un sistema donde no necesitan usarla de forma constante.

Idea clave: No te elevás al nivel de tus metas; caés al nivel de tus sistemas.

2. La arquitectura del hábito
Un hábito no es más que una solución automatizada a un problema recurrente en tu entorno. Tu cerebro busca constantemente ahorrar energía, y para hacerlo, transforma secuencias de acciones en rutinas automáticas.

Todo hábito sigue un bucle de cuatro pasos:
- La Cue (Señal): El detonante ambiental o interno que indica a tu cerebro que inicie una conducta.
- El Anhelo: La motivación o el cambio de estado emocional que esperás obtener.
- La Respuesta: La acción concreta que ejecutás.
- La Recompensa: El premio final que satisface el anhelo y le enseña a tu cerebro a repetir el bucle.

Para cambiar un resultado en tu vida, no intentes cambiar tu "deseo". Tenés que intervenir en la Señal y en la Respuesta. Si querés reducir una conducta negativa, hacela invisible o difícil. Si querés integrar una conducta positiva, hacela obvia e inevitable.

3. Del objetivo a la identidad
El motivo principal por el cual la mayoría de los propósitos de año nuevo fracasan antes de febrero es que están enfocados en el resultado ("Quiero bajar 5 kilos", "Quiero escribir un libro") en lugar de la identidad ("Soy una persona que cuida su cuerpo", "Soy un escritor").

Cada acción que tomás es un voto a favor del tipo de persona en la que te querés convertir.
- Un día de entrenamiento no te transforma físicamente, pero es un voto a favor de tu identidad como atleta.
- Escribir dos páginas no completa un libro, pero es un voto a favor de tu identidad como creador.

Cuando cambiás el enfoque de lo que querés lograr a quién querés ser, el hábito deja de ser un esfuerzo y se convierte en una expresión de tu identidad.

4. Ejercicio Práctico: Tu Auditoría de Sistemas
- Identificá tu meta principal: Escribí en una oración qué querés lograr en los próximos 90 días.
- Desarmá la meta en una acción de 2 minutos: ¿Cuál es la versión mínima e indivisible de esa meta que podés hacer hoy sin excusas?
- Diseñá la señal: Vincularás este nuevo micro-hábito a una rutina que ya hacés automáticamente. Completa la frase: "Inmediatamente después de [HÁBITO ACTUAL], voy a [NUEVO MICRO-HÁBITO DE 2 MINUTOS]."
"""),
        ("Capítulo 2: Auditar el Entorno — Eliminar la Fricción Invisible", """
1. El entorno es la mano invisible
Tendemos a creer que somos los arquitectos absolutos de nuestras elecciones, pero la psicología conductual demuestra lo contrario: en la mayoría de los casos, somos simplemente el producto del entorno en el que nos movemos.

Si dejas el celular en la mesa de noche al lado de tu cama, lo vas a revisar apenas abras los ojos. No es una falla de tu carácter; es una respuesta automática a un estímulo accesible. Si en tu escritorio hay comida ultraprocesada a la vista, terminarás comiéndola cuando tu energía baje a mitad de la tarde.

Los hábitos no nacen del deseo interno, sino de la conveniencia visual y espacial. Quienes muestran un autocontrol excepcional en apariencia no son mártires de la disciplina; son arquitectos de entornos donde las tentaciones son invisibles y las buenas acciones son inevitables.

Principio clave: Es más fácil cambiar tu entorno que cambiar tu nivel de disciplina frente a una tentación constante.

2. La Ley de la Fricción
La fricción es la cantidad de energía, pasos o tiempo necesarios para ejecutar una acción. Tu cerebro está programado evolutivamente para seguir la ley del mínimo esfuerzo: entre dos opciones, siempre elegirá la que requiera menor resistencia inmediata.

Para tomar el control de tus hábitos, tenés que aplicar esta regla de forma estratégica:
- Para construir hábitos positivos: Reducí la fricción. Eliminá la mayor cantidad de pasos entre vos y la conducta deseada.
- Para eliminar hábitos negativos: Aumentá la fricción. Añadí pasos adicionales entre vos y la conducta que querés evitar.

3. Rediseño de Espacios: Una función por lugar
El cerebro asocia los espacios con las acciones que se realizan repetidamente en ellos. Para maximizar el enfoque y la claridad mental, aplicá el principio de un solo contexto por espacio:
- Zona de Trabajo Profundo: Un espacio destinado exclusivamente a producir, estudiar o crear.
- Zona de Descanso: Tu cama o un sillón destinado solo a dormir y relajarte.
- Zona Social / Ocio: Lugares destinados a comer, conversar o entretenerte.

4. Ejercicio Práctico: Limpieza de Fricción
- Auditoría de Distracciones: Identificá la mayor distracción que destruye tu concentración diaria.
- Aumentá 3 pasos de fricción: Escribí e implementá hoy mismo 3 barreras para esa distracción.
- Optimizá un hábito positivo: Elegí la acción más importante que querés hacer mañana y prepará todo hoy a la noche.
"""),
        ("Capítulo 3: La Regla del Micro-Aumento — La Ciencia de la Consistencia Ininterrumpida", """
1. El peligro del entusiasmo inicial
El mayor enemigo de la consistencia a largo plazo no es la pereza, sino el exceso de ambición inicial.

Cuando alguien decide cambiar su vida, suele hacerlo motivado por un pico de frustración o inspiración. En ese estado emocional, establece metas desproporcionadas. Durante los primeros tres o cuatro días, la energía alcanza para sostener el esfuerzo. Pero al primer imprevisto, el sistema colapsa.

Principio clave: Un hábito debe estar consolidado antes de ser optimizado. No podés mejorar algo que no existe de forma constante.

2. La Regla del 1% y el Interés Compuesto
Los cambios significativos no son el resultado de transformaciones drásticas de la noche a la mañana, sino del efecto acumulativo de pequeñas mejoras diarias.
Si mejorás un 1% cada día en una disciplina, al cabo de un año serás 37 veces mejor en esa área.
Mejora acumulada: (1.01)^365 = 37.78

3. El umbral del esfuerzo mínimo (La regla de no romper la cadena)
Para garantizar la consistencia ininterrumpida, cada hábito de tu vida debe tener dos versiones:
- La versión ideal: Lo que hacés en un día perfecto con tiempo, energía y motivación alta.
- La versión de emergencia (El Micro-Aumento): Lo mínimo indispensable que hacés en un día caótico, enfermo o saturado.

Regla de oro: Nunca te saltes dos días seguidos.

4. Ejercicio Práctico: Definir tus Niveles de Ejecución
- Nivel A (Día Ideal): Lo que hacés con todo el tiempo y energía disponible.
- Nivel B (Día Normal): Estándar aceptable cuando tu día transcurre de forma habitual.
- Nivel C (Día de Emergencia): Versión reducida al 5% para no romper la cadena.
"""),
        ("Capítulo 4: Gestión de la Energía — Enfoque Profundo en un Mundo Hiperconectado", """
1. El tiempo no se gestiona; se gestiona la energía
El error más común de la productividad moderna es tratar el tiempo como el recurso principal. Todos tenemos exactamente 24 horas al día, pero no todas las horas son iguales. Una hora de trabajo a las 9:00 AM con la mente fresca no produce el mismo impacto que una hora a las 11:00 PM tras un día agotador.

Principio clave: No necesitas más horas en tu día; necesitas más horas de atención indivisible.

2. Los cuatro bloques de energía biológica
- Bloque de Enfoque Profundo (Deep Work): 90 a 120 minutos dedicados exclusivamente a la tarea de mayor impacto.
- Bloque Operativo: Tareas secundarias de baja exigencia cognitiva.
- Bloque de Recarga: Pausas activas sin pantallas.
- Bloque de Desconexión: Límite claro al final de la jornada laboral.

3. Dieta dopaminérgica: Reducir el ruido digital
- Elimina las alertas pasivas.
- Agrupa las revisiones.
- Entrena el aburrimiento.

4. Ejercicio Práctico: Tu Protocolo de Enfoque Profundo
- Define un único objetivo: "En los próximos 60 minutos, lo único que voy a terminar es [TAREA CLAVE]".
- Aísla el entorno digital.
- Establece un temporizador visible (Pomodoro).
"""),
        ("Capítulo 5: El Manual de Ejecución — Tu Protocolo Diario para Mantener el Rumbo a Largo Plazo", """
1. La diferencia entre planificar y ejecutar
Tener una visión clara o un libro lleno de buenas intenciones no sirve de nada si no existe una estructura diaria que traduzca esas ideas en acciones tangibles.

Principio clave: La claridad elimina la indecisión. Cuando sabes exactamente qué hacer en cada momento, la fricción mental desaparece.

2. El Ritual de Apertura (Los primeros 30 minutos)
- Hidratación y movimiento mínimo.
- Revisión del objetivo del día (La Roca del Día).
- Bloqueo de agenda.

3. El Ritual de Cierre y la Auditoría Semanal
- Vaciar la mente.
- Preparar el entorno para mañana.
- Cierre consciente.
- Auditoría Semanal de 15 minutos.

4. Plantilla del Sistema: Tu Checklist de Ejecución Diaria
[ ] RITUAL DE APERTURA
    - Hidratación + movimiento
    - Definir la Roca del Día: [ _____________________________________ ]
    - Bloque de Enfoque Profundo agendado (90 min)

[ ] EJECUCIÓN
    - Cumplir el Bloque de Enfoque Profundo (sin teléfono ni distracciones)
    - Ejecutar la versión reducida (Nivel C) si el día se complica

[ ] RITUAL DE CIERRE
    - Procesar pendientes y vaciar la mente
    - Dejar el entorno preparado para mañana
    - Desconexión digital nocturna
""")
    ]
    
    b_id = 1
    for title, text in sections_data:
        p_head = doc.add_paragraph()
        add_bookmark(p_head, f"Chapter_{b_id}", b_id)
        
        run_h = p_head.add_run(title)
        run_h.font.name = "Arial"
        run_h.font.size = Pt(16)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(14, 165, 233)
            
        end_bookmark(p_head, b_id)
        b_id += 1
            
        if text.strip():
            p_body = doc.add_paragraph(text.strip())
            p_body.style.font.size = Pt(11)
            
    target_path = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\el-algoritmo-personal\libro.docx"
    doc.save(target_path)
    print(f"File created successfully: {target_path}")

if __name__ == "__main__":
    create_ebook_docx()
