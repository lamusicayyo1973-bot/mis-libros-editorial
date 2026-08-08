import os
import glob
import shutil
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_text, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_text)
    tag.append(start)

def end_bookmark(paragraph, bookmark_id):
    run = paragraph.add_run()
    tag = run._r
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)

src_dir = r"C:\Users\nicol\.gemini\antigravity\brain\6adf8ce5-9839-4292-a8a1-57beed4c3827"
target_root = r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-4"
target_sys = r"c:\Users\nicol\Downloads\MIS LIBROS\sistema_editorial\libros\oni-no-ketsuryu-volumen-4"

os.makedirs(target_root, exist_ok=True)
os.makedirs(target_sys, exist_ok=True)

matches = glob.glob(os.path.join(src_dir, 'oni_vol*.jpg'))
if matches:
    matches.sort(key=os.path.getmtime, reverse=True)
    latest_img = matches[0]
    for fn in ["portada.jpg", "thumbnail.jpg", "banner.jpg", "escena_1.jpg", "escena_climax.jpg"]:
        shutil.copy2(latest_img, os.path.join(target_root, fn))
        shutil.copy2(latest_img, os.path.join(target_sys, fn))

def create_oni_vol4_docx():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("ONI NO KETSURYŪ\n(鬼の血流 - La Estirpe de la Sangre)\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(220, 38, 38)
    
    run_sub = title_p.add_run("Volumen 4: El Distrito de los Espejos y la Mariposa de la Sombra\n\nMANGA / LIGHT NOVEL • FANTASÍA OSCURA & ACCIÓN SENGOKU\n\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(185, 28, 28)
    
    run_author = title_p.add_run("Por Nicolás Noguera\n\n\n")
    run_author.font.name = "Arial"
    run_author.font.size = Pt(14)
    run_author.font.bold = True
    run_author.font.color.rgb = RGBColor(71, 85, 105)
    
    img_cover_path = os.path.join(target_root, "portada.jpg")
    if os.path.exists(img_cover_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_cover_path, width=Inches(4.5))
        
    doc.add_page_break()
    
    # TOC Header
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = toc_p.add_run("TABLA DE CONTENIDOS\n")
    r_toc.font.name = "Arial"
    r_toc.font.size = Pt(20)
    r_toc.font.bold = True
    r_toc.font.color.rgb = RGBColor(220, 38, 38)
    
    chapters_toc = [
        ("Capítulo 1", "La Finca del Fuego", "chap1"),
        ("Capítulo 2", "La Ciudad que Nunca Duerme", "chap2"),
        ("Capítulo 3", "El Doble Estirpe", "chap3"),
        ("Capítulo 4", "La Purificación de la Sangre", "chap4"),
        ("Capítulo 5", "Las Cenizas de la Misericordia (Clímax del Volumen 4)", "chap5")
    ]
    
    for c_num, c_title, c_bkm in chapters_toc:
        p_t = doc.add_paragraph()
        r_num = p_t.add_run(f"• {c_num}: ")
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(185, 28, 28)
        r_name = p_t.add_run(c_title)
        r_name.font.italic = True
        
    doc.add_page_break()
    
    # Full Content Text
    content = [
        {
            "id": "chap1",
            "number": "Capítulo 1",
            "title": "La Finca del Fuego",
            "scenes": [
                {
                    "num": "Escena 1: Las Memorias del Primer Guerrero de la Hermandad",
                    "text": """Ren llegó a la Finca del Fuego al atardecer. La residencia, antes llena de vida, permanecía en un silencio sepulcral. El padre de Kenshin, un antiguo Maestro Celestial retirado y consumido por la bebida, lo recibió en el patio central.

Al ver la guarda con forma de llama en la katana negra de Ren y la marca carmesí en su mejilla, el anciano dejó caer su copa de sake.

—Esa marca... —murmuró el anciano con voz temblorosa—. No es una simple cicatriz de venas demoníacas. Es la Marca del Sol, el sello de los primeros espadachines que casi destruyen al Rey Oni hace mil años.

El anciano lo guio hasta la biblioteca secreta de la finca y le entregó un volumen antiguo con páginas de papel arroz desgastado: Los Diarios del Sol."""
                },
                {
                    "num": "Escena 2: La Danza del Sol",
                    "text": """Ren pasó dos días sin comer ni dormir dentro de la biblioteca, estudiando los grabados del diario. Las ilustraciones no mostraban posturas de esgrima complejas, sino una secuencia de doce movimientos de danza ritual que los herreros antiguos realizaban desde el amanecer hasta el anochecer para pedir la bendición del fuego.

Al cerrar los ojos, Ren recordó a su propio padre ejecutando esa misma danza frente al horno de la herrería cada año nuevo.

—No era solo un ritual... —comprendió Ren, poniéndose de pie con la katana en la mano—. La Estilo del Fuego Carmesí de Herrería que me enseñó mi padre era la versión disfrazada de la Danza del Sol.

Al ejecutar el primer movimiento dentro del dojo, una ola de calor puro envolvió la habitación, volviendo la hoja negra de su katana de un color dorado incandescente."""
                },
                {
                    "num": "Escena 3: La Petición de la Mariposa",
                    "img": "escena_1.jpg",
                    "text": """Al salir de la finca, una mariposa de luz morada revoloteó sobre la cabeza de Ren. Detrás de él reapareció la chica silenciosa de la Selección Final —la joven de la máscara de zorro, cuyo nombre era Aoi, la nueva discípula del Maestro Celestial del Ingesta y Veneno.

—El Gremio requiere tu presencia en el Distrito de las Luces —dijo Aoi con voz suave pero firme—. Un demonio de nivel Quinto Estirpe de Sangre ha convertido el distrito de entretenimientos en su terreno de cacería. Los cazadores enviados anteriormente han desaparecido sin dejar rastro.

Ren miró la caja de madera en su espalda, donde Miyuki descansaba.

—Nos ponemos en marcha —respondió Ren—. Esta vez no dejaremos que ningún inocente caiga."""
                }
            ]
        },
        {
            "id": "chap2",
            "number": "Capítulo 2",
            "title": "La Ciudad que Nunca Duerme",
            "scenes": [
                {
                    "num": "Escena 1: El Lujo de las Sombras",
                    "text": """El Distrito de las Luces era una metrópolis nocturna deslumbrante. Gigantescas linternas de seda roja, casas de te con cortinas de bambú y puentes de madera sobre canales de agua iluminaban la noche. Cientos de personas caminaban festejando con música de samisen.

Pero para los ojos rúnicos de Ren y el olfato desarrollado de Aoi, la ciudad olía a un veneno dulce y denso.

—El demonio no ataca en los callejones —explicó Aoi, ajustando las agujas de veneno de glicina en sus mangas—. Vive dentro de una de las casas de entretenimientos más prestigiosas como una matrona de alta sociedad.

Para infiltrarse sin llamar la atención de los guardias, Ren, Aoi y el muchacho del haori amarillo (que los alcanzó en el camino) tuvieron que disfrazarse con kimonos tradicionales para ingresar a la casa principal."""
                },
                {
                    "num": "Escena 2: La Matrona del Espejo",
                    "text": """Dentro de la casa Yoshiwara, la matrona principal —una mujer de belleza deslumbrante llamada Daki— descansaba sobre un diván de seda. Llevaba una faja obi de seda de diez metros de largo decorada con patrones de flores de cerezo.

Sin embargo, detrás de su reflejo en los espejos de bronce de la habitación, su verdadera forma se revelaba: una criatura de piel pálida con venas verdes y la marca del Quinto Estirpe de Sangre en ambos ojos.

—Guerreros de la Hermandad disfrazados... qué falta de respeto al arte —siseó Daki con voz melodiosa.

Con un movimiento de sus dedos, la faja obi de seda cobró vida propia, transformándose en láminas de acero flexible que atravesaron las paredes de madera del dojo."""
                },
                {
                    "num": "Escena 3: La Trampa de la Seda",
                    "text": """La faja obi se dividió en seis cintas independientes que atacaron desde todos los ángulos.

Aoi se desplazó con una agilidad impresionante, utilizando el Estilo del Veneno para cortar las cintas con sus dagas cortas llenas de veneno de glicina. Sin embargo, al cortar la seda, las cintas no sangraban: liberaban a las personas que Daki había atrapado y almacenado vivas dentro del tejido.

—¡Están atrapadas dentro de la tela! —gritó Aoi—. ¡Si cortamos a ciegas, mataremos a los rehenes!

Ren reaccionó al instante. Usó la guarda de llama de Kenshin para desviar los cortes de seda sin usar el filo de su katana, mientras protegía a las mujeres liberadas.

De pronto, la pared trasera del edificio colapsó cuando la caja de madera de Miyuki saltó a la acción."""
                }
            ]
        },
        {
            "id": "chap3",
            "number": "Capítulo 3",
            "title": "El Doble Estirpe",
            "scenes": [
                {
                    "num": "Escena 1: La Fuerza de Miyuki",
                    "text": """Miyuki se lanzó directo contra Daki. Con una patada impregnada de llamas púrpuras, destruyó la faja obi principal y arrojó a la matrona demonio a través del techo hacia los tejados del distrito.

En los tejados, bajo la luz de la luna llena, la batalla se volvió abierta.

Daki gritó de rabia al ver su rostro arañado por las garras de Miyuki. La herida en su mejilla comenzó a arder por la sangre de la joven demonio.

—¡Monstruo asqueroso! ¡Arruinaste mi rostro! —chilló Daki, mientras su cuerpo comenzaba a cambiar de forma.

Ren saltó al tejado al lado de su hermana, ejecutando el primer movimiento de la Danza del Sol. Su katana se envolvió en una llamarada de fuego dorado brillante que decapitó a Daki de un solo tajo impecable.

La cabeza de la matrona rodó por las tejas de madera."""
                },
                {
                    "num": "Escena 2: El Segundo Hermano",
                    "text": """Pero la cabeza de Daki no se convirtió en cenizas. Comenzó a llorar ruidosamente como una niña pequeña.

—¡Hermano! ¡Me cortaron la cabeza! ¡Ayúdame! —gritó Daki hacia la sombra de su propio cuerpo decapitado.

De la espalda del cuerpo de Daki emergió una figura esquelética y aterradora: un hombre encorvado de piel verdosa que sostenía dos hoces de hueso negro bañadas en veneno mortal. Era Gyutaro, el verdadero núcleo del Quinto Estirpe de Sangre.

—Nadie hace llorar a mi hermana pequeña... —gruñó Gyutaro con una voz profunda y rasposa—. ¿Crees que por tener esa marca en la cara eres especial, cazador?

Gyutaro cruzó sus hoces de hueso, desatando una ráfaga de hoces de sangre venenosa que destruyó tres manzanas del distrito de un solo golpe."""
                },
                {
                    "num": "Escena 3: La Regla de las Dos Cabezas",
                    "text": """Aoi y el muchacho del haori amarillo llegaron al tejado para reunirse con Ren.

—¡Tienen que cortar la cabeza de los dos hermanos al mismo tiempo! —advirtió Aoi, analizando la estructura del arte demoníaco—. Si solo decapitan a uno, ninguno morirá.

Gyutaro se movió a una velocidad que superaba la percepción humana. Clavó una de sus hoces en el hombro de Ren, inyectando un veneno letal que comenzó a volver negras las venas de su cuello.

Ren cayó de rodillas sobre las tejas, sintiendo cómo el veneno paralizaba su sistema nervioso.

—Te quedan cinco minutos de vida, muchacho —se burló Gyutaro, alzando la hoz para el golpe final."""
                }
            ]
        },
        {
            "id": "chap4",
            "number": "Capítulo 4",
            "title": "La Purificación de la Sangre",
            "scenes": [
                {
                    "num": "Escena 1: El Veneno Frenado",
                    "text": """Antes de que la hoz de Gyutaro alcanzara el cuello de Ren, Miyuki se interpuso.

La joven agarró la hoja de hueso con las manos desnudas. Su propia sangre demoníaca brotó de las palmas y quemó el veneno de Gyutaro con sus llamas púrpuras. Luego, Miyuki tocó la herida del hombro de Ren: sus llamas purificaron el veneno del cuerpo de su hermano en cuestión de segundos, restaurando su movilidad.

—Miyuki... gracias —dijo Ren, poniéndose de pie con renovada energía.

El chico del haori amarillo, dormido en su estado de trance del trueno, se encargó de mantener a Daki aislada en los tejados vecinos, evitando que los dos hermanos demonio pudieran reunirse.

Aoi aplicó su veneno de glicina en las hoces de Gyutaro para ralentizar su regeneración.

El contraataque de los cazadores había comenzado."""
                },
                {
                    "num": "Escena 2: La Danza del Fuego Solar",
                    "text": """Ren ajustó la postura de sus pies sobre las tejas de madera. Inhaló aire según la secuencia de los Diarios del Sol, coordinando su ritmo cardíaco con los doce movimientos rituales de la herrería.

Su katana negra no emitió fuego común: la hoja se tiñó de una luz blanca y dorada tan intensa que iluminó todo el Distrito de las Luces como si fuera pleno mediodía.

—Danza del Sol... Cuarta Postura: arco del Sol Poniente.

Ren se convirtió en un fénix de luz dorada. Esquivó las ráfagas de hoces de sangre de Gyutaro, cortando los brazos del demonio uno por uno mientras avanzaba hacia su cuello.

Al mismo tiempo, en el tejado vecino, el chico del haori amarillo ejecutó su Destello del Trueno para cortar la cabeza de Daki por segunda vez."""
                },
                {
                    "num": "Escena 3: La Decapitamiento Simultáneo",
                    "text": """Las dos katanas alcanzaron sus objetivos al mismo milisegundo.

La hoja dorada de Ren atravesó el cuello de Gyutaro, mientras las katanas de trueno cortaban el cuello de Daki. El impacto doble generó una explosión de luz solar que arrasó la estructura principal de la casa de entretenimientos, reduciéndola a escombros de madera y cenizas.

Las cabezas de los dos hermanos demonio rodaron juntas sobre la calle de piedra.

Esta vez, sus cuerpos comenzaron a desintegrarse en cenizas rojas que el viento de la noche dispersó sobre los canales de agua.

Por primera vez en más de cien años, un Estirpe de Sangre de los seis superiores había sido destruido por el Gremio Cuervo."""
                }
            ]
        },
        {
            "id": "chap5",
            "number": "Capítulo 5",
            "title": "Las Cenizas de la Misericordia (Clímax del Volumen 4)",
            "scenes": [
                {
                    "num": "Escena 1: Las Almas del Pasado",
                    "text": """Antes de que las cabezas de Gyutaro y Daki se convirtieran en polvo, sus espíritus permanecieron flotando en la penumbra de la calle abandonada.

Daki, reducida a la forma de la niña humana que solía ser antes de convertirse en demonio, lloraba desconsoladamente, culpando a su hermano por haberlos llevado a la pobreza en su vida humana.

—¡Te odio! ¡Ojalá nunca hubieras sido mi hermano! —chilló la niña.

Gyutaro, con el corazón roto, bajó la mirada.

—Tiene razón... si no hubiera nacido yo, ella habría tenido una vida próspera —susurró el demonio, listo para caminar solo hacia el infierno.

Ren se acercó despacio a las cenizas y se arrodilló entre ellos."""
                },
                {
                    "num": "Escena 2: El Abrazo en el Fuego",
                    "text": """—No mientas —dijo Ren con amabilidad—. Los hermanos no se odian. Sé lo que es querer proteger a tu hermana a cualquier costo. Ella te ama, solo tiene miedo.

Daki se congeló al escuchar las palabras de Ren. Corrió hacia Gyutaro y se abrazó a su cuello, llorando sobre su pecho.

—¡Lo siento, hermano! ¡No me dejes sola! ¡Iré contigo a donde sea, incluso al infierno! —sollozó la niña.

Gyutaro la cargó en su espalda, tal como lo hacía cuando eran humanos, y ambos caminaron juntos hacia las llamas del juicio espiritual, sonriendo en paz por primera vez en siglos.

Miyuki observó la escena desde la sombra, tomando la mano de Ren con ternura."""
                },
                {
                    "num": "Escena 3: La Ira de Kageyama (Cierre del Tomo 4)",
                    "img": "escena_climax.jpg",
                    "text": """En las profundidades del Castillo Infinito —un palacio flotante fuera de la dimensión humana—, la noticia de la caída del Quinto Estirpe de Sangre sacudió la corte demoníaca.

Kageyama, el Rey Oni, en su forma de hombre elegante con ojos de gato carmesí, destruyó su laboratorio de alquimia de un solo golpe de rabia.

Alrededor de su trono, los restantes cuatro Estirpes de Sangre Superiores se arrodillaron en pánico.

—Los cazadores han despertado la Marca del Sol... —dijo Kageyama con una voz helada que hizo temblar la dimensión del castillo—. No habrá más cacerías individuales. Desplieguen a los tres Demonios del Abismo. Destruyan la Aldea de los Herreros y maten al chico de la marca.

[ CONTINUARÁ EN EL VOLUMEN 5 ]"""
                }
            ]
        }
    ]
    
    bkm_id = 1
    for chap in content:
        p_c = doc.add_paragraph()
        add_bookmark(p_c, chap["id"], bkm_id)
        
        r_cnum = p_c.add_run(f"{chap['number']}\n")
        r_cnum.font.name = "Arial"
        r_cnum.font.size = Pt(14)
        r_cnum.font.bold = True
        r_cnum.font.color.rgb = RGBColor(185, 28, 28)
        
        r_ctitle = p_c.add_run(chap["title"])
        r_ctitle.font.name = "Arial"
        r_ctitle.font.size = Pt(20)
        r_ctitle.font.bold = True
        r_ctitle.font.color.rgb = RGBColor(220, 38, 38)
        
        end_bookmark(p_c, bkm_id)
        bkm_id += 1
        
        for sc in chap["scenes"]:
            p_s = doc.add_paragraph()
            r_snum = p_s.add_run(f"\n◆ {sc['num']}\n")
            r_snum.font.name = "Arial"
            r_snum.font.size = Pt(13)
            r_snum.font.bold = True
            r_snum.font.color.rgb = RGBColor(71, 85, 105)
            
            p_stext = doc.add_paragraph()
            r_stext = p_stext.add_run(sc["text"])
            r_stext.font.name = "Georgia"
            r_stext.font.size = Pt(11)
            
            if "img" in sc:
                img_p = os.path.join(target_root, sc["img"])
                if os.path.exists(img_p):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_p, width=Inches(5))
                    
        doc.add_page_break()
        
    doc_dest_root = os.path.join(target_root, "libro.docx")
    doc_dest_sys = os.path.join(target_sys, "libro.docx")
    doc.save(doc_dest_root)
    shutil.copy2(doc_dest_root, doc_dest_sys)
    print(f"Generated docx at {doc_dest_root}")

create_oni_vol4_docx()

# Generate ficha_producto.json
ficha_data = {
    "titulo": "Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre) - Volumen 4: El Distrito de los Espejos y la Mariposa de la Sombra",
    "autor": "Nicolás Noguera",
    "precio": 20.00,
    "moneda": "USD",
    "genero": "Manga / Light Novel / Fantasía Oscura / Acción Sengoku",
    "headline": "El descubrimiento de Los Diarios del Sol. La infiltración en el Distrito de las Luces. Y la histórica victoria contra el Quinto Estirpe de Sangre.",
    "descripcion": "En la biblioteca secreta de la Finca del Fuego, Ren descubre la verdad tras la Estilo de Esgrima de su padre: la legendaria Danza del Sol. Enviado junto a Aoi al deslumbrante Distrito de las Luces, los cazadores enfrentan a Daki y Gyutaro, los hermanos que componen el Quinto Estirpe de Sangre. Una batalla en los tejados a doble decapitamiento simultáneo que provocará la ira desmedida del Rey Oni Kageyama en el Castillo Infinito.",
    "beneficios": [
        "Manuscrito oficial ilustrado completo en formato .docx listo para eReaders y Amazon KDP.",
        "Ilustraciones de alta definición en estética anime/manga dark fantasy.",
        "Cuarta entrega llena de acción y revelaciones históricas de la saga de Nicolás Noguera."
    ],
    "capitulos": [
        "Capítulo 1: La Finca del Fuego",
        "Capítulo 2: La Ciudad que Nunca Duerme",
        "Capítulo 3: El Doble Estirpe",
        "Capítulo 4: La Purificación de la Sangre",
        "Capítulo 5: Las Cenizas de la Misericordia"
    ]
}

with open(os.path.join(target_root, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)
with open(os.path.join(target_sys, "ficha_producto.json"), "w", encoding="utf-8") as f:
    json.dump(ficha_data, f, indent=2, ensure_ascii=False)

# Generate index.html landing page for Oni no Ketsuryu Vol 4
html_content = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Oni no Ketsuryū Vol 4 - Nicolás Noguera | Tienda Oficial</title>
    <style>
        :root {
            --bg-color: #0b0f19;
            --card-bg: #111827;
            --accent: #dc2626;
            --accent-hover: #b91c1c;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
        }
        body {
            font-family: 'Segoe UI', system-ui, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 0;
            line-height: 1.6;
        }
        .header-banner {
            width: 100%;
            max-height: 400px;
            object-fit: cover;
            border-bottom: 3px solid var(--accent);
        }
        .container {
            max-width: 1000px;
            margin: 0 auto;
            padding: 40px 20px;
        }
        .product-grid {
            display: grid;
            grid-template-columns: 1fr 2fr;
            gap: 40px;
            margin-bottom: 50px;
        }
        @media (max-width: 768px) {
            .product-grid { grid-template-columns: 1fr; }
        }
        .cover-img {
            width: 100%;
            border-radius: 12px;
            box-shadow: 0 10px 30px rgba(220, 38, 38, 0.3);
            border: 1px solid rgba(220, 38, 38, 0.3);
        }
        .badge {
            background: var(--accent);
            color: white;
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            display: inline-block;
            margin-bottom: 15px;
        }
        h1 { font-size: 2.2rem; margin: 10px 0; color: #fff; }
        .author { color: var(--accent); font-weight: 600; font-size: 1.1rem; margin-bottom: 20px; }
        .price-tag {
            font-size: 2rem;
            font-weight: 800;
            color: #fff;
            margin: 20px 0;
        }
        .price-tag span { font-size: 1rem; color: var(--text-muted); }
        .buy-btn {
            background: linear-gradient(135deg, #dc2626, #991b1b);
            color: white;
            text-decoration: none;
            padding: 16px 36px;
            border-radius: 8px;
            font-weight: bold;
            font-size: 1.2rem;
            display: inline-block;
            box-shadow: 0 6px 20px rgba(220, 38, 38, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
        }
        .buy-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(220, 38, 38, 0.6);
        }
        .section-title {
            font-size: 1.5rem;
            border-left: 4px solid var(--accent);
            padding-left: 12px;
            margin: 40px 0 20px;
            color: #fff;
        }
        .chapters-list {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 25px;
            border: 1px solid rgba(255,255,255,0.05);
        }
        .chapters-list li {
            margin-bottom: 12px;
            color: var(--text-main);
        }
        .gallery-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin-top: 20px;
        }
        .gallery-grid img {
            width: 100%;
            border-radius: 8px;
            border: 1px solid rgba(255,255,255,0.1);
        }
    </style>
</head>
<body>
    <img src="banner.jpg" alt="Banner Oni no Ketsuryu Vol 4" class="header-banner">
    
    <div class="container">
        <div class="product-grid">
            <div>
                <img src="portada.jpg" alt="Portada Oni no Ketsuryu Vol 4" class="cover-img">
            </div>
            <div>
                <span class="badge">MANGA / LIGHT NOVEL</span>
                <h1>Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)</h1>
                <div class="author">Volumen 4: El Distrito de los Espejos • Por Nicolás Noguera</div>
                <p>El descubrimiento de Los Diarios del Sol. La infiltración en el Distrito de las Luces. Y la histórica victoria contra el Quinto Estirpe de Sangre.</p>
                <p>En la biblioteca secreta de la Finca del Fuego, Ren descubre la verdad tras la Estilo de Esgrima de su padre: la legendaria Danza del Sol. Enviado junto a Aoi al deslumbrante Distrito de las Luces, los cazadores enfrentan a Daki y Gyutaro en una batalla a doble decapitamiento simultáneo.</p>
                
                <div class="price-tag">$20.00 <span>USD</span></div>
                <a href="#" class="buy-btn">COMPRAR AHORA ($20 USD)</a>
            </div>
        </div>

        <h2 class="section-title">Contenido del Volumen 4</h2>
        <div class="chapters-list">
            <ul>
                <li><strong>Capítulo 1:</strong> La Finca del Fuego (Escenas 1-3)</li>
                <li><strong>Capítulo 2:</strong> La Ciudad que Nunca Duerme (Escenas 1-3)</li>
                <li><strong>Capítulo 3:</strong> El Doble Estirpe (Escenas 1-3)</li>
                <li><strong>Capítulo 4:</strong> La Purificación de la Sangre (Escenas 1-3)</li>
                <li><strong>Capítulo 5:</strong> Las Cenizas de la Misericordia (Clímax del Volumen 4 - Escenas 1-3)</li>
            </ul>
        </div>

        <h2 class="section-title">Ilustraciones Interiores Destacadas</h2>
        <div class="gallery-grid">
            <img src="escena_1.jpg" alt="La Danza del Sol">
            <img src="escena_climax.jpg" alt="La Ira de Kageyama en el Castillo Infinito">
        </div>
    </div>
</body>
</html>
"""

with open(os.path.join(target_root, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)
with open(os.path.join(target_sys, "index.html"), "w", encoding="utf-8") as f:
    f.write(html_content)

print("Generated HTML landing pages for Vol 4 successfully")
