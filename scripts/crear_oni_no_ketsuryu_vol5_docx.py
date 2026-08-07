# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = docx.Document()

# Margenes
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

# Titulo Principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("Oni no Ketsuryū (鬼の血流 - La Estirpe de la Sangre)\nVolumen 5: La Aldea de los Herreros Olvidados")
run_title.font.name = "Arial"
run_title.font.size = Pt(24)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(180, 0, 0)

# Subtitulo
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Obra Oficial por Nicolás Noguera\nEdición Digital Ilustrada")
run_sub.font.name = "Arial"
run_sub.font.size = Pt(14)
run_sub.font.italic = True

doc.add_page_break()

base_dir = Path(r"C:\Proyectos\mis-libros-editorial\libros\oni-no-ketsuryu-volumen-5")

# Portada
portada_path = base_dir / "portada.jpg"
if portada_path.exists():
    doc.add_picture(str(portada_path), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

# Texto del libro completo
capitulos = [
    {
        "titulo": "Capítulo 1: El Valle Oculto de la Forja",
        "escenas": [
            {
                "sub": "Escena 1: El Viaje a Ciegas",
                "text": "Tras la batalla en el Distrito de las Luces, la katana negra de Ren quedó severamente agrietada debido a las temperaturas extremas de la Danza del Sol. Al ser un arma con una aleación única, ningún artesano común podía repararla.\n\nPara proteger la ubicación del secreto más valioso del Gremio Cuervo, Ren tuvo que hacer el viaje hacia la Aldea de los Herreros con los ojos vendados, tapones en los oídos y transportado en relevos por diferentes mensajeros del gremio que cambiaban cada tres kilómetros.\n\nAl retirarse la venda, Ren se encontró en un valle secreto escondido entre profundas montañas volcánicas, lleno de aguas termales y chozas de madera con chimeneas humeantes.\n\nTodos los habitantes de la aldea vestían máscaras ceremoniales de Hyottoko para proteger sus identidades.",
                "img": "escena_c1_e1.jpg"
            },
            {
                "sub": "Escena 2: El Jefe de la Aldea y la Marioneta de Entrenamiento",
                "text": "Ren fue recibido por Tecchin, el anciano jefe de la aldea, un hombre de baja estatura pero de aura autoritaria.\n\n—La katana que llevas no está rota por descuido —dijo Tecchin, examinando las muescas de la hoja negra—. La sometiste al calor de la respiración original. Para que esta espada sostenga la Danza del Sol, debes entrenar tu velocidad con el mecanismo ancestral de la aldea.\n\nTecchin lo llevó al centro del bosque de bambú, donde se encontraba Yoriichi Modelo Zeroth: una marioneta de entrenamiento mecánica de seis brazos construida hace más de cuatrocientos años por los fundadores de la herrería.\n\nEl mecanismo se activó con un engranaje seco, desenvainando seis katanas de madera a una velocidad sobrehumana.",
                "img": "escena_c1_e2.jpg"
            },
            {
                "sub": "Escena 3: La Katana de Cuatrocientos Años",
                "text": "Ren entrenó sin descanso durante tres días contra la marioneta de seis brazos, usando la Respiración de Sangre para afinar sus reflejos hasta adelantarse a los movimientos mecánicos del autómata.\n\nAl dar el golpe final que destruyó la armadura de madera de la marioneta, el torso del autómata se abrió por la mitad.\n\nEn su interior, escondida dentro del mecanismo central desde la era Sengoku, descansaba una katana antigua sin empuñadura. A pesar de haber pasado cuatro siglos encerrada, la hoja de acero oscuro no tenía una sola mancha de óxido y emitía un calor sutil que hizo vibrar la marca de la cara de Ren.\n\n—La espada del primer usuario de la Respiración del Sol... —susurró el jefe de la aldea con veneración.",
                "img": "escena_c1_e3.jpg"
            }
        ]
    },
    {
        "titulo": "Capítulo 2: La Invasión del Cuarto y Segundo Lunar",
        "escenas": [
            {
                "sub": "Escena 1: El Ataque a las Nieblas",
                "text": "Mientras el maestro herrero Haganezuka comenzaba el proceso de pulido de tres días para adaptar la antigua espada al cuerpo de Ren, el cielo sobre el valle volcánico se volvió rojo sangre.\n\nLas barreras mágicas de glicina que ocultaban la aldea fueron destruidas desde el aire.\n\nDesde las nubes descendieron dos sombras colosales: el Cuarto Lunar Rojo (un demonio anciano y decrepito llamado Hantengu que flotaba sobre una vasija) y el Segundo Lunar Rojo (Doma, un demonio de apariencia aristocrática con ojos multicolores y abanicos de cristal que sonreía con falsa amabilidad).\n\n—Qué lugar tan pintoresco... —dijo Doma, desplegando su abanico de cristal—. Mi Señor desea que la forja de las katanas del Gremio sea reducida a cenizas hoy mismo.",
                "img": "escena_c2_e1.jpg"
            },
            {
                "sub": "Escena 2: El Abanico de Hielo",
                "text": "Doma agitó su abanico de cristal congelado. Una ventisca helada a temperatura bajo cero arrasó las chozas de madera del sector norte de la aldea, congelando instantáneamente las aguas termales y a varios herreros que intentaban proteger sus talleres.\n\nRen corrió al combate llevando la caja de Miyuki en su espalda, desenvainando la katana de la marioneta aún sin pulir por completo.\n\nUn rayo de niebla blanca cortó el ataque de hielo de Doma.\n\nEl Pilar de la Niebla —Muichiro, un joven espadachín prodigio de catorce años con largo cabello negro y verde— se interpuso en el camino, desenvainando su katana con una calma absoluta.\n\n—Ren... llévate al demonio anciano lejos de la zona de las forjas —ordenó Muichiro sin cambiar la expresión de su rostro—. Yo me encargaré del hombre de los abanicos.",
                "img": "escena_c2_e2.jpg"
            },
            {
                "sub": "Escena 3: Las Seis Emociones de la Ira",
                "text": "Ren persiguió al anciano Hantengu hacia el bosque contiguo a las termas. Al acorralar al demonio cobarde y cortar su cabeza con un tajo rápido, ocurrió algo insólito:\n\nEl cuerpo decapitado y la cabeza no murieron. Se dividieron en dos demonios jóvenes independientes, representando la Ira (un guerrero con un abanico de hoja que generaba tornados) y la Alegría (un demonio con alas que emitía ondas de sonido destructivas).\n\nAl cortar a esos dos, se dividieron nuevamente en cuatro entidades representando la Tristeza y el Placer.\n\n—¡Cada vez que lo decapito se multiplica y se vuelve más fuerte! —comprendió Ren, esquivando una ráfaga de truenos que destruyó los árboles a su alrededor.",
                "img": "escena_c2_e3.jpg"
            }
        ]
    },
    {
        "titulo": "Capítulo 3: La Batalla en la Niebla y la Madera",
        "escenas": [
            {
                "sub": "Escena 1: El Dominio de la Niebla",
                "text": "En la aldea congelada, la batalla entre el Pilar de la Niebla y el Segundo Lunar Rojo alcanzó una escala devastadora. Doma creó varios clones de hielo en forma de estatuas de Buda que lanzaban escarcha mortal capaz de congelar los pulmones de cualquier cazador que respirara el aire.\n\nMuichiro inhaló profundamente, activando la marca de la niebla en su frente con forma de nubes oscuras.\n\n—Respiración de la Niebla... Séptima Postura: Nubes Cambiantes.\n\nMuichiro desapareció en una ilusión de niebla densa. Su velocidad cambió de ritmo constantemente, confundiendo la percepción espacial de Doma. Con un movimiento rápido como un suspiro, Muichiro cortó los brazos de cristal del Segundo Lunar y destruyó sus clones de hielo.\n\nDoma sonrió con fascinación mientras se regeneraba al instante.\n\n—Eres muy rápido, muchacho... pero tus pulmones ya han absorbido mi veneno helado.",
                "img": "escena_c3_e1.jpg"
            },
            {
                "sub": "Escena 2: El Dragón de Madera de las Cinco Emociones",
                "text": "Mientras tanto, en el bosque, las cuatro emociones de Hantengu se fusionaron en un solo ser definitivo: Zohakuten, un demonio con aspecto de niño guerrero rodado de tambores tradicionales flotantes.\n\nZohakuten golpeó los tambores de su espalda, invocando gigantescos dragones de madera de cinco cabezas que emergieron de la tierra para aplastar a Ren y a Miyuki.\n\nMiyuki usó sus garras para cortar la madera, pero los dragones se regeneraban usando la energía del bosque volcánico.\n\nRen fue acorralado contra un barranco de piedra por las mordeduras de los dragones de madera, sintiendo que sus piernas no respondían por el agotamiento de la Respiración de Sangre.\n\n—Ustedes, los cazadores, son los verdaderos malvados... —dijo Zohakuten con una voz infantil y desprovista de empatía—. Atacan a un anciano débil e indefenso... no tienen piedad.",
                "img": "escena_c3_e2.jpg"
            },
            {
                "sub": "Escena 3: La Llegada del Pilar del Amor",
                "text": "Cuando una de las cabezas de madera estaba a milímetros de tragar a Ren, un látigo de acero flexible y rosado cortó la madera en cien pedazos en una fracción de segundo.\n\nMitsuri, la Pilar del Amor —una mujer de gran agilidad que vestía un haori blanco y llevaba una katana-látigo ultraligera—, saltó al barranco con una sonrisa confiada.\n\n—¡Lamento la demora! —exclamó Mitsuri, ejecutando giros acrobáticos en el aire que destruyeron a los restantes dragones de madera—. ¡Nadie toca a los aprendices del gremio mientras yo esté aquí!\n\nCon la llegada de Mitsuri, la batalla en el bosque tomó un segundo aire.",
                "img": "escena_c3_e3.jpg"
            }
        ]
    },
    {
        "titulo": "Capítulo 4: La Hoja del Fuego Solar",
        "escenas": [
            {
                "sub": "Escena 1: El Verdadero Núcleo Oculto",
                "text": "Mitsuri mantuvo a Zohakuten ocupado en un combate acrobático de alta velocidad. Mientras sus látigos de acero cortaban los ataques de madera y rayo, Ren usó su olfato desarrollado para buscar la verdadera ubicación del demonio original.\n\n—Zohakuten solo es una proyección... —comprendió Ren—. El cuerpo real del anciano Hantengu sigue siendo del tamaño de un ratón y está escondido dentro del corazón del dragón de madera principal.\n\nRen corrió hacia el dragón central, pero la katana de la marioneta comenzó a agrietarse bajo la presión del fuego solar.\n\nEn ese instante, el maestro herrero Haganezuka emergió de un taller en llamas, corriendo hacia Ren mientras esquivaba los escombros y sosteniendo la katana antigua pulida con su acabado definitivo.\n\n—¡Tómala, Ren! —gritó Haganezuka, arrojando el arma al aire—. ¡Esta es la Katana del Sol de la Era Sengoku!",
                "img": "escena_c4_e1.jpg"
            },
            {
                "sub": "Escena 2: La Transformación de la Espada",
                "text": "Ren atrapó la nueva katana en pleno vuelo.\n\nAl contacto con la marca de su cara y la Respiración de Sangre, el acero oscuro absorbió el calor de su cuerpo. La hoja no solo se volvió incandescente: el grabado antiguo en la base de la espada reveló el carácter kanji de 'Destruir' (Metsu).\n\nMiyuki corrió al lado de su hermano y sujetó el filo de la espada con sus manos desnudas, bañando la hoja con su propia sangre de demonio.\n\nEl fuego dorado de la Danza del Sol y las llamas púrpuras de Miyuki se combinaron en una sola llamarada bicolor que iluminó todo el bosque de la aldea.\n\n—Danza del Sol Combinada... Décima Postura: Dragón del Fuego Solar.\n\nRen se desplazó como una ráfaga de luz que atravesó el cuerpo del gigante de madera en línea recta.",
                "img": "escena_c4_e2.jpg"
            },
            {
                "sub": "Escena 3: El Corte del Duende",
                "text": "La hoja de llamas doradas y púrpuras cortó la cabeza del pequeño Hantengu escondido dentro del corazón del dragón de madera.\n\nEl Cuarto Lunar Rojo emitió un alarido de agonía mientras su cuerpo gigante de madera y sus cuatro emociones se desintegraban simultáneamente en cenizas brillantes.\n\nEn el otro lado de la aldea, al sentir la muerte de Hantengu, el Segundo Lunar Rojo (Doma) cubrió su rostro con sus abanicos de cristal y sonrió con falsa cortesía.\n\n—Vaya... parece que mi compañero ha caído —dijo Doma, creando una barrera de viento helado para retirarse—. No tiene sentido seguir aquí solo. Nos veremos en la cacería final, pequeño cazador.\n\nDoma desapareció en un portal de humo negro, dejando la aldea en un silencio paulatino mientras los primeros rayos del amanecer comenzaban a cruzar las montañas volcánicas.",
                "img": "escena_c4_e3.jpg"
            }
        ]
    },
    {
        "titulo": "Capítulo 5: El Milagro bajo el Sol (Clímax del Volumen 5)",
        "escenas": [
            {
                "sub": "Escena 1: El Sol en el Bosque",
                "text": "El sol de la mañana iluminó por completo las ruinas de la Aldea de los Herreros.\n\nRen cayó sobre sus rodillas, apoyando la katana del sol en la hierba húmeda, exhausto por el gasto de energía de la combinación de respiraciones.\n\nDe pronto, un pensamiento de terror puro cruzó su mente.\n\n—¡Miyuki! —gritó Ren, girándose hacia el bosque.\n\nMiyuki permanecía de pie en medio del claro. La luz directa del sol matutino bañaba su cuerpo por completo. Ren corrió desesperado con su capa en las manos para cubrirla antes de que los rayos solares la quemaran y la convirtieran en cenizas como a cualquier otro demonio.\n\nPero cuando Ren llegó a su lado, la capa se cayó de sus manos por la impresión.",
                "img": "escena_c5_e1.jpg"
            },
            {
                "sub": "Escena 2: La Conquista del Sol",
                "text": "Miyuki no se estaba disolviendo.\n\nLos pequeños cuernos de hueso de su frente se retrayeron bajo su piel. Sus pupilas rasgadas de bestia volvieron a dilatarse, tomando la forma de los ojos humanos y cálidos de su infancia.\n\nMiyuki se retiró despacio el bambú de la boca con su propia mano y miró a Ren a los ojos.\n\n—Buenos... días... hermano... —dijo Miyuki con una voz suave que Ren no había escuchado en más de tres años.\n\nMiyuki se había convertido en la primera entidad demoníaca en la historia en conquistar la luz del sol, manteniendo su cuerpo inmune a los rayos solares mientras conservaba sus habilidades sobrehumanas.\n\nRen la abrazó llorando con fuerza sobre la hierba, mientras el Pilar de la Niebla y la Pilar del Amor observaban la escena con lágrimas en los ojos.",
                "img": "escena_c5_e2.jpg"
            },
            {
                "sub": "Escena 3: La Declaración de Guerra (Cierre del Tomo 5)",
                "text": "A miles de kilómetros de distancia, dentro del Castillo Infinito, Muzan sintió la conquista del sol en la sangre de Miyuki a través del vínculo demoníaco.\n\nEl Rey Oni dejó caer la copa de cristal de su mano, y por primera vez en mil años, sus ojos carmesí mostraron una emoción que nunca antes había sentido: obsesión pura.\n\n—Ya no necesito buscar la Lirios del Sol Azul... —dijo Muzan, mientras una sonrisa macabra se dibujaba en su rostro—. Si devoro a esa niña y absorbo su sangre, me convertiré en un ser inmune al sol. ¡Seré el Dios perfecto de este mundo!\n\nMuzan dio la orden final a todos los demonios restantes del planeta:\n\n—¡Reúnan a las tropas! ¡Ataquen la sede central del Gremio Cuervo! ¡Tráiganme a la chica del sol!\n\n\n                  [ CONTINUARÁ EN EL VOLUMEN 6 ]\n                  [ INICIO DEL ARCO DE LA GUERRA FINAL ]",
                "img": "escena_c5_e3.jpg"
            }
        ]
    }
]

for cap in capitulos:
    p_c = doc.add_paragraph()
    run_c = p_c.add_run(cap["titulo"])
    run_c.font.name = "Arial"
    run_c.font.size = Pt(18)
    run_c.font.bold = True
    run_c.font.color.rgb = RGBColor(160, 0, 0)
    
    for esc in cap["escenas"]:
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(esc["sub"])
        run_sub.font.name = "Arial"
        run_sub.font.size = Pt(14)
        run_sub.font.bold = True
        run_sub.font.color.rgb = RGBColor(50, 50, 50)
        
        p_t = doc.add_paragraph()
        run_t = p_t.add_run(esc["text"])
        run_t.font.name = "Calibri"
        run_t.font.size = Pt(12)
        
        img_path = base_dir / esc["img"]
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
    doc.add_page_break()

docx_out1 = base_dir / "libro.docx"
docx_out2 = Path(r"c:\Users\nicol\Downloads\MIS LIBROS\libros\oni-no-ketsuryu-volumen-5\libro.docx")

doc.save(str(docx_out1))
doc.save(str(docx_out2))
print(f"Generated libro.docx successfully at {docx_out1}")
